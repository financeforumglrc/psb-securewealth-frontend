#!/usr/bin/env python3
"""Generate balanced 10-min pitch script: Kunal reduced, others get more."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

OUT_PATH = "PSB_Hackathon_10Min_Balanced_Pitch_Script.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    return h


def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_line(doc, speaker, text, demo=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p.add_run(f"{speaker}: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    if demo:
        r3 = p.add_run(f"\n    [LIVE DEMO: {demo}]")
        r3.bold = True
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x00, 0x66, 0x33)
        r3.font.size = Pt(10)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(f"💡 Tip: {text}")
    r.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
    r.font.size = Pt(10)
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PSB SecureWealth Twin\n10-Minute Balanced Pitch Script")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    t.paragraph_format.space_after = Pt(12)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Kunal Load Reduced | Rikshita, Deepanshu & Mrigesh Get More Time")
    r.font.size = Pt(14)
    r.italic = True
    st.paragraph_format.space_after = Pt(18)

    add_para(doc,
        "This version balances speaking time. Kunal handles one core solution, one innovation, and the closing. "
        "Rikshita, Deepanshu, and Mrigesh carry the problem, future scope, credibility, and multiple innovations. "
        "Live demo cues are in [LIVE DEMO: ...] brackets.",
        italic=True, space_after=12)

    add_heading(doc, "Speaker Time Distribution", level=1)
    add_para(doc, "Rikshita Barua: 2 min 35 sec (Opening + Problem + Innovation 1 + Future Scope)")
    add_para(doc, "Deepanshu Sharma: 1 min 40 sec (Solution 1 + Innovation 2 + Credibility)")
    add_para(doc, "Mrigesh Mohanty: 1 min 40 sec (Solution 2 + Innovation 3 + Credibility)")
    add_para(doc, "Ishita Anand: 1 min 30 sec (Solution 3 + Innovation 4)")
    add_para(doc, "Tripti Jain: 45 sec (Solution 4)")
    add_para(doc, "Kunal Saxena: 2 min (Solution 5 + Innovation 5 + Closing)")
    add_para(doc, "Total: 10 minutes")

    doc.add_page_break()

    add_heading(doc, "0:00 – 1:15 | Opening + Problem Statement", level=2)
    add_line(doc, "Rikshita", "Good morning judges. I am Rikshita, and we are Team DS Financial Solutions.")
    add_line(doc, "Rikshita", "Today, we are not showing just a website. We are selling a product — a product that can change how Punjab & Sind Bank serves every customer.")
    add_line(doc, "Rikshita", "Our product is PSB SecureWealth Twin. Our tagline is: One Bank. One Twin. Infinite Possibilities.")
    add_line(doc, "Rikshita", "Let me start with the problem. Today, an average Indian customer has money in many places — savings account, FD, loan, insurance, mutual funds, gold. But when they open their bank app, they only see one savings balance.")
    add_line(doc, "Rikshita", "Because of this, three big problems happen. First, idle money sleeps in savings accounts and loses value. Second, people miss tax-saving deadlines. Third, fraudsters take advantage of this confusion through fake calls and phishing.")
    add_line(doc, "Rikshita", "So the real problem is: banks today store money, but they do not actively grow, protect, and simplify their customers' financial lives. Let me hand over to Deepanshu for the first solution.")
    add_note(doc, "Rikshita sets the energy. Speak slowly and make eye contact with all judges.")

    doc.add_page_break()

    add_heading(doc, "1:15 – 2:00 | Solution 1: Unified Wealth Twin", level=2)
    add_line(doc, "Deepanshu", "Thank you, Rikshita. I am Deepanshu. Let me show you how we solve the scattered-money problem.",
             demo="Open psb-securewealth-frontend.onrender.com. Click 'Demo Mode'. Select 'Neha — Salaried Professional'. Wait for the Account Aggregator animation.")
    add_line(doc, "Deepanshu", "This is the Wealth Twin Dashboard. On one screen, the customer sees all accounts, FDs, loans, insurance, gold, mutual funds, and real estate.",
             demo="Point to the Net Worth card, Asset Breakdown chart, and Recent Transactions list.")
    add_line(doc, "Deepanshu", "The best part is the AI Action Cards. They do not just show data — they tell the user what to do. For example, here it says: 'You have two lakh rupees idle for thirty days. Sweep it to FD ladder in one tap.'")
    add_line(doc, "Deepanshu", "So our first solution gives the customer a complete financial picture and turns it into action. Mrigesh will now show you how we protect all of this.")

    doc.add_page_break()

    add_heading(doc, "2:00 – 2:45 | Solution 2: Security Beast", level=2)
    add_line(doc, "Mrigesh", "Thanks, Deepanshu. I am Mrigesh. When customers link all their money in one place, security becomes the most important thing.",
             demo="Navigate to the Security Beast or Fraud Protection section from the sidebar.")
    add_line(doc, "Mrigesh", "We built a Security Beast with four layers. Layer one is Rakshak AI. It learns what is normal for each user — when they transact, how much, from which device.")
    add_line(doc, "Mrigesh", "If it sees something unusual, like a transaction at 3 AM from a new city, it blocks it instantly and asks for biometric verification.",
             demo="Show the Rakshak AI risk score card or fraud alert simulation.")
    add_line(doc, "Mrigesh", "Layer two is Deepfake Voice Shield. Many frauds now use cloned voices. Our system checks voice print and liveness before allowing high-risk actions. Layer three is Decoy Account for forced unlock. Layer four is Ghost Mode to hide balances in public.")
    add_line(doc, "Mrigesh", "So our product protects money before it is lost. Ishita will now explain how we help businesses.")

    doc.add_page_break()

    add_heading(doc, "2:45 – 3:30 | Solution 3: SME Centre + Tax + Inclusion", level=2)
    add_line(doc, "Ishita", "Thanks, Mrigesh. I am Ishita. A bank serves not just individuals, but also lakhs of small businesses.",
             demo="Click on 'SME Centre' or 'Business Mode' from the navigation menu.")
    add_line(doc, "Ishita", "Our SME Centre has three tools. First, the Cash Flow Timeline shows monthly inflow and outflow, and highlights months where the business may run short of money.")
    add_line(doc, "Ishita", "Second, Working Capital Health gives a score based on current ratio, quick ratio, and payment speed.",
             demo="Point to the Working Capital Health score and colour-coded ratios.")
    add_line(doc, "Ishita", "Third, the Surplus Fund Advisor recommends FD sweep, liquid fund, or early vendor payment — all with projected savings.")
    add_line(doc, "Ishita", "For retail users, we have an Advanced Tax Centre with Old vs New regime calculator, 80C tracker, and tax deadline calendar. We also have Senior Mode with big fonts and Face Login.")
    add_line(doc, "Ishita", "Tripti will now explain the macro view.")

    doc.add_page_break()

    add_heading(doc, "3:30 – 4:15 | Solution 4: Macro Signal Tower + Compliance", level=2)
    add_line(doc, "Tripti", "Thank you, Ishita. I am Tripti. Most banks give advice using only customer data. We also use the economy.",
             demo="Go back to Dashboard. Scroll to the Macro Signal Tower card.")
    add_line(doc, "Tripti", "Our Macro Signal Tower tracks RBI repo rate, inflation, USD-INR, and gold prices. It then converts these into simple actions.")
    add_line(doc, "Tripti", "For example, if repo rate is rising, it says: shift to floating-rate FD. If inflation is above six percent, it says: reduce debt duration and increase equity SIP. If gold is rising, it says: book partial profit and move to FD.",
             demo="Point to each signal and its recommended action on the card.")
    add_line(doc, "Tripti", "And because we show projections, every screen has a clear disclaimer — this is a simulation, not a guaranteed return. Kunal will now show you proactive intelligence.")

    doc.add_page_break()

    add_heading(doc, "4:15 – 5:00 | Solution 5: AI Action Cards + Smart Sweep", level=2)
    add_line(doc, "Kunal", "Thank you, Tripti. I am Kunal. Let me show you how the product thinks for the customer.",
             demo="Click on 'Smart Sweep' or 'AI Actions' from the dashboard. Show the list of recommendations.")
    add_line(doc, "Kunal", "Smart Sweep scans all linked accounts and finds idle money. Then it recommends the best place to park it — FD ladder, liquid fund, or overnight fund — based on how soon the user needs the money.",
             demo="Click one recommendation. Show the projected interest comparison.")
    add_line(doc, "Kunal", "We also have behavioural nudges. If a user is overspending, it suggests pausing a subscription. If a SIP is due, it reminds them. If tax deadline is near, it prompts 80C investment.")
    add_line(doc, "Kunal", "This makes banking proactive. Rikshita will now take you through our first major innovation.")

    doc.add_page_break()

    add_heading(doc, "5:00 – 8:00 | Innovation Showcase (Shared Equally)", level=2)
    add_line(doc, "Rikshita", "We solved the core problems. Now let me show you the innovations that make this product exciting and different.",
             demo="Stay on the live website. You will demo 5 innovations one by one.")

    add_heading(doc, "5:00 – 5:45 | Innovation 1: Life-Shock Simulator", level=3)
    add_line(doc, "Rikshita", "First, the Life-Shock Simulator. Most people do not plan for emergencies. This tool lets the user test what happens if they lose their job, face a medical emergency, or see a market crash.",
             demo="Navigate to Wealth Twin → Life-Shock Simulator. Select 'Job Loss'. Show impact on net worth and goals. Show recommended action.")
    add_line(doc, "Rikshita", "It shows the impact on net worth and goals, and recommends insurance or contingency fund actions.")

    add_heading(doc, "5:45 – 6:30 | Innovation 2: Generational Wealth Slider", level=3)
    add_line(doc, "Deepanshu", "Second, the Generational Wealth Slider. A user can slide across years and see how their retirement corpus grows, how inflation reduces purchasing power, and whether goals are on track.",
             demo="Navigate to Goals or Wealth Twin. Use the Generational Wealth Slider. Move from 2026 to 2056 and show the projection graph.")
    add_line(doc, "Deepanshu", "It makes long-term planning visual and easy, even for someone who does not understand finance.")

    add_heading(doc, "6:30 – 7:15 | Innovation 3: Deepfake Voice Shield + Decoy Mode", level=3)
    add_line(doc, "Mrigesh", "Third, our security innovations. The Deepfake Voice Shield checks voice print and liveness before any high-risk action. If the voice does not match, the transaction is blocked.",
             demo="Show the Security Beast section. Toggle Decoy Account and Ghost Mode settings. If voice demo is live, show the voice verification screen.")
    add_line(doc, "Mrigesh", "Decoy Account shows a fake low-balance account when a duress PIN is entered. Ghost Mode hides real balances in public places.")

    add_heading(doc, "7:15 – 8:00 | Innovation 4: CreditBridge AI for MSMEs", level=3)
    add_line(doc, "Ishita", "Fourth, CreditBridge AI. Many small businesses do not have formal credit history, so banks reject their loan applications.",
             demo="Navigate to SME Centre → CreditBridge AI. Show the credit health score and loan product recommendation.")
    add_line(doc, "Ishita", "Our AI analyses cash flow, GST turnover, and UPI receipts to create a credit health score and suggest the right PSB loan product.")

    doc.add_page_break()

    add_heading(doc, "8:00 – 8:30 | Innovation 5: More Unique Features", level=3)
    add_line(doc, "Kunal", "Fifth, a quick look at more unique features. NRI Mode for FEMA-aware tracking. Fantasy League of Wealth for family financial literacy. Sovereign Vault for long-term high-security holdings.",
             demo="Quickly show Settings or NRI Mode toggle.")
    add_line(doc, "Kunal", "These features show we have thought about safety, inclusion, and engagement deeply.")

    doc.add_page_break()

    add_heading(doc, "8:30 – 8:50 | Credibility", level=2)
    add_line(doc, "Deepanshu", "Judges, what you saw is not just a college project. It is backed by real research.",
             demo="Switch to Credibility slide.")
    add_line(doc, "Deepanshu", "Our Ethical Algorithm work is published in Infinity. This means our AI is built to be fair and transparent, with no bias against any customer group.")
    add_line(doc, "Mrigesh", "We were also finalists at NMIMS Mumbai for our MSME-focused work — specifically on helping small businesses get formal credit using cash-flow data.")
    add_line(doc, "Deepanshu", "And technically, the product is solid — 19 frontend tests and 89 backend tests are passing. It is already deployed on Render with a live backend and frontend.")

    doc.add_page_break()

    add_heading(doc, "8:50 – 9:25 | Future Scope", level=2)
    add_line(doc, "Rikshita", "Looking ahead, our roadmap is clear.",
             demo="Switch to Roadmap slide.")
    add_line(doc, "Rikshita", "Phase 1: Pilot with one PSB branch — one thousand retail customers and one hundred SMEs.")
    add_line(doc, "Rikshita", "Phase 2: Connect live Account Aggregator partners like PhonePe and Finvu.")
    add_line(doc, "Rikshita", "Phase 3: Fine-tune our AI using anonymized PSB transaction data.")
    add_line(doc, "Rikshita", "Phase 4: Roll out to all PSB internet banking users.")

    doc.add_page_break()

    add_heading(doc, "9:25 – 10:00 | Impact + Closing", level=2)
    add_line(doc, "Kunal", "The impact we expect is clear.",
             demo="Switch to Impact Metrics slide.")
    add_line(doc, "Kunal", "Three times increase in monthly active users. Fifteen percent growth in FD and recurring deposit bookings. Forty percent reduction in reported digital fraud. And a twenty-point improvement in customer satisfaction score.")
    add_line(doc, "Kunal", "To conclude, PSB SecureWealth Twin is not just an internet banking upgrade. It is a product that turns a bank account into a personal wealth companion.",
             demo="Switch to Thank You slide with QR code.")
    add_line(doc, "Kunal", "It unifies scattered money. It protects against fraud. It serves individuals, SMEs, and NRIs. It is compliant, inclusive, and ready to deploy today. We are Team DS Financial Solutions, ready to build this for Punjab & Sind Bank. Thank you, judges. We would love to take your questions.")
    add_note(doc, "All team members step forward, smile, and stand ready for Q&A. End exactly at 10:00.")

    doc.add_page_break()

    add_heading(doc, "PPT Slide Outline", level=1)
    slides = [
        ("1", "Title", "PSB SecureWealth Twin | One Bank. One Twin. Infinite Possibilities. | Team DS Financial Solutions"),
        ("2", "The Problem", "Scattered money, missed deadlines, rising fraud — one image of confused customer"),
        ("3", "Solution 1 — Wealth Twin", "Unified dashboard screenshot"),
        ("4", "Solution 2 — Security Beast", "Rakshak AI, Deepfake Shield, Decoy, Ghost Mode icons"),
        ("5", "Solution 3 — SME & Tax", "Cash Flow Timeline and Old vs New Tax calculator screenshots"),
        ("6", "Solution 4 — Macro Signals", "Macro Signal Tower card screenshot"),
        ("7", "Solution 5 — Smart AI", "AI Action Cards and Smart Sweep screenshot"),
        ("8", "Innovation 1 — Life-Shock Simulator", "Job-loss scenario screenshot"),
        ("9", "Innovation 2 — Generational Slider", "Wealth projection slider screenshot"),
        ("10", "Innovation 3 — Security Innovations", "Voice shield + Decoy + Ghost Mode"),
        ("11", "Innovation 4 — CreditBridge AI", "MSME credit score dashboard"),
        ("12", "Innovation 5 — More Features", "NRI Mode, Fantasy League, Sovereign Vault"),
        ("13", "Credibility", "Infinity publication + NMIMS finalist + test counts"),
        ("14", "Roadmap", "Phase 1 to Phase 4"),
        ("15", "Impact", "3x MAU, 15% FD growth, 40% fraud reduction, +20 NPS"),
        ("16", "Thank You", "QR code to live demo + contact"),
    ]
    for num, title, content in slides:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"Slide {num}: {title} — ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(content)
        r2.font.size = Pt(11)

    doc.add_page_break()

    add_heading(doc, "Live Demo Checklist", level=1)
    add_para(doc, "□ Open psb-securewealth-frontend.onrender.com in Chrome before the pitch starts.")
    add_para(doc, "□ Log in once with a demo profile to warm up the backend.")
    add_para(doc, "□ Set zoom to 125% so judges can read text.")
    add_para(doc, "□ Keep slides open in another tab or second laptop.")
    add_para(doc, "□ Have screenshots as backup for every screen.")
    add_para(doc, "□ Assign one teammate to click during demo; speakers face judges.")
    add_para(doc, "□ Close unrelated tabs and notifications.")

    add_heading(doc, "If Running Short On Time", level=1)
    add_para(doc, "1. Skip Innovation 5 (More Unique Features).")
    add_para(doc, "2. Skip detailed tax calculator demo in Solution 3.")
    add_para(doc, "3. Shorten Future Scope from four phases to two.")
    add_para(doc, "4. Never cut opening hook, closing, or credibility slide.")

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
