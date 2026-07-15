#!/usr/bin/env python3
"""Generate a single-person 10-minute live website demo script."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

OUT_PATH = "PSB_Solo_10Min_Live_Demo_Script.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    return h


def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_demo_step(doc, time, spoken, action=None, note=None):
    add_heading(doc, time, level=3)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("Say: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    r1.font.size = Pt(11)
    r2 = p.add_run(spoken)
    r2.font.size = Pt(11)
    if action:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.left_indent = Inches(0.2)
        r3 = p2.add_run(f"[SCREEN: {action}]")
        r3.bold = True
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x00, 0x66, 0x33)
        r3.font.size = Pt(10)
    if note:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(6)
        p3.paragraph_format.left_indent = Inches(0.2)
        r4 = p3.add_run(f"💡 Tip: {note}")
        r4.italic = True
        r4.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
        r4.font.size = Pt(10)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PSB SecureWealth Twin\nSingle-Person 10-Minute Live Demo Script")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    t.paragraph_format.space_after = Pt(12)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Website: psb-securewealth-frontend.onrender.com")
    r.font.size = Pt(14)
    r.italic = True
    st.paragraph_format.space_after = Pt(18)

    add_para(doc,
        "This script is for ONE speaker giving a 10-minute live product demo. Every line you speak is written below. "
        "Every click, scroll, and toggle is in [SCREEN: ...] brackets. Practice twice with a timer.",
        italic=True, space_after=12)

    add_heading(doc, "Before You Start", level=1)
    add_para(doc, "□ Open the website in Chrome. Zoom to 125%.")
    add_para(doc, "□ Log out if already logged in. You must start from the login page.")
    add_para(doc, "□ Close all other tabs and notifications.")
    add_para(doc, "□ Keep a second device with screenshots as backup.")
    add_para(doc, "□ Speak slowly. Judges need to see AND hear.")

    doc.add_page_break()

    add_heading(doc, "0:00 – 0:45 | Opening on Login Page", level=2)
    add_demo_step(doc, "0:00 – 0:15",
        "Good morning judges. I am going to show you PSB SecureWealth Twin — a product that turns a normal bank account into a personal wealth companion.",
        action="Face the judges. The login page should be visible on screen.",
        note="Do not click yet. Let judges absorb the login page.")
    add_demo_step(doc, "0:15 – 0:30",
        "Our tagline is simple: One Bank. One Twin. Infinite Possibilities. Right now you are seeing the login page. Users can sign in normally, use Face Login, use passkey, or use one of these demo profiles for quick access.",
        action="Point cursor to the demo profile list on the right side of the login page. Slowly scroll through 2-3 profiles.",
        note="Show that there are multiple realistic personas.")
    add_demo_step(doc, "0:30 – 0:45",
        "I will log in as Neha, a salaried professional. Watch what happens — this is not just a login, this is where the magic starts.",
        action="Click on Neha's profile.",
        note="Wait for the Account Aggregator animation to start.")

    doc.add_page_break()

    add_heading(doc, "0:45 – 1:45 | Account Aggregator Animation", level=2)
    add_demo_step(doc, "0:45 – 1:00",
        "You see this animation? This is the Account Aggregator connecting Neha's bank accounts, FDs, mutual funds, and insurance into one secure view. In real life, this uses RBI's Account Aggregator framework with user consent.",
        action="Let the AA animation play. Point to each institution as it connects.",
        note="Do not skip the animation. It builds credibility.")
    add_demo_step(doc, "1:00 – 1:20",
        "Now we are on the Wealth Twin Dashboard. This single screen solves the biggest problem in Indian banking today — scattered money. Customers have accounts everywhere, but they never see the full picture.",
        action="After animation finishes, the dashboard loads. Pause and let judges see the full screen.",
        note="This is the problem statement moment.")
    add_demo_step(doc, "1:20 – 1:45",
        "Look at the top. Net worth. Asset breakdown. Recent transactions. Liability summary. Everything is here. No more opening five different apps. No more guessing where the money is. This is the first solution — unified financial view.",
        action="Point to Net Worth card, then Asset Breakdown chart, then Liability section, then Recent Transactions.",
        note="Move mouse slowly. Let eyes follow.")

    doc.add_page_break()

    add_heading(doc, "1:45 – 3:00 | Problem Statement + Dashboard Deep-Dive", level=2)
    add_demo_step(doc, "1:45 – 2:10",
        "Let me explain the problem more clearly with this screen. Neha earns well, but her money is sleeping in a savings account. Inflation is eating it. She does not know how much to invest, where to invest, or when to act. This is the second problem — banks show balance, but they do not give advice.",
        action="Scroll down slightly to show the AI Action Cards section.",
        note="Connect the visual to the problem.")
    add_demo_step(doc, "2:10 – 2:35",
        "Here is the answer. AI Action Cards. This card says Neha has two lakh rupees idle for thirty days. The system is suggesting a one-tap FD ladder. Another card says her emergency fund is only sixty percent complete and suggests a small monthly top-up.",
        action="Point to the first AI Action Card, then the second. Click the first card if it expands.",
        note="Show that advice is specific and actionable.")
    add_demo_step(doc, "2:35 – 3:00",
        "This is the difference between old banking and new banking. Old banking says 'your balance is this much.' New banking says 'here is what you should do with it.' Let me show you Smart Sweep, which actually moves idle money automatically.",
        action="Click on 'Smart Sweep' or the first AI Action Card that leads to Smart Sweep.",
        note="Transition smoothly to the next feature.")

    doc.add_page_break()

    add_heading(doc, "3:00 – 4:30 | Smart Sweep + Macro Signals", level=2)
    add_demo_step(doc, "3:00 – 3:20",
        "Smart Sweep scans all linked accounts and finds idle money. Then it tells Neha the best place to park it — FD ladder, liquid fund, or overnight fund — based on when she needs the money.",
        action="Show the Smart Sweep screen. Point to the linked accounts and idle balance.",
        note="This directly answers 'sell gold, shift to FD' type problem.")
    add_demo_step(doc, "3:20 – 3:45",
        "The recommendation is not random. It looks at the macro economy too. Let me go back to the dashboard and show you the Macro Signal Tower.",
        action="Go back to Dashboard. Scroll to the Macro Signal Tower card.",
        note="Use the browser back button or sidebar.")
    add_demo_step(doc, "3:45 – 4:10",
        "Here we track RBI repo rate, inflation, USD-INR, and gold trend. The system converts these into simple advice. For example, because repo rate is rising, it recommends floating-rate FD. Because gold is high, it suggests booking partial profit.",
        action="Point to each signal and its recommended action.",
        note="This shows economic awareness.")
    add_demo_step(doc, "4:10 – 4:30",
        "So we are not just giving advice based on the user's data. We are giving advice based on the real economy. And every projection carries this disclaimer — it is a simulation, not a guaranteed return. Now let me show you security.",
        action="Point to the disclaimer text at the bottom of the card or screen.",
        note="Compliance moment — judges will notice.")

    doc.add_page_break()

    add_heading(doc, "4:30 – 6:00 | Security Beast", level=2)
    add_demo_step(doc, "4:30 – 4:50",
        "When all this financial data lives in one place, security is everything. We have built the Security Beast. Let me open it.",
        action="Click on 'Security Beast' or 'Fraud Protection' in the sidebar.",
        note="Security is a strong judging point.")
    add_demo_step(doc, "4:50 – 5:15",
        "First layer is Rakshak AI. It learns what is normal for Neha — when she pays, how much she pays, from which device. If something unusual happens, like a transaction at 3 AM from a new device, it blocks it before money leaves the account.",
        action="Point to the Rakshak AI risk score or fraud alert simulation.",
        note="Use a demo fraud alert if available.")
    add_demo_step(doc, "5:15 – 5:35",
        "Second layer is Deepfake Voice Shield. Fraudsters are cloning voices now. Before any high-risk action, the user speaks a rotating phrase. The system matches voice print and checks liveness.",
        action="If voice verification screen is available, show it. Otherwise, show the feature listed in Security Beast.",
        note="This is an innovation highlight.")
    add_demo_step(doc, "5:35 – 6:00",
        "Third and fourth layers are Decoy Account and Ghost Mode. If someone forces Neha to unlock the app, she can enter a duress PIN. The app opens a fake low-balance account. Ghost Mode hides real balances when she is in public.",
        action="Show the Decoy Account and Ghost Mode toggles in Settings or Security section.",
        note="These are wow features.")

    doc.add_page_break()

    add_heading(doc, "6:00 – 7:30 | SME Centre + Tax + Inclusion", level=2)
    add_demo_step(doc, "6:00 – 6:15",
        "This product is not just for individuals. We also serve small businesses. Let me switch to SME Centre.",
        action="Click on 'SME Centre' or 'Business Mode' in the navigation.",
        note="Switch user or mode if the demo allows. Otherwise, show the SME dashboard directly.")
    add_demo_step(doc, "6:15 – 6:35",
        "Here a business owner sees Cash Flow Timeline — monthly money in and money out. Red months show where the business may run short. Green months show surplus.",
        action="Point to the Cash Flow Timeline graph. Highlight one red month and one green month.",
        note="Make it relatable for business judges.")
    add_demo_step(doc, "6:35 – 6:55",
        "Working Capital Health score shows current ratio, quick ratio, and days to collect payment versus days to pay suppliers. It is colour-coded green, amber, red.",
        action="Point to the Working Capital Health score and ratios.",
        note="This shows financial depth.")
    add_demo_step(doc, "6:55 – 7:15",
        "If the business has surplus cash, Surplus Fund Advisor suggests FD sweep, liquid fund, or early vendor payment — all with projected savings. One click, and the advice becomes action.",
        action="Show a surplus recommendation and its projected savings.",
        note="Connect back to Smart Sweep idea.")
    add_demo_step(doc, "7:15 – 7:30",
        "For retail users, we have Advanced Tax Centre with Old vs New regime calculator, 80C tracker, and deadline calendar. We also have Senior Mode, Face Login, and vernacular support for inclusion. Now let me show you the innovations.",
        action="Quickly show the Tax Centre or Senior Mode toggle.",
        note="Fast transition to innovation section.")

    doc.add_page_break()

    add_heading(doc, "7:30 – 9:00 | Innovation Showcase", level=2)
    add_demo_step(doc, "7:30 – 7:50",
        "First innovation: Life-Shock Simulator. Most people do not plan for emergencies. This lets Neha test job loss, medical emergency, or market crash.",
        action="Navigate to Wealth Twin → Life-Shock Simulator. Select 'Job Loss' scenario.",
        note="This is emotionally powerful.")
    add_demo_step(doc, "7:50 – 8:05",
        "See how the net worth drops and goals get delayed? The system then recommends increasing emergency fund or buying term insurance.",
        action="Show the impact graph and recommended action.",
        note="Pause for judges to see the value.")
    add_demo_step(doc, "8:05 – 8:25",
        "Second innovation: Generational Wealth Slider. Neha can slide from today to 2056 and see her retirement corpus grow, inflation impact, and whether she is on track.",
        action="Navigate to Goals or Wealth Twin. Move the Generational Wealth Slider from 2026 to 2056.",
        note="Visual wow moment.")
    add_demo_step(doc, "8:25 – 8:45",
        "Third innovation: CreditBridge AI for MSMEs. Many small businesses have no credit history. Our AI analyses cash flow, GST, and UPI receipts to give a credit score and suggest the right PSB loan.",
        action="Navigate to SME Centre → CreditBridge AI. Show the credit score and recommended loan product.",
        note="Connects to NMIMS MSME work.")
    add_demo_step(doc, "8:45 – 9:00",
        "We also have NRI Mode, Fantasy League of Wealth for family learning, and Sovereign Vault for long-term secure holdings. These make the product complete.",
        action="Quickly toggle NRI Mode or show the feature list.",
        note="End innovation section strongly.")

    doc.add_page_break()

    add_heading(doc, "9:00 – 10:00 | Credibility + Closing", level=2)
    add_demo_step(doc, "9:00 – 9:15",
        "What you saw is not just a demo. It is backed by real research. Our Ethical Algorithm work is published in Infinity, ensuring fair and transparent AI recommendations.",
        action="Switch to slide showing Infinity publication and credibility points.",
        note="Use slides for credibility since numbers are easier to read.")
    add_demo_step(doc, "9:15 – 9:30",
        "We were also finalists at NMIMS Mumbai for our MSME credit work. And technically, this product has 19 frontend tests and 89 backend tests passing, already deployed live.",
        action="Show the NMIMS finalist mention and test counts on the slide.",
        note="Build trust.")
    add_demo_step(doc, "9:30 – 9:50",
        "With PSB SecureWealth Twin, Punjab & Sind Bank can increase monthly active users by three times, grow FD bookings by fifteen percent, reduce fraud by forty percent, and improve customer satisfaction by twenty points.",
        action="Switch to impact metrics slide.",
        note="End with numbers.")
    add_demo_step(doc, "9:50 – 10:00",
        "It unifies scattered money, protects users, serves SMEs, and makes banking proactive. Thank you, judges. I would be happy to take your questions.",
        action="Switch to Thank You slide with QR code. Step back and smile.",
        note="Finish exactly at 10 minutes.")

    doc.add_page_break()

    add_heading(doc, "Demo Navigation Cheat Sheet", level=1)
    add_para(doc, "Keep this page visible on a separate screen or print it.")
    add_para(doc, "1. Login page → Click Neha profile → Wait for AA animation")
    add_para(doc, "2. Dashboard → Net Worth → Asset Breakdown → AI Action Cards")
    add_para(doc, "3. Smart Sweep → Back to Dashboard → Macro Signal Tower")
    add_para(doc, "4. Security Beast → Rakshak AI → Deepfake Voice → Decoy/Ghost Mode")
    add_para(doc, "5. SME Centre → Cash Flow → Working Capital → Surplus Advisor")
    add_para(doc, "6. Wealth Twin → Life-Shock Simulator → Generational Slider")
    add_para(doc, "7. SME Centre → CreditBridge AI")
    add_para(doc, "8. Slides: Credibility → Impact → Thank You")

    add_heading(doc, "If Something Breaks On Stage", level=1)
    add_para(doc, "□ If website is slow: say 'Let me show you a screenshot while the live page loads' and switch to backup screenshots.")
    add_para(doc, "□ If a click does not work: do not click ten times. Move to the next feature smoothly.")
    add_para(doc, "□ If you forget a line: pause, look at the screen, and describe what you see. That is enough.")
    add_para(doc, "□ If you run short on time: skip SME Centre details and jump directly to innovations.")

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
