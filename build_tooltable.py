# -*- coding: utf-8 -*-
"""10-Minute Demo Tool Table (Hinglish) -> Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0B,0x2B,0x52); BLUE=RGBColor(0x1E,0x5A,0x9C); RED=RGBColor(0xB3,0x1B,0x1B)
GREEN=RGBColor(0x1B,0x7A,0x3D); GREY=RGBColor(0x55,0x55,0x55); WHITE=RGBColor(0xFF,0xFF,0xFF)

doc=Document()
sec=doc.sections[0]
sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.top_margin=Inches(0.4); sec.bottom_margin=Inches(0.4); sec.left_margin=Inches(0.4); sec.right_margin=Inches(0.4)
doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(8.5)

def shade(cell,hexcolor):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexcolor); tcPr.append(sh)

def title(text,size,color,after=2,it=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(after)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color; r.italic=it
    return p

title('SECUREWEALTH TWIN — 10 MINUTE DEMO TOOL TABLE',16,NAVY)
title('Kaunsa tool dikhana hai, kab, kyun (problem statement), kaise bana, aur kaunsa score milega',9.5,GREY,after=6,it=True)

# ---- main table ----
headers=['#','Tool (Component)','Time','PS Relevance (kya cover karta hai)','Kaise Bana (library/method)','Score','Pri']
widths=[0.3,1.9,0.65,3.1,2.9,0.8,0.55]
rows=[
 ('1','Architecture slide + live app','0:45','Bank integration, real prototype','PDF flow diagram; deployed Render+Surge','M6,7','***'),
 ('2','DashboardView + NetWorthCard','1:30','Net worth, spending/income (PS #1,5)','React+TS, Context state; balances+assets sum','M1,3','***'),
 ('3','AccountAggregatorFull + AAFetchAnimation','2:00','AA integration (PS #4)','React consent flow; aaBanks data; simulated fetch','M1,Comp','****'),
 ('4','ManualAssetForm / PhysicalAssetIntelligence','2:20','Gold/property/vehicle -> net worth (PS #5)','Form -> backend /banking/assets (per-user SQLite)','M1','**'),
 ('5','GoalTracker','2:45','Goal planning "Save Rs.X more" (Outcome #3)','Component + useRecommendationEngine hook','M1','***'),
 ('6','AIRecommendationsView + NBAInsights','3:00','Timely suggestion WITH reason (PS #6)','aiOrchestrator (multi-provider) + explainable card','M1,4','****'),
 ('7','DeviceStatusCard / DeviceFingerprintPanel','3:30','Trusted vs new device (Protection #1)','fingerprintService.ts -> device id, trust flag','M2','*****'),
 ('8','timingCheck (<10s flag)','4:00','Rushed action after login (Protection #2)','Express middleware timingCheck.js, timestamp diff','M2','*****'),
 ('9','RiskMeter (amount vs history)','4:45','Unusual amount (Protection #3)','fraudDetectionService.ts - rule-based weighted scoring','M2','*****'),
 ('10','OTPSimulation','5:15','OTP misuse pattern (Protection #4)','Component + backend /otp route, retry counter','M2','****'),
 ('11','Risk Score -> Low/Med/High','5:30','Wealth Protection Risk Score (PS exact)','Sum of signal weights + thresholds','M2','*****'),
 ('12','TransactionGuardModal + CoolingVaultModal','6:00','Allow / Warn / Block + cooling-off (PS exact)','React modals gated by score','M2','*****'),
 ('13','DuressMode + DuressPinSetup','6:30','Coerced transactions (PS - the WOW)','duressService.ts: secret PIN -> fake success + lock','M2,5','*****'),
 ('14','AuditLog','6:45','Audit layer (PS architecture)','auditLogger middleware + SQLite records','M2,6','***'),
 ('15','ExplainableTooltip / AIDecisionLog','7:15','Explainable AI (Compliance #3)','Tooltip components surfacing rule reasons','M4,Comp','****'),
 ('16','ConsentModal + PrivacyCenter','7:45','Consent + data transparency (Comp #1)','Modal + privacy audit panel','Comp','***'),
 ('17','KYCModal','8:00','KYC before investment (Comp #5)','Component + backend /kyc routes','Comp','**'),
 ('18','InnovationLab (Bhavishya/FutureSelf/Business)','9:00','Prediction, scenario sim, corporate','Client-side Monte-Carlo style simulation','M5','***'),
 ('19','Close + scale line','10:00','Scalability + integration (PS #7)','React + Node microservice + Docker + mock CBS','M6,7','***'),
]
t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
hc=t.rows[0].cells
for i,h in enumerate(headers):
    shade(hc[i],'0B2B52'); pr=hc[i].paragraphs[0]; rn=pr.add_run(h); rn.bold=True; rn.font.color.rgb=WHITE; rn.font.size=Pt(8.5)
for row in rows:
    cells=t.add_row().cells
    is_core = row[5]=='M2'
    for i,val in enumerate(row):
        if i==1 and is_core: shade(cells[i],'FBE9E9')
        if i==6 and val=='*****': shade(cells[i],'FFF3CD')
        pr=cells[i].paragraphs[0]; rn=pr.add_run(val); rn.font.size=Pt(8.2)
        if i in (0,6): pr.alignment=WD_ALIGN_PARAGRAPH.CENTER
        if i==1: rn.bold=True
for r in t.rows:
    for i,w in enumerate(widths): r.cells[i].width=Inches(w)

p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
r=p.add_run('Note: Red rows = mandatory Protection Layer (M2) = sabse zyada marks. Pri = Priority (zyada * = zyada zaroori).')
r.italic=True; r.font.size=Pt(8); r.font.color.rgb=GREY

# ---- Top 5 ----
title('TOP 5 MUST-SHOW (time kam pade to yeh chhodna mat)',12,RED,after=3)
t2=doc.add_table(rows=1,cols=3); t2.style='Table Grid'
for i,h in enumerate(['Rank','Tool','Kyun #1 priority']):
    shade(t2.rows[0].cells[i],'B31B1B'); pr=t2.rows[0].cells[i].paragraphs[0]; rn=pr.add_run(h); rn.bold=True; rn.font.color.rgb=WHITE; rn.font.size=Pt(9)
top=[
 ('1','Risk Score -> Allow/Warn/Block (#11,#12)','PS ka mandatory CORE - sabse zyada marks'),
 ('2','DuressMode (#13)','Coerced transaction = unique, judges yaad rakhenge'),
 ('3','timingCheck <10s (#8)','PS ne exactly yeh maanga, humne literally banaya'),
 ('4','AIRecommendations with reason (#6)','Intelligence + Explainability ek saath'),
 ('5','Account Aggregator (#3)','Zyadatar teams skip karti hain - differentiator'),
]
for rank,tool,why in top:
    c=t2.add_row().cells
    c[0].paragraphs[0].add_run(rank).bold=True
    rn=c[1].paragraphs[0].add_run(tool); rn.bold=True; rn.font.size=Pt(9)
    c[2].paragraphs[0].add_run(why).font.size=Pt(9)
for r in t2.rows:
    r.cells[0].width=Inches(0.5); r.cells[1].width=Inches(3.3); r.cells[2].width=Inches(5.4)

# ---- Key points ----
title('YAAD RAKHNE WALI BAATEIN',12,GREEN,after=3)
pts=[
 ('Protection block (3:00-6:45) pe sabse zyada time do', ' - ye ~40% score hai; domain hi "Cyber Security & Fraud" hai.'),
 ('Har tool ke saath ek "reason" dikhao', ' - explainability har metric mein bonus deti hai.'),
 ('"Kaise bana" pucha jaaye to: ', 'Frontend = React+TypeScript+Vite+Context API; Backend = Node+Express+SQLite+JWT(HS256)+bcrypt; AI = multi-provider orchestrator + circuit breaker; Protection = rule-based weighted signals.'),
 ('Do tabs khol ke rakho: ', 'timingCheck.js + fraudDetectionService.ts - "show me the code" pe turant dikha do.'),
 ('Mat dikhana: ', 'Scenario/NLP what-if (throw ho sakta), exact tax rupee figures, GST/DCF tools, aur "47 patents/military-grade" wording.'),
]
for lead,rest in pts:
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(lead); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=NAVY
    p.add_run(rest).font.size=Pt(9)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6)
r=p.add_run('WIN FORMULA: Protection layer flawless dikhao + clean & focused raho + har metric pe ek line + koi jhootha claim nahi.')
r.bold=True; r.font.size=Pt(10); r.font.color.rgb=GREEN

doc.save(r'G:/PSB/DS_Financial/frontend-render/SecureWealth_Twin_Tool_Table.docx')
print('SAVED: SecureWealth_Twin_Tool_Table.docx')
