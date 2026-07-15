#!/usr/bin/env python3
"""Generate updated 10-min pitch script with equal focus on innovation + live demo cues."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "PSB_Hackathon_10Min_Pitch_Product_Demo_Script.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    return h


def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_line(doc, speaker, text, demo=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p.add_run(f"{speaker}: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    if demo:
        r3 = p.add_run(f"\n    [LIVE DEMO: {demo}]")
        r3.bold = True
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x00, 0x66, 0x33)
        r3.font.size = Pt(10)
    return p


def add_slide(doc, slide_no, title, content, speaker=None):
    add_heading(doc, f"Slide {slide_no}: {title}", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if speaker:
        r = p.add_run(f"Speaker: {speaker}\n")
        r.bold = True
        r.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
        r.font.size = Pt(11)
    r2 = p.add_run(f"Visual: {content}")
    r2.font.size = Pt(11)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(f"💡 Tip: {text}")
    r.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
    r.font.size = Pt(10)
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PSB SecureWealth Twin\n10-Minute Product Pitch + Live Demo Script")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    t.paragraph_format.space_after = Pt(12)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Product Pitch Format | Equal Focus: Problem + Solutions + Innovation")
    r.font.size = Pt(14)
    r.italic = True
    st.paragraph_format.space_after = Pt(18)

    add_para(doc,
        "This version treats the website as a product being sold to Punjab & Sind Bank. "
        "Live demo cues are written in [LIVE DEMO: ...] brackets. "
        "Innovation section has equal weight as problem+solutions. Language is simple and professional.",
        italic=True, space_after=12)

    add_heading(doc, "Speaker Roles", level=1)
    add_para(doc, "1. Rikshita Barua — Opening hook + Problem statement (0:00–1:15)")
    add_para(doc, "2. Kunal Saxena — Solution 1: Wealth Twin Dashboard (1:15–2:00)")
    add_para(doc, "3. Deepanshu Sharma — Solution 2: AI Action Cards + Smart Sweep (2:00–2:45)")
    add_para(doc, "4. Mrigesh Mohanty — Solution 3: Security Beast (2:45–3:30)")
    add_para(doc, "5. Ishita Anand — Solution 4: SME Centre + Tax + Inclusion (3:30–4:15)")
    add_para(doc, "6. Tripti Jain — Solution 5: Macro Signal Tower + Compliance (4:15–5:00)")
    add_para(doc, "7. Kunal Saxena — Innovation showcase + Credibility + Future + Closing (5:00–10:00)")

    add_heading(doc, "Time Balance", level=1)
    add_para(doc, "Problem statement: 1 min 15 sec")
    add_para(doc, "Core solutions: 3 min 45 sec")
    add_para(doc, "Innovation showcase: 3 min")
    add_para(doc, "Credibility + Future + Closing: 2 min")
    add_para(doc, "Total: 10 minutes")

    doc.add_page_break()

    # PPT OUTLINE
    add_heading(doc, "PPT Slide Outline", level=1)
    add_slide(doc, "1", "Title", "Product name: PSB SecureWealth Twin. Tagline: One Bank. One Twin. Infinite Possibilities. Team DS Financial Solutions.", "Rikshita")
    add_slide(doc, "2", "The Problem", "One image: confused customer looking at multiple banking apps. Text: Money is scattered. Advice is missing. Fraud is rising.", "Rikshita")
    add_slide(doc, "3", "Solution 1 — Unified Wealth Twin", "Screenshot of Wealth Twin Dashboard with net worth, asset breakdown, and AI Action Cards.", "Kunal")
    add_slide(doc, "4", "Solution 2 — Smart AI Recommendations", "Screenshot of Smart Sweep and AI Action Cards.", "Deepanshu")
    add_slide(doc, "5", "Solution 3 — Security Beast", "Icons: Rakshak AI, Deepfake Voice Shield, Decoy Account, Ghost Mode.", "Mrigesh")
    add_slide(doc, "6", "Solution 4 — SME & Tax Centre", "Screenshot of SME Cash Flow Timeline and Old vs New Tax Regime calculator.", "Ishita")
    add_slide(doc, "7", "Solution 5 — Macro Signal Tower", "Screenshot of macro signals card and compliance disclaimer.", "Tripti")
    add_slide(doc, "8", "Innovation 1 — Life-Shock Simulator", "Screenshot of simulator showing job-loss scenario.", "Kunal")
    add_slide(doc, "9", "Innovation 2 — Deepfake Voice Shield", "Screenshot or flow diagram of voice verification.", "Kunal")
    add_slide(doc, "10", "Innovation 3 — CreditBridge AI", "Screenshot of MSME credit score dashboard.", "Kunal")
    add_slide(doc, "11", "Innovation 4 — Generational Wealth Slider", "Screenshot of slider projecting corpus over decades.", "Kunal")
    add_slide(doc, "12", "Innovation 5 — More Beast Features", "Decoy Account, Ghost Mode, NRI Mode, Fantasy League of Wealth.", "Kunal")
    add_slide(doc, "13", "Credibility", "Infinity publication logo, NMIMS Mumbai finalist mention, test counts: 19 frontend + 89 backend tests passing.", "Kunal")
    add_slide(doc, "14", "Roadmap", "Phase 1 Pilot → Phase 2 AA Integration → Phase 3 AI Fine-tuning → Phase 4 Full Rollout.", "Kunal")
    add_slide(doc, "15", "Impact", "Metrics: 3x MAU, 15% FD growth, 40% fraud reduction, +20 NPS points.", "Kunal")
    add_slide(doc, "16", "Thank You / Demo Link", "QR code to live website + GitHub repo + contact emails.", "Kunal")

    doc.add_page_break()

    # SCRIPT START
    add_heading(doc, "Detailed 10-Minute Script with Live Demo Cues", level=1)

    add_heading(doc, "0:00 – 1:15 | Opening + Problem Statement", level=2)
    add_line(doc, "Rikshita", "Good morning judges. I am Rikshita, and we are Team DS Financial Solutions.")
    add_line(doc, "Rikshita", "Today, we are not showing you just a website. We are selling you a product — a product that can change how Punjab & Sind Bank serves its customers.")
    add_line(doc, "Rikshita", "Our product is called PSB SecureWealth Twin. Our tagline is: One Bank. One Twin. Infinite Possibilities.")
    add_line(doc, "Rikshita", "Let me tell you the problem first. Today, an average Indian customer has money in many places — savings account, FD, loan, insurance, mutual funds, gold.")
    add_line(doc, "Rikshita", "But when they open their bank app, they only see one savings account balance. They cannot see the full picture.")
    add_line(doc, "Rikshita", "Because of this, three big problems happen. One, idle money sleeps in savings accounts and loses value. Two, people miss tax-saving deadlines. Three, fraudsters take advantage of this confusion through fake calls and phishing.")
    add_line(doc, "Rikshita", "So the real problem is this: banks today store money, but they do not actively grow, protect, and simplify their customers' financial lives.")
    add_line(doc, "Rikshita", "Our product fixes exactly this. I will hand over to Kunal, who will show you the first solution.")
    add_note(doc, "Rikshita should be confident. This opening sets the tone. Make eye contact with all judges.")

    doc.add_page_break()

    add_heading(doc, "1:15 – 2:00 | Solution 1: Unified Wealth Twin", level=2)
    add_line(doc, "Kunal", "Thank you, Rikshita. I am Kunal. Let me show you how PSB SecureWealth Twin solves the scattered-money problem.",
             demo="Open the website psb-securewealth-frontend.onrender.com. Click 'Demo Mode'. Select 'Neha — Salaried Professional'. Wait for the Account Aggregator animation to finish.")
    add_line(doc, "Kunal", "This is the Wealth Twin Dashboard. It is a single screen where the customer sees everything — bank accounts, FDs, loans, insurance, gold, mutual funds, and real estate.")
    add_line(doc, "Kunal", "At the top, you see the net worth. On the left, the asset breakdown. On the right, recent transactions.",
             demo="Point the cursor at the Net Worth card, then the Asset Breakdown chart, then the Recent Transactions list.")
    add_line(doc, "Kunal", "The best part is the AI Action Cards. They do not just show data — they tell the user what to do. For example, here the card says: 'You have two lakh rupees idle for thirty days. Sweep it to FD ladder in one tap.'")
    add_line(doc, "Kunal", "So our first solution gives the customer a complete financial picture and turns it into action. Deepanshu will now show you the intelligence layer.")

    doc.add_page_break()

    add_heading(doc, "2:00 – 2:45 | Solution 2: AI Action Cards + Smart Sweep", level=2)
    add_line(doc, "Deepanshu", "Thanks, Kunal. I am Deepanshu. Let me show you how the product thinks for the customer.",
             demo="Click on the 'Smart Sweep' or 'AI Actions' section from the dashboard. Show the list of recommendations.")
    add_line(doc, "Deepanshu", "Smart Sweep uses the Account Aggregator framework. It scans all linked accounts and finds idle money.")
    add_line(doc, "Deepanshu", "Then it recommends the best place to park it — FD ladder, liquid fund, or overnight fund — based on how soon the user needs the money.",
             demo="Click one recommendation. Show the projected interest or return comparison.")
    add_line(doc, "Deepanshu", "We also have behavioural nudges. If the user is spending too much, it suggests pausing one subscription. If a SIP is due, it reminds them. If tax deadline is near, it prompts 80C investment.")
    add_line(doc, "Deepanshu", "This makes banking proactive, not reactive. Mrigesh will now show you how we protect all of this.")

    doc.add_page_break()

    add_heading(doc, "2:45 – 3:30 | Solution 3: Security Beast", level=2)
    add_line(doc, "Mrigesh", "Thank you, Deepanshu. I am Mrigesh. When customers link all their money in one place, security becomes the most important thing.",
             demo="Navigate to the Security Beast or Fraud Protection section from the sidebar.")
    add_line(doc, "Mrigesh", "We built a Security Beast with four layers. Layer one is Rakshak AI. It learns what is normal for each user — when they transact, how much, from which device, and from where.")
    add_line(doc, "Mrigesh", "If it sees something unusual, like a transaction at 3 AM from a new city, it blocks it instantly and asks for biometric verification.",
             demo="Show the Rakshak AI risk score card or fraud alert simulation if available.")
    add_line(doc, "Mrigesh", "Layer two is Deepfake Voice Shield. Many frauds now use cloned voices. Our system checks voice print + liveness before allowing high-risk actions.")
    add_line(doc, "Mrigesh", "Layer three is Decoy Account. If someone forces the user to open the app, a duress PIN shows a fake account. Layer four is Ghost Mode, which hides real balances in public.")
    add_line(doc, "Mrigesh", "So our product protects money before it is lost. Ishita will now explain how we help businesses and taxpayers.")

    doc.add_page_break()

    add_heading(doc, "3:30 – 4:15 | Solution 4: SME Centre + Tax + Inclusion", level=2)
    add_line(doc, "Ishita", "Thanks, Mrigesh. I am Ishita. A bank is not just for individuals. It also serves lakhs of small businesses.",
             demo="Click on 'SME Centre' or 'Business Mode' from the navigation menu.")
    add_line(doc, "Ishita", "Our SME Centre has three tools. First, the Cash Flow Timeline. It shows monthly inflow and outflow, and highlights months where the business may run short of money.")
    add_line(doc, "Ishita", "Second, Working Capital Health. It gives a score based on current ratio, quick ratio, and how fast the business collects payments versus how fast it pays suppliers.",
             demo="Point to the Working Capital Health score and the colour-coded ratios.")
    add_line(doc, "Ishita", "Third, the Surplus Fund Advisor. If the business has extra cash, it recommends FD sweep, liquid fund, or early vendor payment — all with projected savings.")
    add_line(doc, "Ishita", "For retail users, we have an Advanced Tax Centre with Old vs New regime calculator, 80C tracker, and tax deadline calendar.")
    add_line(doc, "Ishita", "We also have Senior Mode with big fonts and voice prompts, Face Login, and vernacular language support. Tripti will now explain the macro view.")

    doc.add_page_break()

    add_heading(doc, "4:15 – 5:00 | Solution 5: Macro Signal Tower + Compliance", level=2)
    add_line(doc, "Tripti", "Thank you, Ishita. I am Tripti. Most banks give advice using only customer data. We also use the economy.",
             demo="Go back to Dashboard. Scroll to the Macro Signal Tower card.")
    add_line(doc, "Tripti", "Our Macro Signal Tower tracks RBI repo rate, inflation, USD-INR, and gold prices. It then converts these into simple actions.")
    add_line(doc, "Tripti", "For example, if repo rate is rising, it says: shift to floating-rate FD. If inflation is above six percent, it says: reduce debt duration and increase equity SIP. If gold is rising, it says: book partial profit and move to FD.",
             demo="Point to each signal and its recommended action on the card.")
    add_line(doc, "Tripti", "This directly matches the problem statement example: sell gold, shift to FD based on global indicators.")
    add_line(doc, "Tripti", "And because we show projections, every screen has a clear disclaimer — this is a simulation, not a guaranteed return. Kunal will now show you the innovations that make this product truly powerful.")

    doc.add_page_break()

    add_heading(doc, "5:00 – 8:00 | Innovation Showcase (Equal Focus)", level=2)
    add_line(doc, "Kunal", "Thank you, Tripti. We solved the core problems. Now let me show you what makes this product exciting and different.",
             demo="Stay on the live website. You will demo 5 innovations one by one.")

    add_heading(doc, "Innovation 1: Life-Shock Simulator (5:00–5:35)", level=3)
    add_line(doc, "Kunal", "First, the Life-Shock Simulator. Most people do not plan for emergencies. This tool lets the user test what happens if they lose their job, face a medical emergency, or see a market crash.",
             demo="Navigate to Wealth Twin → Life-Shock Simulator. Select 'Job Loss'. Show how net worth and goals are impacted. Then show the recommended action.")
    add_line(doc, "Kunal", "It shows the impact on net worth and goals, and recommends insurance or contingency fund actions.")

    add_heading(doc, "Innovation 2: Deepfake Voice Shield (5:35–6:10)", level=3)
    add_line(doc, "Kunal", "Second, Deepfake Voice Shield. Voice cloning fraud is rising. Before any high-risk action, the user speaks a phrase. The system checks the voice print and liveness.",
             demo="If the feature is live, show the voice verification screen. Otherwise, show the Security Beast section where this is listed.")
    add_line(doc, "Kunal", "If the voice does not match, the transaction is blocked. This protects users from AI-generated fraud calls.")

    add_heading(doc, "Innovation 3: CreditBridge AI (6:10–6:45)", level=3)
    add_line(doc, "Kunal", "Third, CreditBridge AI for MSMEs. Many small businesses do not have formal credit history, so banks reject their loan applications.",
             demo="Navigate to SME Centre → CreditBridge AI. Show the credit health score and loan product recommendation.")
    add_line(doc, "Kunal", "Our AI analyses cash flow, GST turnover, and UPI receipts to create a credit health score and suggest the right PSB loan product.")

    add_heading(doc, "Innovation 4: Generational Wealth Slider (6:45–7:20)", level=3)
    add_line(doc, "Kunal", "Fourth, the Generational Wealth Slider. A user can slide across years and see how their retirement corpus grows, how inflation reduces purchasing power, and whether their goals are on track.",
             demo="Navigate to Goals or Wealth Twin. Use the Generational Wealth Slider. Move the slider from 2026 to 2056 and show the projection graph.")
    add_line(doc, "Kunal", "It makes long-term planning visual and easy, even for someone who does not understand finance.")

    add_heading(doc, "Innovation 5: More Beast Features (7:20–8:00)", level=3)
    add_line(doc, "Kunal", "Fifth, a quick look at more unique features. Decoy Account for forced unlock situations. Ghost Mode to hide balances in public. NRI Mode for FEMA-aware tracking. Fantasy League of Wealth for family financial literacy.",
             demo="Quickly show the Settings or Security section for Decoy Account / Ghost Mode toggles.")
    add_line(doc, "Kunal", "These features show that we have thought about safety, inclusion, and engagement deeply.")

    doc.add_page_break()

    add_heading(doc, "8:00 – 8:40 | Credibility", level=2)
    add_line(doc, "Kunal", "Judges, what I just showed is not just a college project. It is backed by real research.",
             demo="Switch to Slide 13: Credibility.")
    add_line(doc, "Kunal", "Our Ethical Algorithm work is published in Infinity. This means our AI is built to be fair and transparent, with no bias against any customer group.")
    add_line(doc, "Kunal", "We were also finalists at NMIMS Mumbai for our MSME-focused work — specifically on helping small businesses get formal credit using cash-flow data.")
    add_line(doc, "Kunal", "And technically, the product is solid — 19 frontend tests and 89 backend tests are passing. It is already deployed on Render with a live backend and frontend.")

    doc.add_page_break()

    add_heading(doc, "8:40 – 9:40 | Future Scope + Impact", level=2)
    add_line(doc, "Kunal", "Now, where do we go from here?",
             demo="Switch to Slide 14: Roadmap.")
    add_line(doc, "Kunal", "Phase 1: Pilot with one PSB branch — one thousand retail customers and one hundred SMEs.")
    add_line(doc, "Kunal", "Phase 2: Connect live Account Aggregator partners like PhonePe and Finvu.")
    add_line(doc, "Kunal", "Phase 3: Fine-tune our AI using anonymized PSB transaction data.")
    add_line(doc, "Kunal", "Phase 4: Roll out to all PSB internet banking users.",
             demo="Switch to Slide 15: Impact Metrics.")
    add_line(doc, "Kunal", "The impact we expect is clear. Three times increase in monthly active users. Fifteen percent growth in FD and recurring deposit bookings. Forty percent reduction in reported digital fraud. And a twenty-point improvement in customer satisfaction score.")

    doc.add_page_break()

    add_heading(doc, "9:40 – 10:00 | Closing", level=2)
    add_line(doc, "Kunal", "To conclude, PSB SecureWealth Twin is not just an internet banking upgrade. It is a product that turns a bank account into a personal wealth companion.",
             demo="Switch to Slide 16: Thank You + QR code.")
    add_line(doc, "Kunal", "It unifies scattered money. It protects against fraud. It serves individuals, SMEs, and NRIs. It is compliant, inclusive, and ready to deploy today.")
    add_line(doc, "Kunal", "We are Team DS Financial Solutions, and we are ready to build this for Punjab & Sind Bank. Thank you, judges. We would love to take your questions.")
    add_note(doc, "All team members step forward, smile, and stand ready for Q&A. End exactly at 10:00.")

    doc.add_page_break()

    # Delivery guide
    add_heading(doc, "Live Demo Checklist (Before Going On Stage)", level=1)
    add_para(doc, "□ Open psb-securewealth-frontend.onrender.com in Chrome. Log in once to warm up the backend.")
    add_para(doc, "□ Keep the website open in one tab and slides in another tab (or second laptop).")
    add_para(doc, "□ Have screenshots of every screen as backup in case internet fails.")
    add_para(doc, "□ Set screen zoom to 125% so judges can read text clearly.")
    add_para(doc, "□ Close all unrelated tabs and notifications.")
    add_para(doc, "□ Practice the transition from website to slides smoothly.")
    add_para(doc, "□ Assign one teammate to click/navigate during demo so the speaker can face judges.")

    add_heading(doc, "If You Are Running Short On Time", level=1)
    add_para(doc, "Priority order of cuts (only cut if absolutely needed):")
    add_para(doc, "1. Skip Fantasy League of Wealth mention in Innovation 5.")
    add_para(doc, "2. Skip detailed tax calculator demo in Solution 4.")
    add_para(doc, "3. Shorten Future Scope from four phases to two phases.")
    add_para(doc, "4. Never cut the opening hook, closing, or credibility slide.")

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
