#!/usr/bin/env python3
"""Generate single-person 10-min live demo script with powerful intro + closing."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

OUT_PATH = "PSB_Solo_10Min_Live_Demo_Script_Powerful.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    return h


def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_demo_step(doc, time, spoken, action=None, note=None):
    add_heading(doc, time, level=3)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("Say: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    r1.font.size = Pt(11)
    r2 = p.add_run(spoken)
    r2.font.size = Pt(11)
    if action:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.left_indent = Inches(0.2)
        r3 = p2.add_run(f"[SCREEN: {action}]")
        r3.bold = True
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x00, 0x66, 0x33)
        r3.font.size = Pt(10)
    if note:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(6)
        p3.paragraph_format.left_indent = Inches(0.2)
        r4 = p3.add_run(f"💡 Tip: {note}")
        r4.italic = True
        r4.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
        r4.font.size = Pt(10)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PSB SecureWealth Twin\nSingle-Person 10-Minute Live Demo Script")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    t.paragraph_format.space_after = Pt(12)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Powerful Intro + Powerful Closing | Designed to Win")
    r.font.size = Pt(14)
    r.italic = True
    st.paragraph_format.space_after = Pt(18)

    add_para(doc,
        "This script is for ONE speaker giving a 10-minute live website demo. The opening is built to grab attention instantly. "
        "The closing is built to leave a lasting impression. Every click is in [SCREEN: ...] brackets. Practice twice with a timer.",
        italic=True, space_after=12)

    add_heading(doc, "Before You Start", level=1)
    add_para(doc, "□ Open the website in Chrome. Zoom to 125%.")
    add_para(doc, "□ Log out if already logged in. Start from the login page.")
    add_para(doc, "□ Close all other tabs and notifications.")
    add_para(doc, "□ Keep backup screenshots on a second device.")
    add_para(doc, "□ Speak slowly. Judges need to see AND hear.")

    doc.add_page_break()

    # POWERFUL INTRO
    add_heading(doc, "0:00 – 0:50 | Powerful Opening Hook", level=2)
    add_demo_step(doc, "0:00 – 0:10",
        "Good morning judges. Let me ask you a simple question: How many banking apps do you have on your phone right now?",
        action="Stand confidently. Make eye contact with different judges. The login page is visible behind you.",
        note="Pause for 2-3 seconds. Let them think. Do not answer for them.")
    add_demo_step(doc, "0:10 – 0:25",
        "Three? Four? Maybe five? Now tell me — does even one of them tell you what to DO with your money? Or do they all just show you a balance and leave you confused?",
        action="Gesture with your hands to show confusion. Point slightly toward the login page.",
        note="This connects the problem to everyone's daily life.")
    add_demo_step(doc, "0:25 – 0:40",
        "That is exactly the problem we are solving today. I am going to show you PSB SecureWealth Twin — a product that does not just store your money, but actively grows it, protects it, and simplifies your entire financial life.",
        action="Turn slightly toward the screen. Smile.",
        note="This is your value proposition. Say it with conviction.")
    add_demo_step(doc, "0:40 – 0:50",
        "Our tagline is: One Bank. One Twin. Infinite Possibilities. Let me prove it to you in the next ten minutes.",
        action="Move the cursor to the demo profile list on the login page.",
        note="The word 'prove' builds curiosity.")

    doc.add_page_break()

    add_heading(doc, "0:50 – 1:50 | Login + Account Aggregator", level=2)
    add_demo_step(doc, "0:50 – 1:05",
        "Right now you see the login page. Users can sign in normally, use Face Login, passkey, or pick a demo profile. I will log in as Neha — a regular salaried professional like millions of PSB customers.",
        action="Click on Neha's demo profile.",
        note="Use the profile that has good data loaded.")
    add_demo_step(doc, "1:05 – 1:25",
        "Watch this animation carefully. This is the Account Aggregator connecting Neha's savings account, FD, mutual funds, insurance, and gold into one secure view. In real life, this uses RBI's Account Aggregator framework with full user consent.",
        action="Let the AA animation play. Point to each institution as it connects.",
        note="Do not skip this. It builds trust and explains the tech.")
    add_demo_step(doc, "1:25 – 1:50",
        "And here we are — the Wealth Twin Dashboard. This single screen solves the biggest problem in Indian banking: scattered money. Customers have accounts everywhere, but they never see the full picture. Today, they will.",
        action="After animation finishes, pause and let the full dashboard load.",
        note="This is the first 'wow' visual. Let judges absorb it.")

    doc.add_page_break()

    add_heading(doc, "1:50 – 3:10 | Problem Statement + Dashboard", level=2)
    add_demo_step(doc, "1:50 – 2:15",
        "Look at the top. Net worth. Asset breakdown. Liabilities. Recent transactions. Everything is here. No more opening five apps. No more guessing. No more Excel sheets.",
        action="Point to Net Worth card, then Asset Breakdown chart, then Liabilities, then Recent Transactions.",
        note="Move mouse slowly. Let eyes follow.")
    add_demo_step(doc, "2:15 – 2:40",
        "But here is the real issue. Neha earns well, yet her money is sleeping in a savings account. Inflation is quietly eating it. She does not know how much to invest, where to invest, or when to act. Banks show balance. We give answers.",
        action="Scroll down slightly to show the AI Action Cards section.",
        note="Connect the visual directly to the problem.")
    add_demo_step(doc, "2:40 – 3:10",
        "These are AI Action Cards. This one says Neha has two lakh rupees idle for thirty days and suggests a one-tap FD ladder. This one says her emergency fund is only sixty percent complete. This is not data — this is action.",
        action="Point to the first AI Action Card, then the second. Click the first card if it expands.",
        note="Show that the product thinks, not just displays.")

    doc.add_page_break()

    add_heading(doc, "3:10 – 4:30 | Smart Sweep + Macro Signals", level=2)
    add_demo_step(doc, "3:10 – 3:35",
        "Let me show you Smart Sweep. It scans all linked accounts, finds idle money, and tells Neha the best place to park it — FD ladder, liquid fund, or overnight fund — based on when she needs the money.",
        action="Click on 'Smart Sweep' or the first AI Action Card.",
        note="This is a practical, relatable feature.")
    add_demo_step(doc, "3:35 – 4:00",
        "The advice is not random. It also looks at the macro economy. Let me go back to the dashboard and show you the Macro Signal Tower.",
        action="Go back to Dashboard. Scroll to the Macro Signal Tower card.",
        note="Use sidebar or browser back.")
    add_demo_step(doc, "4:00 – 4:30",
        "We track RBI repo rate, inflation, USD-INR, and gold. If repo rate is rising, we suggest floating-rate FD. If gold is high, we suggest booking partial profit. Every recommendation carries a clear disclaimer — simulation only, no guaranteed returns.",
        action="Point to each signal and its recommended action. Then point to the disclaimer.",
        note="This answers the problem-statement example directly.")

    doc.add_page_break()

    add_heading(doc, "4:30 – 6:00 | Security Beast", level=2)
    add_demo_step(doc, "4:30 – 4:50",
        "Now, when all this money lives in one place, security is everything. We have built the Security Beast. Let me open it.",
        action="Click on 'Security Beast' or 'Fraud Protection' in the sidebar.",
        note="Change your tone slightly — more serious.")
    add_demo_step(doc, "4:50 – 5:15",
        "First layer: Rakshak AI. It learns what is normal for Neha — when she pays, how much, from which device. If it sees something unusual, like a 3 AM transaction from a new city, it blocks it before money leaves.",
        action="Point to the Rakshak AI risk score or fraud alert simulation.",
        note="Use demo fraud alert if available.")
    add_demo_step(doc, "5:15 – 5:35",
        "Second layer: Deepfake Voice Shield. Fraudsters clone voices now. Before any high-risk action, the user speaks a rotating phrase. We match voice print and check liveness.",
        action="Show voice verification screen or the feature listed.",
        note="Strong innovation highlight.")
    add_demo_step(doc, "5:35 – 6:00",
        "Third and fourth layers: Decoy Account and Ghost Mode. If someone forces Neha to unlock the app, a duress PIN shows a fake low-balance account. Ghost Mode hides real balances in public. We protect money before it is lost.",
        action="Show Decoy Account and Ghost Mode toggles.",
        note="These are memorable 'wow' features.")

    doc.add_page_break()

    add_heading(doc, "6:00 – 7:20 | SME Centre + Tax + Inclusion", level=2)
    add_demo_step(doc, "6:00 – 6:15",
        "This product is not just for individuals. We also serve small businesses. Let me switch to SME Centre.",
        action="Click on 'SME Centre' or 'Business Mode'.",
        note="Switch user if needed.")
    add_demo_step(doc, "6:15 – 6:35",
        "Here a business owner sees the Cash Flow Timeline — monthly money in and money out. Red months show where the business may run short. Green months show surplus.",
        action="Point to the Cash Flow Timeline graph. Highlight one red and one green month.",
        note="Relatable for business judges.")
    add_demo_step(doc, "6:35 – 6:55",
        "Working Capital Health gives a score using current ratio, quick ratio, and payment collection speed. It is colour-coded green, amber, red.",
        action="Point to the Working Capital Health score and ratios.",
        note="Shows financial depth.")
    add_demo_step(doc, "6:55 – 7:20",
        "If there is surplus cash, Surplus Fund Advisor suggests FD sweep, liquid fund, or early vendor payment. For retail users, we also have Advanced Tax Centre with Old vs New regime calculator, 80C tracker, and deadline calendar. Plus Senior Mode, Face Login, and vernacular support.",
        action="Show a surplus recommendation, then quickly show Tax Centre or Senior Mode toggle.",
        note="Fast transition to innovations coming next.")

    doc.add_page_break()

    add_heading(doc, "7:20 – 8:50 | Innovation Showcase", level=2)
    add_demo_step(doc, "7:20 – 7:40",
        "Now the innovations. First: Life-Shock Simulator. Most people do not plan for emergencies. Neha can test job loss, medical emergency, or market crash.",
        action="Navigate to Wealth Twin → Life-Shock Simulator. Select 'Job Loss'.",
        note="Emotionally powerful feature.")
    add_demo_step(doc, "7:40 – 7:55",
        "See how net worth drops and goals get delayed? The system then recommends increasing emergency fund or buying term insurance.",
        action="Show impact graph and recommended action.",
        note="Pause for judges to see the value.")
    add_demo_step(doc, "7:55 – 8:15",
        "Second innovation: Generational Wealth Slider. Neha slides from 2026 to 2056 and sees her retirement corpus, inflation impact, and whether she is on track.",
        action="Navigate to Goals/Wealth Twin. Move slider from 2026 to 2056.",
        note="Visual wow moment.")
    add_demo_step(doc, "8:15 – 8:35",
        "Third innovation: CreditBridge AI for MSMEs. Many small businesses have no credit history, so banks reject them. Our AI analyses cash flow, GST, and UPI receipts to generate a credit score and suggest the right PSB loan.",
        action="Navigate to SME Centre → CreditBridge AI. Show credit score and loan recommendation.",
        note="Connects to NMIMS MSME work.")
    add_demo_step(doc, "8:35 – 8:50",
        "We also have NRI Mode, Fantasy League of Wealth for family learning, and Sovereign Vault for long-term secure holdings. This is not a feature list. This is a complete financial operating system.",
        action="Quickly toggle NRI Mode or show feature list.",
        note="End innovation section strongly.")

    doc.add_page_break()

    add_heading(doc, "8:50 – 9:30 | Credibility", level=2)
    add_demo_step(doc, "8:50 – 9:05",
        "What you saw is not just a demo. It is backed by real research. Our Ethical Algorithm work is published in Infinity, ensuring our AI is fair, transparent, and unbiased.",
        action="Switch to slide showing Infinity publication.",
        note="Use slides for credibility.")
    add_demo_step(doc, "9:05 – 9:20",
        "We were finalists at NMIMS Mumbai for our MSME credit work. And technically, this product has 19 frontend tests and 89 backend tests passing, already deployed live on Render.",
        action="Show NMIMS finalist mention and test counts on the slide.",
        note="Builds trust.")
    add_demo_step(doc, "9:20 – 9:30",
        "So the platform works. It is tested. It is deployed. And it is ready for Punjab & Sind Bank.",
        action="Pause. Look at judges.",
        note="Set up the closing.")

    doc.add_page_break()

    # POWERFUL CLOSING
    add_heading(doc, "9:30 – 10:00 | Powerful Closing", level=2)
    add_demo_step(doc, "9:30 – 9:40",
        "With PSB SecureWealth Twin, Punjab & Sind Bank can triple monthly active users, grow FD bookings by fifteen percent, reduce fraud by forty percent, and improve customer satisfaction by twenty points.",
        action="Switch to impact metrics slide.",
        note="End with clear numbers.")
    add_demo_step(doc, "9:40 – 9:50",
        "But beyond the numbers, this is about something bigger. It is about giving every Indian — whether a salaried employee, a small business owner, or a retired parent — a bank that truly understands them.",
        action="Step slightly forward. Make eye contact.",
        note="Emotional peak before final line.")
    add_demo_step(doc, "9:50 – 10:00",
        "We did not just build a banking app. We built a wealth twin. Thank you, judges. I would be honoured to answer your questions.",
        action="Switch to Thank You slide with QR code. Step back and smile.",
        note="Final line must be confident and memorable.")

    doc.add_page_break()

    add_heading(doc, "Why This Intro Works", level=1)
    add_para(doc, "1. It starts with a question — judges engage mentally.")
    add_para(doc, "2. It exposes a daily pain — too many apps, no guidance.")
    add_para(doc, "3. It promises proof — 'I will prove it to you in ten minutes' — creates curiosity.")
    add_para(doc, "4. It introduces the product name and tagline cleanly.")

    add_heading(doc, "Why This Closing Works", level=1)
    add_para(doc, "1. It repeats hard numbers — impact is measurable.")
    add_para(doc, "2. It raises the emotional stakes — 'every Indian' makes it bigger than tech.")
    add_para(doc, "3. It ends with a contrast — 'not just a banking app, a wealth twin' — reinforces differentiation.")
    add_para(doc, "4. It finishes with gratitude and confidence, not apology.")

    add_heading(doc, "Demo Navigation Cheat Sheet", level=1)
    add_para(doc, "1. Login page → Click Neha → AA animation")
    add_para(doc, "2. Dashboard → Net Worth → Asset Breakdown → AI Action Cards")
    add_para(doc, "3. Smart Sweep → Dashboard → Macro Signal Tower")
    add_para(doc, "4. Security Beast → Rakshak AI → Deepfake Voice → Decoy/Ghost")
    add_para(doc, "5. SME Centre → Cash Flow → Working Capital → Surplus Advisor")
    add_para(doc, "6. Wealth Twin → Life-Shock Simulator → Generational Slider")
    add_para(doc, "7. SME Centre → CreditBridge AI")
    add_para(doc, "8. Slides: Credibility → Impact → Thank You")

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
