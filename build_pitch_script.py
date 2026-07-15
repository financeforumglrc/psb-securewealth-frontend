#!/usr/bin/env python3
"""Generate the 10-minute PSB Hackathon pitch script as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "PSB_Hackathon_10Min_Pitch_Script.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    return h


def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_block(doc, label, text, speaker=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if speaker:
        r = p.add_run(f"[{speaker}] ")
        r.bold = True
        r.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
        r.font.size = Pt(11)
    r2 = p.add_run(f"{label}: ")
    r2.bold = True
    r2.font.size = Pt(11)
    r3 = p.add_run(text)
    r3.font.size = Pt(11)
    return p


def add_line(doc, speaker, text, action=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p.add_run(f"{speaker}: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    if action:
        r3 = p.add_run(f"\n    [Action: {action}]")
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r3.font.size = Pt(10)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(f"Note: {text}")
    r.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
    r.font.size = Pt(10)
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PSB SecureWealth Twin\n10-Minute Pitch Script")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    t.paragraph_format.space_after = Pt(12)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("PSB Hackathon 2026 — Internet Banking Reimagined")
    r.font.size = Pt(14)
    r.italic = True
    st.paragraph_format.space_after = Pt(18)

    add_para(doc,
        "This script is designed for a technical + product judge panel, with mostly live demo and a few slides. "
        "Each speaker has exact time cues. Speak slowly, clearly, and transition with confidence. Total time: 10 minutes.",
        italic=True, space_after=12)

    add_heading(doc, "Speaker Order & Roles", level=1)
    add_para(doc, "1. Rikshita Barua — Opening hook + Problem statement (0:00–1:40)")
    add_para(doc, "2. Deepanshu Sharma — Solution 1: Unified Wealth Twin (1:40–2:50)")
    add_para(doc, "3. Mrigesh Mohanty — Solution 2: Security Beast Layer (2:50–4:00)")
    add_para(doc, "4. Ishita Anand — Solution 3: SME Centre + Tax + Inclusion (4:00–5:00)")
    add_para(doc, "5. Tripti Jain — Solution 4: Macro Signals + Compliance (5:00–6:00)")
    add_para(doc, "6. Kunal Saxena — Innovation demo + Credibility + Future scope + Closing (6:00–10:00)")

    doc.add_page_break()

    # Section 1
    add_heading(doc, "0:00 – 1:40 | Opening Hook & Problem Statement", level=1)
    add_block(doc, "Slide", "Slide 1: Title — PSB SecureWealth Twin", speaker="Screen")
    add_line(doc, "Rikshita", "Good morning judges. I am Rikshita, and we are Team DS Financial Solutions.",
             action="Smile, make eye contact with all judges")
    add_line(doc, "Rikshita", "Today, we are not just building another banking app. We are giving every PSB customer a personal wealth twin.")
    add_line(doc, "Rikshita", "Our tagline is simple: \"One Bank. One Twin. Infinite Possibilities.\"")
    add_block(doc, "Slide", "Slide 2: The Problem — A passive bank account", speaker="Screen")
    add_line(doc, "Rikshita", "Let me start with a question. Most of us have a savings account. But does our bank actually help us grow our money?")
    add_line(doc, "Rikshita", "The answer is no. Today, internet banking in India is mostly transactional — check balance, transfer money, pay bills.")
    add_line(doc, "Rikshita", "But customers actually face a much bigger problem. Their money is scattered across multiple accounts, FDs, loans, insurance, and investments. They have no single view.")
    add_line(doc, "Rikshita", "Because of this, three things happen: one, idle money sits in savings accounts and loses value to inflation. Two, users miss tax-saving deadlines. And three, fraudsters exploit this confusion through fake calls and phishing.")
    add_line(doc, "Rikshita", "In short, the problem is this: customers need a bank that does not just store money, but actively protects, grows, and simplifies their financial life.")
    add_line(doc, "Rikshita", "Our solution is PSB SecureWealth Twin. Let me hand over to Deepanshu, who will show you the core platform.")
    add_note(doc, "Rikshita should speak with energy. This is the first impression. Do not rush the tagline.")

    doc.add_page_break()

    # Section 2
    add_heading(doc, "1:40 – 2:50 | Solution 1: Unified Wealth Twin", level=1)
    add_block(doc, "Demo", "Login with demo profile → open Wealth Twin Dashboard", speaker="Screen")
    add_line(doc, "Deepanshu", "Thank you, Rikshita. I am Deepanshu. Let me show you how we solve the scattered-money problem.")
    add_line(doc, "Deepanshu", "This is the Wealth Twin Dashboard. On one screen, the customer can see all bank accounts, FDs, loans, insurance, gold, mutual funds, and real estate — all linked through the Account Aggregator framework.")
    add_line(doc, "Deepanshu", "The net-worth card updates automatically. The asset-liability view shows exactly where the user stands today.")
    add_line(doc, "Deepanshu", "But we did not stop at visibility. We built intelligence on top. The AI Action Cards here suggest real actions — sweep idle surplus into an FD, top up an emergency fund, or rebalance equity exposure.")
    add_line(doc, "Deepanshu", "For example, if the system detects that a customer has two lakh rupees idle in a savings account for thirty days, it recommends a one-tap FD ladder or liquid fund sweep.")
    add_line(doc, "Deepanshu", "So the first solution is: unified view plus proactive recommendations. Next, Mrigesh will show you how we protect all of this.")
    add_note(doc, "Click slowly. Pause after each card so judges can absorb the UI.")

    doc.add_page_break()

    # Section 3
    add_heading(doc, "2:50 – 4:00 | Solution 2: Security Beast Layer", level=1)
    add_block(doc, "Demo", "Navigate to Security Beast / Fraud Protection section", speaker="Screen")
    add_line(doc, "Mrigesh", "Thanks, Deepanshu. I am Mrigesh. Now, when users link so much financial data, security becomes critical.")
    add_line(doc, "Mrigesh", "We built the Security Beast — a multi-layered fraud shield. The first layer is Rakshak AI. It learns the user's normal behaviour — transaction time, amount patterns, device fingerprint, and location.")
    add_line(doc, "Mrigesh", "If something looks wrong, Rakshak intervenes before the money leaves the account. It can soft-block the transaction, ask for biometric re-verification, or directly freeze the account.")
    add_line(doc, "Mrigesh", "The second layer is our Deepfake Voice Shield. In India, voice-cloning fraud is rising. If a fraudster calls pretending to be the customer, our system compares the live voice print against the registered voice print and detects liveness.")
    add_line(doc, "Mrigesh", "We also have Decoy Account and Ghost Mode. If someone is forced to unlock the app under pressure, they can enter a duress PIN. The app opens a fake low-balance account instead. Ghost Mode hides real balances in public places.")
    add_line(doc, "Mrigesh", "So security is not an afterthought — it is built into every action. Ishita will now explain how we serve businesses and taxpayers.")
    add_note(doc, "Use the fraud simulation if available. A live demo of Rakshak blocking a fake transaction has high impact.")

    doc.add_page_break()

    # Section 4
    add_heading(doc, "4:00 – 5:00 | Solution 3: SME Centre + Tax + Inclusion", level=1)
    add_block(doc, "Demo", "Open SME Centre → show Cash Flow Timeline and Working Capital Health", speaker="Screen")
    add_line(doc, "Ishita", "Thank you, Mrigesh. I am Ishita. A bank serves not just individuals, but also lakhs of small businesses.")
    add_line(doc, "Ishita", "Our SME Centre gives business owners three powerful tools. First, a Cash Flow Timeline that shows monthly inflows, outflows, and negative gaps before they become crises.")
    add_line(doc, "Ishita", "Second, a Working Capital Health score with current ratio, quick ratio, and days payable versus receivable. It is colour-coded green, amber, or red.")
    add_line(doc, "Ishita", "Third, the Surplus Fund Advisor. If the business has idle money, it recommends FD sweep, liquid fund, or vendor prepayment — with projected interest or opportunity cost.")
    add_block(doc, "Demo", "Switch to Tax Centre tab", speaker="Screen")
    add_line(doc, "Ishita", "For retail users, we have an Advanced Tax Centre. It compares Old versus New tax regime, tracks Section 80C utilization, and shows a tax deadline calendar so users never miss March 31 or advance-tax dates.")
    add_line(doc, "Ishita", "Finally, inclusion. We have Senior Mode with large fonts and voice prompts, Face Login for users who cannot type easily, and vernacular language support.")
    add_line(doc, "Ishita", "Tripti will now explain how macro trends and compliance fit into this.")

    doc.add_page_break()

    # Section 5
    add_heading(doc, "5:00 – 6:00 | Solution 4: Macro Signals & Compliance", level=1)
    add_block(doc, "Demo", "Open Macro Signal Tower on Dashboard", speaker="Screen")
    add_line(doc, "Tripti", "Thanks, Ishita. I am Tripti. Most banks give advice based only on the customer's data. We also look at the economy.")
    add_line(doc, "Tripti", "Our Macro Signal Tower tracks RBI repo rate, inflation, USD-INR, and gold trends. It then converts these signals into plain-English actions.")
    add_line(doc, "Tripti", "For example, if the repo rate is rising, it suggests shifting to floating-rate FDs. If inflation crosses six percent, it recommends reducing debt duration and increasing equity SIPs. If gold is rising, it suggests booking partial profit and moving to FDs.")
    add_line(doc, "Tripti", "This directly matches the problem statement example: sell gold, shift to FD based on global indicators.")
    add_line(doc, "Tripti", "And because we deal with projections and AI advice, compliance is built-in. Every screen that shows a forecast carries a clear disclaimer: this is a simulation, not a guaranteed return, and the bank is not a SEBI-registered advisor.")
    add_line(doc, "Tripti", "Now, Kunal will take you through the innovations that make this platform truly powerful.")
    add_note(doc, "Speak the example clearly. Judges remember concrete use cases.")

    doc.add_page_break()

    # Section 6
    add_heading(doc, "6:00 – 7:30 | Innovation Showcase (Live Demo)", level=1)
    add_block(doc, "Demo", "Kunal drives the laptop. Show top 4 innovations quickly.", speaker="Screen")
    add_line(doc, "Kunal", "Thank you, Tripti. I am Kunal. We have solved the core problem. Now let me show you the innovations that make PSB SecureWealth Twin stand out.")
    add_line(doc, "Kunal", "First, the Life-Shock Simulator. Users can stress-test their finances against job loss, medical emergency, market crash, or death. It shows the impact on goals and net worth, and recommends insurance or contingency actions.")
    add_line(doc, "Kunal", "Second, Smart Sweep. Using the Account Aggregator framework, it finds idle balances across linked accounts and recommends one-tap sweep into higher-yield instruments.")
    add_line(doc, "Kunal", "Third, the Generational Wealth Slider. A user can slide across decades and see how their corpus grows, how inflation erodes purchasing power, and whether their retirement goal is on track.")
    add_line(doc, "Kunal", "Fourth, CreditBridge AI for MSMEs. It analyses cash-flow velocity, GST turnover, and UPI receipts to generate a credit health score and suggest the right PSB loan product — even for businesses without formal credit history.")
    add_line(doc, "Kunal", "We also have NRI Mode, Sovereign Vault, Fantasy League of Wealth for financial literacy, and a Collective Immune System that warns all users when a new fraud pattern is detected.")
    add_note(doc, "Demo speed matters here. Spend ~20 seconds per innovation. Do not go too deep.")

    doc.add_page_break()

    # Section 7
    add_heading(doc, "7:30 – 8:10 | Credibility & Validation", level=1)
    add_block(doc, "Slide", "Slide 3: Traction & Recognition", speaker="Screen")
    add_line(doc, "Kunal", "Judges, what I just showed is not just a prototype. It is backed by real research and validation.")
    add_line(doc, "Kunal", "Our Ethical Algorithm work has been published in Infinity. This ensures that our AI recommendations are fair, transparent, and free from bias against any customer segment.")
    add_line(doc, "Kunal", "We were also finalists at NMIMS Mumbai for our MSME-focused work, specifically on credit access for small businesses using surrogate data like cash flow and GST patterns.")
    add_line(doc, "Kunal", "Today, PSB SecureWealth Twin has 19 frontend tests and 89 backend tests passing, and it is already deployed on Render with a live backend and frontend.")
    add_note(doc, "Say this confidently but not arrogantly. Credibility builds trust.")

    doc.add_page_break()

    # Section 8
    add_heading(doc, "8:10 – 9:20 | Future Scope & Impact", level=1)
    add_block(doc, "Slide", "Slide 4: Roadmap & Impact Metrics", speaker="Screen")
    add_line(doc, "Kunal", "Looking ahead, our roadmap is clear.")
    add_line(doc, "Kunal", "Phase 1: Pilot with one PSB branch — one thousand retail customers and one hundred SMEs. Phase 2: Integrate live Account Aggregator partners like PhonePe and Finvu. Phase 3: Fine-tune our recommendation models on anonymized PSB data. Phase 4: Full rollout across all PSB internet banking users.")
    add_line(doc, "Kunal", "The impact we expect is significant: three times increase in monthly active users, fifteen percent growth in FD and recurring deposit bookings, forty percent reduction in reported digital fraud, and a twenty-point improvement in digital banking NPS.")
    add_line(doc, "Kunal", "Beyond numbers, this advances financial inclusion. It protects vulnerable users from fraud, supports SMEs with working capital intelligence, and helps every PSB customer make smarter financial decisions.")

    doc.add_page_break()

    # Section 9
    add_heading(doc, "9:20 – 10:00 | Closing", level=1)
    add_block(doc, "Slide", "Slide 5: Thank You + QR / Links", speaker="Screen")
    add_line(doc, "Kunal", "To conclude, PSB SecureWealth Twin transforms internet banking from a transaction tool into a trusted wealth companion.")
    add_line(doc, "Kunal", "It solves the scattered-money problem. It protects users with intelligent security. It serves individuals, SMEs, and NRIs. It is compliant, inclusive, and demonstrable today.")
    add_line(doc, "Kunal", "We are ready to deploy, scale, and win this for Punjab & Sind Bank.")
    add_line(doc, "Kunal", "Thank you, judges. We would be happy to take your questions.", action="Team steps forward, smiles, ready for Q&A")
    add_note(doc, "End exactly at 10:00. If running early, slow down the closing. If running late, skip one optional innovation.")

    doc.add_page_break()

    # Tips page
    add_heading(doc, "Delivery Tips for the Team", level=1)
    add_para(doc, "1. Practice with a timer. The demo portions are the easiest to overrun.")
    add_para(doc, "2. Rikshita and Kunal have the highest-impact moments — opening and closing. Rehearse these the most.")
    add_para(doc, "3. During demo transitions, the speaker should face judges, not the screen. One teammate can click if needed.")
    add_para(doc, "4. If the website is slow on stage, have screenshots as backup. Do not panic.")
    add_para(doc, "5. Use words like 'we built', 'we solved', 'we tested' — it shows ownership.")
    add_para(doc, "6. For technical judges, mention the test count and deployment. For business judges, emphasize impact metrics.")
    add_para(doc, "7. Speak in short sentences. Avoid jargon. Say 'idle money' instead of 'uninvested surplus'.")
    add_para(doc, "8. Everyone should know the full script so you can cover if someone forgets a line.")

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
