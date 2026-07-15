#!/usr/bin/env python3
"""Generate a 15-page PSB SecureWealth Twin problem+solutions+innovations document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "PSB_SecureWealth_Problem_Statement_Solutions_and_Innovations.docx"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    return h


def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "0B61A4")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PSB SecureWealth Twin")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x0B, 0x61, 0xA4)
    title.paragraph_format.space_after = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Problem Statement, Solutions & Innovation Blueprint")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    subtitle.paragraph_format.space_after = Pt(36)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Punjab & Sind Bank — PSB Hackathon 2026\nDS Financial Solutions")
    run.font.size = Pt(14)
    run.font.italic = True
    meta.paragraph_format.space_after = Pt(12)

    add_para(doc, "", space_after=24)

    summary_box = doc.add_paragraph()
    summary_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary_box.add_run(
        "This document maps every requirement of the PSB Hackathon problem statement to a concrete solution, "
        "then enumerates the cutting-edge innovations built into the SecureWealth Twin platform."
    )
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "Public Sector Banks (PSBs) in India serve hundreds of millions of retail, SME, and rural customers. "
        "Yet most digital banking experiences remain transaction-oriented: check balance, transfer funds, pay bills. "
        "The PSB Hackathon 2026 asks participants to reimagine internet banking as a proactive, intelligent, "
        "and inclusive wealth companion.")
    add_para(doc,
        "PSB SecureWealth Twin is our answer. It is a full-stack, AI-native internet banking platform that transforms "
        "a PSB account from a passive ledger into a dynamic \"wealth twin\" — a digital mirror of the customer's "
        "financial life that thinks ahead, protects against fraud, optimizes taxes, and nudges behaviour toward "
        "long-term prosperity.")
    add_para(doc,
        "The platform is engineered around five design principles that directly address the problem statement:")
    add_bullet(doc, "Unified View: All accounts, goals, assets, liabilities, and cash flows in one dashboard.")
    add_bullet(doc, "Proactive Intelligence: AI-driven recommendations, life-shock simulations, and macro signals.")
    add_bullet(doc, "Trust & Safety: Deepfake voice shield, behavioural biometrics, decoy accounts, and Rakshak AI.")
    add_bullet(doc, "Inclusion: Senior mode, vernacular UI, accessibility, and simplified journeys for non-tech users.")
    add_bullet(doc, "Compliance First: Every projection carries a regulatory disclaimer; no guaranteed returns are shown.")

    add_heading(doc, "1.1 Document Structure", level=2)
    add_para(doc,
        "Section 2 deconstructs the hackathon problem statement into atomic requirements. "
        "Section 3 maps each requirement to a specific product or technical solution. "
        "Section 4 details the innovations built beyond the baseline. "
        "Section 5 covers architecture, security, and deployment. "
        "Section 6 outlines the roadmap and impact metrics.")

    doc.add_page_break()

    # 2. Problem Statement Deconstruction
    add_heading(doc, "2. Problem Statement Deconstruction", level=1)
    add_para(doc,
        "The PSB Hackathon 2026 problem statement for Internet Banking can be summarized as follows: "
        "\"Reimagine the public sector banking experience by building a secure, intuitive, and intelligent "
        "digital platform that empowers customers to manage, grow, and protect their wealth, while ensuring "
        "accessibility, compliance, and trust.\"")

    add_heading(doc, "2.1 Atomic Requirement Map", level=2)
    reqs = [
        ("R1", "Unified Financial View", "Customers hold multiple accounts, FDs, loans, insurance policies, and investments across institutions. They need a single pane of glass."),
        ("R2", "Personalized Insights", "Generic dashboards do not guide behaviour. The platform must offer contextual, actionable advice."),
        ("R3", "Security & Fraud Prevention", "Rising UPI fraud, deepfake calls, and phishing demand multi-layered, intelligent protection."),
        ("R4", "Goal-Based Planning", "Users need to visualize retirement, education, home purchase, and emergency corpus goals."),
        ("R5", "Tax & Regulatory Optimization", "Timely tax-saving suggestions, regime comparison, and deadline awareness are expected."),
        ("R6", "SME & Business Banking", "Micro, small, and medium enterprises need cash-flow visibility, working-capital alerts, and surplus-fund advice."),
        ("R7", "Accessibility & Inclusion", "Elderly, rural, and differently-abled users must be able to use the platform confidently."),
        ("R8", "Macro-Aware Recommendations", "Interest rates, inflation, currency, and gold trends should influence advice."),
        ("R9", "Compliance & Disclaimers", "All projections are simulations; no guaranteed returns; SEBI/bank licensing disclaimers must be visible."),
        ("R10", "Demo-Ready Presentation", "The solution must be demonstrable end-to-end with realistic data and clear user journeys."),
    ]
    add_table(doc, ["ID", "Requirement", "Pain Point Addressed"], reqs)

    add_heading(doc, "2.2 Stakeholder View", level=2)
    add_para(doc,
        "The problem is multi-stakeholder. Retail customers want simplicity and growth. SMEs want liquidity control. "
        "Bank management wants lower cost-to-serve and higher trust. Regulators want transparency and data protection. "
        "Our solution is designed to satisfy all four without compromising any one dimension.")

    doc.add_page_break()

    # 3. Solutions Mapping
    add_heading(doc, "3. Requirement-to-Solution Mapping", level=1)
    add_para(doc,
        "This section maps every atomic requirement to the exact module, component, or API that fulfills it. "
        "All solutions are live in the working demo unless explicitly marked as future enhancement.")

    add_heading(doc, "3.1 R1 — Unified Financial View", level=2)
    add_para(doc,
        "Problem: Customers struggle to see a consolidated picture across savings accounts, FDs, loans, insurance, "
        "mutual funds, gold, real estate, and liabilities.")
    add_para(doc, "Solution: Wealth Twin Dashboard", bold=True)
    add_bullet(doc, "Account Aggregator widget links bank accounts, credit cards, and external holdings (mock AA flow).")
    add_bullet(doc, "Net-worth summary with real-time asset-liability reconciliation.")
    add_bullet(doc, "Holdings breakdown by asset class, liquidity bucket, and risk profile.")
    add_bullet(doc, "Transaction history with intelligent categorization and search.")
    add_bullet(doc, "Backend SQLite/Postgres adapter stores user data; API layer serves it via REST.")

    add_heading(doc, "3.2 R2 — Personalized Insights", level=2)
    add_para(doc,
        "Problem: Customers receive generic product pitches instead of contextual nudges.")
    add_para(doc, "Solution: Agentic AI Action Cards + Bhavishya Engine", bold=True)
    add_bullet(doc, "AI Action Cards surface one-tap recommendations: sweep surplus to FD, top-up emergency fund, rebalance equity.")
    add_bullet(doc, "Bhavishya Engine projects future net worth under multiple scenarios.")
    add_bullet(doc, "Behavioural nudges encourage saving, tax investment, and goal contribution.")
    add_bullet(doc, "Risk-profile-driven suggestions prevent unsuitable products.")

    add_heading(doc, "3.3 R3 — Security & Fraud Prevention", level=2)
    add_para(doc,
        "Problem: Digital fraud is evolving faster than static OTPs and passwords can handle.")
    add_para(doc, "Solution: Security Beast Layer", bold=True)
    add_bullet(doc, "Rakshak AI intervenes when transactions deviate from behavioural baselines.")
    add_bullet(doc, "Deepfake Voice Shield verifies caller identity during high-risk actions.")
    add_bullet(doc, "Behavioural Biometrics tracks typing rhythm, mouse patterns, and device fingerprint.")
    add_bullet(doc, "Decoy Account confuses attackers with a fake low-balance persona.")
    add_bullet(doc, "Ghost Mode hides real balances in public or shared-screen situations.")
    add_bullet(doc, "Duress/Panic PIN silently alerts authorities while appearing to work normally.")

    add_heading(doc, "3.4 R4 — Goal-Based Planning", level=2)
    add_para(doc,
        "Problem: Users lack a visual, interactive way to plan for life goals.")
    add_para(doc, "Solution: Goal Tracker + Life-Shock Simulator", bold=True)
    add_bullet(doc, "Goal Tracker lets users define retirement, education, home, travel, and emergency goals.")
    add_bullet(doc, "Monte Carlo simulator estimates probability of goal achievement.")
    add_bullet(doc, "Life-Shock Simulator models job loss, medical emergency, market crash, and death scenarios.")
    add_bullet(doc, "Generational Wealth Slider shows corpus projection across decades.")

    doc.add_page_break()

    add_heading(doc, "3.5 R5 — Tax & Regulatory Optimization", level=2)
    add_para(doc,
        "Problem: Tax planning is reactive and deadline-driven; users miss 80C, 80D, and advance-tax windows.")
    add_para(doc, "Solution: Advanced Tax Centre", bold=True)
    add_bullet(doc, "Old vs New Regime calculator with live tax computation backend.")
    add_bullet(doc, "Section 80C tracker with risk-profile-based ELSS/PPF/NPS suggestions.")
    add_bullet(doc, "Tax deadline calendar with urgency badges and one-tap actions.")
    add_bullet(doc, "Tax-loss harvesting alerts for equity holdings.")

    add_heading(doc, "3.6 R6 — SME & Business Banking", level=2)
    add_para(doc,
        "Problem: SMEs use personal banking tools for business cash flow because PSB portals lack working-capital intelligence.")
    add_para(doc, "Solution: SME Centre", bold=True)
    add_bullet(doc, "Cash Flow Timeline visualizes monthly inflows, outflows, and negative gaps.")
    add_bullet(doc, "Surplus Fund Advisor recommends FD sweep, liquid funds, or vendor prepayment.")
    add_bullet(doc, "Working Capital Health score tracks current ratio, quick ratio, and days payable/receivable.")
    add_bullet(doc, "Export to PDF/Excel for accountant hand-off.")

    add_heading(doc, "3.7 R7 — Accessibility & Inclusion", level=2)
    add_para(doc,
        "Problem: Elderly and rural customers find modern banking apps overwhelming.")
    add_para(doc, "Solution: Adaptive UI Modes", bold=True)
    add_bullet(doc, "Senior Mode with larger fonts, high contrast, simplified navigation, and voice prompts.")
    add_bullet(doc, "Vernacular support framework (Hindi, Punjabi, Tamil, etc.) via i18n store.")
    add_bullet(doc, "Face Login for users who forget passwords or cannot type easily.")
    add_bullet(doc, "Accessibility settings for screen-reader labels, motion reduction, and dyslexia-friendly fonts.")

    add_heading(doc, "3.8 R8 — Macro-Aware Recommendations", level=2)
    add_para(doc,
        "Problem: Advice ignores the macroeconomic context, e.g., rising repo rates or falling rupee.")
    add_para(doc, "Solution: Macro Signal Tower", bold=True)
    add_bullet(doc, "Live mock feed of RBI repo rate, CPI inflation, USD/INR, and gold trend.")
    add_bullet(doc, "Auto-triggered rules: rising repo → floating-rate FD; inflation >6% → reduce debt duration; falling rupee → trim gold.")
    add_bullet(doc, "Integrated into Dashboard and Wealth Twin tabs.")

    add_heading(doc, "3.9 R9 — Compliance & Disclaimers", level=2)
    add_para(doc,
        "Problem: Projection screens can be misinterpreted as guaranteed advice.")
    add_para(doc, "Solution: Regulatory Disclaimer Component", bold=True)
    add_bullet(doc, "Reusable disclaimer appears on every projection, recommendation, tax, and AI screen.")
    add_bullet(doc, "Global footer clarifies: simulation only, not SEBI-registered, no guaranteed returns.")
    add_bullet(doc, "All demo data is synthetic; no real customer PII is used.")

    add_heading(doc, "3.10 R10 — Demo-Ready Presentation", level=2)
    add_para(doc,
        "Problem: A great backend is invisible unless the demo tells a compelling story.")
    add_para(doc, "Solution: Cinematic Demo Journeys", bold=True)
    add_bullet(doc, "Pre-seeded demo profiles: Neha (salaried), Rahul (business), Arjun (retired), Priya (NRI).")
    add_bullet(doc, "One-click demo login with AA fetch animation.")
    add_bullet(doc, "Pitch deck, demo playbook, cue card, and screenshots auto-generated from the codebase.")

    doc.add_page_break()

    # 4. Innovations
    add_heading(doc, "4. Innovation Catalogue", level=1)
    add_para(doc,
        "Beyond the baseline requirements, PSB SecureWealth Twin introduces several first-of-its-kind features "
        "for a public sector banking context. These are the differentiators that elevate the solution from "
        "a dashboard to a financial operating system.")

    add_heading(doc, "4.1 Wealth Twin — Digital Financial Mirror", level=2)
    add_para(doc,
        "The Wealth Twin is not a static report. It is a living model of the customer's financial universe. "
        "It aggregates accounts, assets, liabilities, goals, and cash flows; then runs simulations to answer "
        "\"what-if\" questions in real time. It is the central metaphor of the product: just as a doctor uses a "
        "scan to diagnose health, the Wealth Twin diagnoses financial health.")
    add_bullet(doc, "Real-time net-worth tracking across asset classes.")
    add_bullet(doc, "Scenario modelling: job loss, medical emergency, market correction, inheritance.")
    add_bullet(doc, "Goal achievement probability via Monte Carlo simulation.")

    add_heading(doc, "4.2 Rakshak AI — Autonomous Fraud Intervention", level=2)
    add_para(doc,
        "Rakshak AI is an autonomous agent that monitors transactions, device behaviour, and voice biometric signals. "
        "When it detects anomalies — an unusual merchant, a new device at 3 AM, a deepfake voice pattern — it intervenes "
        "before money leaves the account.")
    add_bullet(doc, "Behavioural baseline per user.")
    add_bullet(doc, "Real-time risk score with explainable AI reasoning.")
    add_bullet(doc, "Interventions: soft block, biometric re-auth, call verification, freeze.")

    add_heading(doc, "4.3 Deepfake Voice Shield", level=2)
    add_para(doc,
        "Voice cloning fraud is rising in India. The Deepfake Voice Shield adds a voice-print challenge to high-risk "
        "actions. The user speaks a prompt; the system compares the voice print, liveness, and phrase match before "
        "authorizing the transaction.")
    add_bullet(doc, "On-device liveness detection to prevent replay attacks.")
    add_bullet(doc, "Phrases rotate per session.")
    add_bullet(doc, "Fallback to passkey or face login if voice fails.")

    add_heading(doc, "4.4 Smart Sweep — Account Aggregator Arbitrage", level=2)
    add_para(doc,
        "Smart Sweep uses the Account Aggregator framework to find idle balances across linked accounts and sweep them "
        "into higher-yield instruments automatically. It solves the \"money sleeping in savings account\" problem.")
    add_bullet(doc, "Detects idle surplus across linked banks.")
    add_bullet(doc, "Recommends FD ladder, liquid fund, or overnight fund.")
    add_bullet(doc, "One-tap consent and execution.")

    add_heading(doc, "4.5 Life-Shock Simulator", level=2)
    add_para(doc,
        "Users rarely plan for shocks. The Life-Shock Simulator lets them stress-test their finances against job loss, "
        "critical illness, disability, market crash, and death. It visualizes the impact on goals and net worth, "
        "then suggests insurance or contingency actions.")

    add_heading(doc, "4.6 Macro Signal Tower", level=2)
    add_para(doc,
        "The Macro Signal Tower translates RBI policy, inflation, currency, and commodity trends into plain-English "
        "portfolio actions. For example: \"Repo rate rising — shift 20% of debt to floating-rate FD.\"")

    add_heading(doc, "4.7 Decoy Account & Ghost Mode", level=2)
    add_para(doc,
        "These features protect users under coercion or shoulder surfing. Decoy Account shows a fake, low-balance "
        "persona when a duress PIN is entered. Ghost Mode masks real balances in public places.")

    add_heading(doc, "4.8 NRI Mode & Sovereign Vault", level=2)
    add_para(doc,
        "NRI customers get a dedicated mode with FEMA-aware remittance tracking, NRE/NRO bifurcation, and tax residency "
        "insights. Sovereign Vault is a high-security sub-account for large, long-term holdings.")

    add_heading(doc, "4.9 CreditBridge AI — MSME Lending", level=2)
    add_para(doc,
        "CreditBridge AI connects micro-entrepreneurs to formal credit by analysing surrogate data: cash-flow velocity, "
        "GST turnover, UPI receipts, and invoice patterns. It generates a credit health score and suggests the right "
        "PSB loan product.")

    add_heading(doc, "4.10 Fantasy League of Wealth", level=2)
    add_para(doc,
        "A gamified layer where family members or friends can create virtual portfolios and compete on risk-adjusted "
        "returns. This drives financial literacy through friendly competition without risking real money.")

    doc.add_page_break()

    add_heading(doc, "4.11 Neuro-Friction Widget", level=2)
    add_para(doc,
        "Inspired by behavioural economics, this widget introduces intentional friction for impulsive transactions "
        "(large transfers, new merchant payments) while reducing friction for healthy habits (SIP top-ups, goal contributions).")

    add_heading(doc, "4.12 Autonomous Agent Mode", level=2)
    add_para(doc,
        "A future-facing module where users can delegate routine tasks — bill payments, rebalancing, tax harvesting — "
        "to an AI agent operating within guardrails. Every action is logged, reversible, and capped by user-defined limits.")

    add_heading(doc, "4.13 Social Collateral Loan", level=2)
    add_para(doc,
        "Enables community-backed micro-lending within trusted groups. Repayment behaviour is recorded on a private "
        "ledger and can improve the borrower's formal credit score over time.")

    add_heading(doc, "4.14 Collective Immune System", level=2)
    add_para(doc,
        "Anonymized fraud patterns from the entire user base feed a collective model. If a new phishing domain targets "
        "one customer, all customers receive a pre-emptive warning.")

    add_heading(doc, "4.15 Parametric Insurance", level=2)
    add_para(doc,
        "Weather-indexed crop insurance and travel-disruption insurance triggered automatically by verified external "
        "data feeds, reducing claim processing time from weeks to minutes.")

    add_heading(doc, "4.16 Cooling Vault / Dead Man's Switch", level=2)
    add_para(doc,
        "Cooling Vault delays large withdrawals by a user-defined cooling-off period, preventing panic-driven or "
        "coerced losses. Dead Man's Switch transfers assets to nominees if the user fails to check in for a configured period.")

    doc.add_page_break()

    # 5. Architecture & Security
    add_heading(doc, "5. Architecture, Security & Deployment", level=1)
    add_para(doc,
        "The solution is a modern full-stack application designed for hackathon demos and production scalability alike.")

    add_heading(doc, "5.1 Tech Stack", level=2)
    add_table(doc, ["Layer", "Technology", "Purpose"], [
        ("Frontend", "React 19 + TypeScript + Vite", "Responsive, high-performance SPA"),
        ("Styling", "Tailwind CSS v4 + Framer Motion", "Design system and animations"),
        ("State", "Zustand + React Context", "Global and auth state management"),
        ("Backend", "Node.js + Express", "REST API and business logic"),
        ("Database", "SQLite (demo) / PostgreSQL (prod)", "User data, transactions, goals"),
        ("Auth", "Supabase Auth + JWT", "Secure session management"),
        ("AI", "OpenAI / Anthropic / Gemini SDKs", "Recommendations and natural language"),
        ("BI/Charts", "Recharts", "Visualizations and projections"),
        ("Deployment", "Render (Docker + Static)", "Backend service and static frontend"),
        ("Testing", "Vitest + Playwright + Jest", "Unit, component, and E2E tests"),
    ])

    add_heading(doc, "5.2 Security Architecture", level=2)
    add_bullet(doc, "JWT-based authentication with short-lived access tokens.")
    add_bullet(doc, "Rate limiting and Helmet headers on all API routes.")
    add_bullet(doc, "Passwords hashed with bcryptjs.")
    add_bullet(doc, "Sensitive env vars (JWT_SECRET, API keys) managed via Render dashboard.")
    add_bullet(doc, "Device fingerprinting and behavioural biometrics for anomaly detection.")
    add_bullet(doc, "ML-KEM post-quantum key encapsulation for future-proofing.")

    add_heading(doc, "5.3 Deployment Model", level=2)
    add_para(doc,
        "The application is deployed on Render using a Blueprint defined in render.yaml. The backend runs as a Docker "
        "service (psb-securewealth-backend) and the frontend as a static site (psb-securewealth-frontend). "
        "Continuous deployment is enabled: every push to GitHub main triggers a new Render build.")
    add_bullet(doc, "Backend URL: https://psb-securewealth-backend.onrender.com/api/v1")
    add_bullet(doc, "Frontend URL: https://psb-securewealth-frontend.onrender.com")

    doc.add_page_break()

    # 6. Roadmap & Impact
    add_heading(doc, "6. Roadmap, Metrics & Impact", level=1)
    add_para(doc,
        "The platform is demo-ready today. The following roadmap takes it from hackathon prototype to production pilot.")

    add_heading(doc, "6.1 Phased Roadmap", level=2)
    add_numbered(doc, "Phase 0 — Demo Hardening: Fix UI/UX gaps, ensure disclaimers on all screens, and optimize demo login flow.")
    add_numbered(doc, "Phase 1 — Pilot with PSB Branch: Onboard 1,000 retail customers and 100 SMEs in a single branch.")
    add_numbered(doc, "Phase 2 — Account Aggregator Integration: Connect with live AA ecosystem partners (PhonePe, Finvu, etc.).")
    add_numbered(doc, "Phase 3 — AI Model Fine-Tuning: Train recommendation models on anonymized PSB transaction data.")
    add_numbered(doc, "Phase 4 — Full Rollout: Scale to all PSB internet banking users with regional language support.")

    add_heading(doc, "6.2 Success Metrics", level=2)
    add_table(doc, ["Metric", "Target"], [
        ("Customer engagement", "3x increase in monthly active internet banking users"),
        ("Savings mobilization", "15% increase in FD/recurring deposit bookings"),
        ("Fraud reduction", "40% reduction in reported digital fraud incidents"),
        ("SME adoption", "500+ businesses using SME Centre in pilot quarter"),
        ("NPS improvement", "+20 point uplift in digital banking NPS"),
        ("Support cost", "25% reduction in branch support tickets for balance/fund queries"),
    ])

    add_heading(doc, "6.3 Social Impact", level=2)
    add_para(doc,
        "By making wealth management accessible to non-English, non-tech-savvy users, PSB SecureWealth Twin advances "
        "financial inclusion. By protecting vulnerable users from fraud, it builds trust in public sector banking. "
        "By serving SMEs with working-capital intelligence, it supports the real economy at the grassroots level.")

    add_heading(doc, "6.4 Conclusion", level=2)
    add_para(doc,
        "PSB SecureWealth Twin is more than an internet banking upgrade. It is a reimagining of the relationship "
        "between a public sector bank and its customers — from transactional to transformational. Every requirement "
        "of the hackathon problem statement has a mapped solution, and every solution is enriched by innovations "
        "that are practical, demonstrable, and scalable. We are ready to deploy, demonstrate, and win.")

    # Save
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
