# -*- coding: utf-8 -*-
"""One-page printable demo cue card for the screen driver."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0B, 0x2B, 0x52)
RED  = RGBColor(0xB3, 0x1B, 0x1B)
GREEN= RGBColor(0x1B, 0x7A, 0x3D)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE= RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.4); sec.bottom_margin = Inches(0.4)
sec.left_margin = Inches(0.5); sec.right_margin = Inches(0.5)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(9)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexcolor); tcPr.append(sh)

# Title
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(1)
r = p.add_run('SECUREWEALTH TWIN — 10-MIN DEMO CUE CARD'); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=NAVY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
r = p.add_run('Driver: follow top to bottom. Spend the MOST time on the red block.'); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=GREY

# Pitch box
t = doc.add_table(rows=1, cols=1); t.style='Table Grid'
c = t.rows[0].cells[0]; shade(c,'EAF1FB')
pp = c.paragraphs[0]; pp.paragraph_format.space_after=Pt(1)
rr = pp.add_run('OPENING LINE: '); rr.bold=True; rr.font.size=Pt(9); rr.font.color.rgb=NAVY
rr2 = pp.add_run('"SecureWealth Twin grows wealth intelligently, and a mandatory cyber-protection layer guards every critical money action — Allow, Warn, or Block."'); rr2.italic=True; rr2.font.size=Pt(9)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

rows = [
 ('0:45','HOOK + ARCH', 'NAVY',
   'Open live app + 1 architecture slide',
   '"This is our flow: Twin -> Protection Check -> Decision -> Action -> Audit."'),
 ('3:00','WEALTH TWIN', 'NAVY',
   'Dashboard/NetWorth > AccountAggregator > Add Asset (gold) > GoalTracker > AIRecommendations',
   '"The twin builds a full picture and gives advice WITH a reason."'),
 ('6:30','PROTECTION LAYER  (SPEND MOST TIME)', 'RED',
   'Start high-value action > DeviceStatus(new) > <10s flag > RiskMeter(amount) > OTP retry > SCORE Low/Med/High > CoolingVault/Guard > DURESS PIN (fake success+lock) > AuditLog',
   '"Each signal adds weight -> risk score -> Allow / Warn-cooling-off / Block. Duress PIN fakes success and locks the account."'),
 ('8:00','RESPONSIBLE AI', 'NAVY',
   'ExplainableTooltip/AIDecisionLog > ConsentModal/Privacy > KYCModal > "Demo only" badge',
   '"Consent-first, explainable, KYC before invest, no false claims."'),
 ('9:00','INNOVATION (60s flash)', 'NAVY',
   'InnovationLab: BhavishyaEngine/CrisisPredictor, FutureSelfSimulator, BusinessMode',
   '"40+ modules — lightweight by default, deep on demand."'),
 ('10:00','SCALE + CLOSE', 'GREEN',
   'One line: React + Node microservice + Docker + AA + mock CBS',
   '"Wealth created AND safeguarded — exactly the dual mandate."'),
]
colormap = {'NAVY':'0B2B52','RED':'B31B1B','GREEN':'1B7A3D'}
tbl = doc.add_table(rows=1, cols=4); tbl.style='Table Grid'
hdr = tbl.rows[0].cells
for i,h in enumerate(['@','BLOCK','CLICK ORDER (left to right)','SAY THIS']):
    shade(hdr[i],'0B2B52'); pr=hdr[i].paragraphs[0]; rn=pr.add_run(h); rn.bold=True; rn.font.color.rgb=WHITE; rn.font.size=Pt(8.5)
for time,block,clr,clicks,say in rows:
    cells = tbl.add_row().cells
    # time
    pr=cells[0].paragraphs[0]; rn=pr.add_run(time); rn.bold=True; rn.font.size=Pt(8.5)
    # block
    shade(cells[1], colormap[clr])
    pr=cells[1].paragraphs[0]; rn=pr.add_run(block); rn.bold=True; rn.font.color.rgb=WHITE; rn.font.size=Pt(8.5)
    # clicks
    pr=cells[2].paragraphs[0]; rn=pr.add_run(clicks); rn.font.size=Pt(8.5)
    # say
    pr=cells[3].paragraphs[0]; rn=pr.add_run(say); rn.italic=True; rn.font.size=Pt(8.5)
# column widths
for row in tbl.rows:
    row.cells[0].width=Inches(0.5); row.cells[1].width=Inches(1.35)
    row.cells[2].width=Inches(3.4); row.cells[3].width=Inches(3.0)

doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Bottom: DO NOT + roles
b = doc.add_table(rows=1, cols=2); b.style='Table Grid'
left = b.rows[0].cells[0]; right = b.rows[0].cells[1]
shade(left,'FBE9E9')
pl=left.paragraphs[0]; rl=pl.add_run('DO NOT OPEN: '); rl.bold=True; rl.font.color.rgb=RED; rl.font.size=Pt(8.5)
left.add_paragraph('- Scenario/NLP what-if (can throw)', style=None)
left.add_paragraph('- Exact tax rupee figures (Tax view)')
left.add_paragraph('- GST/ITC/DCF/LBO tools (off-topic)')
left.add_paragraph('- "47 patents / military-grade / unbreakable"')
for pp in left.paragraphs[1:]:
    pp.runs[0].font.size=Pt(8.5); pp.paragraph_format.space_after=Pt(0)
shade(right,'E9F4EC')
prr=right.paragraphs[0]; rr=prr.add_run('IF A FEATURE BREAKS: '); rr.bold=True; rr.font.color.rgb=GREEN; rr.font.size=Pt(8.5)
right.add_paragraph('Stay calm: "That module is a work-in-progress — let me show the core protection flow," then go to the red block.')
right.add_paragraph('Have open in tabs: timingCheck.js + fraudDetectionService.ts (for "show me the code").')
for pp in right.paragraphs[1:]:
    pp.runs[0].font.size=Pt(8.5); pp.paragraph_format.space_after=Pt(1)
left.width=Inches(4.3); right.width=Inches(4.3)

doc.save(r'G:/PSB/DS_Financial/frontend-render/SecureWealth_Twin_Cue_Card.docx')
print('SAVED: SecureWealth_Twin_Cue_Card.docx')
