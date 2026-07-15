# -*- coding: utf-8 -*-
"""Generates a Word document that can be used as source material for a PPT."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

OUTPUT = "SecureWealth_Twin_PPT_Source.docx"
DIAGRAM = "securewealth_architecture.png"

NAVY = RGBColor(0x0B, 0x2B, 0x52)
BLUE = RGBColor(0x1E, 0x5A, 0x9C)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
RED = RGBColor(0xB3, 0x1B, 0x1B)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(10.5)
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def H1(text, color=NAVY):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.font.size = Pt(20)
    r.bold = True
    return p


def H2(text, color=BLUE):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.font.size = Pt(14)
    r.bold = True
    return p


def H3(text, color=NAVY):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.font.size = Pt(12)
    r.bold = True
    return p


def para(text, size=10.5, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.add_run(text)
    return p


def table(headers, rows, widths=None, header_fill='0B2B52'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], header_fill)
        pr = hdr[i].paragraphs[0]
        rn = pr.add_run(h)
        rn.bold = True
        rn.font.color.rgb = WHITE
        rn.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            pr = cells[i].paragraphs[0]
            rn = pr.add_run(str(val))
            rn.font.size = Pt(9)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    return t


def add_code_block(path, lines=35):
    if not os.path.exists(path):
        doc.add_paragraph("(file not found)")
        return
    with open(path, 'r', encoding='utf-8') as f:
        code = ''.join(f.readlines()[:lines])
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(code)
    r.font.name = 'Courier New'
    r.font.size = Pt(8)
    r.font.color.rgb = GREY


# ===================== DOCUMENT CONTENT =====================

# Title
para("SecureWealth Twin — PPT Source Document", size=22, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Prepared for PSB Hackathon Series 2026", size=12, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Audience: Technical Judges / Hackathon Evaluators", size=11, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# 1. Problem Statement
H1("1. Problem Statement")
para("The bank has a dual responsibility: help customers grow wealth intelligently, and protect that wealth from fraud, misuse, and risky actions.")
bullet("Customers increasingly use digital channels for savings, investments, and financial planning.")
bullet("AI-powered wealth advisory can guide better outcomes, but digital convenience exposes users to social-engineering fraud, coerced transactions, OTP misuse, and impulsive high-value actions.")
bullet("The problem statement asks for a unified system: an AI-powered digital wealth intelligence system that acts as a virtual financial twin, embedding a mandatory, lightweight cyber-security and fraud protection layer around every critical wealth action.")
doc.add_paragraph()

# 2. Solution Overview
H1("2. Solution Overview")
para("SecureWealth Twin is a full-stack web application that combines:")
bullet("Digital Wealth Intelligence — learns from spending/saving/investment habits, understands income/goal/risk changes, studies market trends, and gives personalized, explainable recommendations.")
bullet("Mandatory Wealth Protection Layer — evaluates risk signals (device trust, OTP patterns, amount vs. history, speed/retries) and generates a Wealth Protection Risk Score that leads to Allow / Warn / Block decisions.")
bullet("Account Aggregator Integration — consent-based multi-bank data pull for a complete financial picture.")
bullet("Responsible AI & Compliance — consent-first data handling, explainable recommendations, KYC gating, and no misleading claims.")
doc.add_paragraph()

# 3. Live Deployment
H2("3. Live Deployment URLs")
table(
    ["Component", "URL"],
    [
        ["Frontend (Surge)", "https://psb-securewealth-2026-new.surge.sh"],
        ["Frontend (Render mirror)", "https://psb-securewealth-frontend.onrender.com"],
        ["Backend API", "https://psb-securewealth-backend.onrender.com/api/v1"],
        ["Health Check", "https://psb-securewealth-backend.onrender.com/api/v1/health"],
        ["Admin Status", "https://psb-securewealth-backend.onrender.com/admin/status"],
    ],
    widths=[2.5, 4.5],
)
doc.add_paragraph()

# 4. Complete Feature Inventory
H1("4. Complete Feature Inventory")

H2("4.1 Wealth Intelligence")
table(
    ["Feature", "Description"],
    [
        ["Dashboard", "Net worth, financial pulse, market hero, next-best-actions, wealth DNA, notifications."],
        ["Wealth Twin", "Digital financial twin with Overview, Goals, Rebalance, Retirement, Tax, What-If tabs and conversational AI."],
        ["BHAVISHYA AI", "Predictive life-cycle engine forecasting life events and financial shocks."],
        ["Goals", "Goal tracker with AI coaching, boost cards, conflict intelligence, celebration animations."],
        ["Portfolio", "Holdings, allocation, rebalancing, ESG scoring."],
        ["Assets", "Bank + physical asset aggregation; manual asset form, AA account linking, physical-asset intelligence, vision appraisal."],
        ["Market", "Live indices, global macro radar, market news, strategist signals, smart triggers."],
        ["Forecast", "Scenario planning, Monte Carlo simulation, what-if forecasting."],
    ],
    widths=[1.8, 5.2],
)

H2("4.2 Payments, Transactions & Fraud Protection")
table(
    ["Feature", "Description"],
    [
        ["Payments", "UPI/NEFT/card/QR simulator, bill splitter, group jar, voice payment, MPIN, rewards, streaks, cashback wallet."],
        ["Transactions", "AI categorization, emotion check-in, receipt scanning, duplicate detection, smart tagging."],
        ["Protection", "Core fraud hub: risk meter, fraud engine, OTP simulation, duress mode, panic button, cooling vault, threat intel, behavioral biometrics, URL safety."],
        ["Security Beast", "10-layer security dashboard: TPM, eBPF, honeytokens, passkeys, post-quantum crypto (ML-KEM), behavioral biometrics, decentralized identity, transaction traps, secure enclave, blockchain audit."],
        ["Privacy", "Privacy audit panel, consent controls, data transparency."],
        ["Audit Log", "Immutable ledger of security events and protected actions."],
    ],
    widths=[1.8, 5.2],
)

H2("4.3 Financial Tools")
table(
    ["Feature", "Description"],
    [
        ["Tax Planner", "Old vs. new regime comparison, Section 80C tracker, tax calculator, deadline calendar, reports."],
        ["Calculators", "EMI, SIP, FD, retirement, rent-vs-buy calculators."],
        ["Bill Calendar", "Due dates, autopay reminders, recurring bill tracking."],
        ["Recurring Payments", "Standing instructions and e-mandates."],
        ["Account Statement", "Statement download and reconciliation."],
        ["Loan Center", "Loan application and tracking."],
    ],
    widths=[1.8, 5.2],
)

H2("4.4 Loans & Credit")
table(
    ["Feature", "Description"],
    [
        ["Credit Health", "CIBIL-style score, factors, and improvement tips."],
        ["CreditBridge AI", "Explainable, bias-audited retail & MSME credit scoring with reason codes, counterfactual simulator, ethics dashboard, and lender marketplace."],
        ["MSME CreditBridge", "Collateral-free MSME credit scoring and alternative-data underwriting."],
        ["Loans Hub", "Landing page for all loan and credit products."],
        ["Social Collateral Loan", "Community-backed trust-circle lending."],
        ["Loan Impact Simulator", "Estimates real-economy impact of MSME credit disbursement."],
        ["Loan Research Showcase", "Published research on algorithmic accountability in AI credit scoring."],
    ],
    widths=[1.8, 5.2],
)

H2("4.5 Family, Lifestyle & Gamification")
table(
    ["Feature", "Description"],
    [
        ["Family Dashboard", "Shared wallets, allowances, family goals."],
        ["Digital Gold", "Buy, sell, and gift 24K gold digitally."],
        ["Subscriptions", "Track and cancel recurring subscriptions."],
        ["Challenges", "Gamified savings and wealth challenges."],
        ["NRI Center", "NRE/NRO, remittance, and FEMA tools."],
        ["SME Centre", "Cash-flow timeline, surplus-fund advisor, working-capital health."],
        ["Kids Mode", "Safe money lessons for children."],
        ["Senior Mode", "Simplified, accessible UI for senior citizens."],
        ["Fantasy League", "Portfolio fantasy league and gamified portfolio."],
    ],
    widths=[1.8, 5.2],
)

H2("4.6 Innovation Lab, AI & Administration")
table(
    ["Feature", "Description"],
    [
        ["Innovation Lab", "Experimental hub: future self simulator, crisis predictor, macro-shock simulator, wealth weather, quantum lock, sovereign vault, etc."],
        ["AI Recommendations", "Explainable next-best-actions with reasoning tooltips and decision logs."],
        ["Values Alignment", "Invest according to personal values."],
        ["Pitch Deck", "Judge-facing pitch strategies and closing script."],
        ["Admin Panel", "Standalone /admin portal with Fraud Intelligence Center, case explorer, heatmap, trace graph, live simulation."],
        ["Onboarding Wizard", "First-time user setup."],
        ["Account Aggregator", "RBI AA consent flow, fetch animation, callback handler."],
        ["Demo Mode", "Public /demo showcase, judge tour, cinematic intro, demo personas."],
        ["Accessibility", "Accessibility settings, i18n, voice narration."],
    ],
    widths=[1.8, 5.2],
)

# 5. Flagship: CreditBridge AI
H1("5. Flagship Feature Deep-Dive: CreditBridge AI")
para("CreditBridge AI addresses the credit gap for thin-file retail borrowers and MSMEs using alternate data and explainable AI.")
H2("5.1 Phase 1 — Scoring Engine")
bullet("Dual mode: Retail/Individual and MSME.")
bullet("Transparent 300–900 score with exposed weights and clamps.")
bullet("No protected attributes (gender, caste, religion, location) are used.")
H2("5.2 Phase 2 — Explainability (XAI)")
bullet("SHAP-style factor breakdown with per-factor point impact.")
bullet("Rejection reason codes (CB-R01 to CB-R07) with actionable steps.")
bullet("Counterfactual ‘What-If’ simulator.")
bullet("CreditBridge vs Bureau-Only comparison and score waterfall chart.")
bullet("PDF explanation report export.")
H2("5.3 Phase 3 — Ethics & Accountability")
bullet("Model governance card with version, audit date, approved-by, data sources.")
bullet("Consent & data provenance ledger.")
bullet("Human review queue simulation with Approve/Reject/Escalate.")
bullet("Auto-generated adverse action notice with applicant rights.")
bullet("Accountability scorecard across explainability, fairness, oversight, regulatory alignment.")
H2("5.4 Phase 4 — Lender Marketplace")
bullet("Product catalog with purpose, rates, tenure, collateral, and sector/women-led boosts.")
bullet("Smart match scoring beyond a simple score cutoff.")
bullet("Filters by purpose, collateral, women-led.")
bullet("Offer detail modal with live EMI calculator and eligibility checklist.")
H2("5.5 Phase 5 — UX Polish")
bullet("Demo persona quick-fill buttons (Strong / Average / Risky).")
bullet("Input summary chips and real toast notifications.")
bullet("Accessibility improvements on toggle switches.")

# 6. Why This Tool Was Built (Flowchart)
H1("6. Why Was This Tool Built?")
para("The following flowchart maps the problem-to-solution journey:")
table(
    ["Step", "Node", "Detail"],
    [
        ["1", "Customer Pain Point", "Digital wealth users face fraud, impulsive actions, and a lack of personalized guidance."],
        ["2", "Bank Responsibility", "Banks must grow customer wealth AND safeguard it during critical money actions."],
        ["3", "Regulatory & Trust Need", "DPDP Act, RBI digital lending guidelines, and customer demand for explainable, consent-first AI."],
        ["4", "SecureWealth Twin", "AI-powered virtual financial twin + mandatory cyber-fraud protection layer."],
        ["5", "Outcome", "Wealth created responsibly, protected proportionally, and explained transparently."],
    ],
    widths=[0.6, 1.8, 5.6],
)

# 7. How It Was Built
H1("7. How Was It Built?")
H2("7.1 Tech Stack")
table(
    ["Layer", "Technology"],
    [
        ["Frontend", "React 19.2.5, TypeScript, Vite 8, Tailwind CSS v4, Framer Motion, Recharts, Lucide React, Zustand"],
        ["Backend", "Node.js, Express, SQLite (better-sqlite3), JWT, bcryptjs, Helmet, express-rate-limit, Winston, WebSocket"],
        ["AI", "Multi-provider orchestrator: Gemini, Groq, OpenAI, Anthropic, Cohere, Mistral, DeepSeek, Grok, NVIDIA, Hugging Face, OpenRouter"],
        ["Security", "ML-KEM-768, WebAuthn/Passkeys, FingerprintJS, behavioral biometrics, duress PIN, transaction trap"],
        ["Data Ingestion", "Yahoo Finance, RBI/World Bank, Alpha Vantage, Kaggle metadata"],
        ["Deployment", "Render (backend + static), Surge (frontend), Docker, GitHub Actions"],
    ],
    widths=[1.8, 5.2],
)

H2("7.2 Python Code Used")
para("Three Python scripts at the project root generate supporting Word documents for the demo team. They are not part of the runtime app but are part of the solution/documentation workflow.")

H3("build_cuecard.py")
para("Generates a one-page printable 10-minute demo cue card with click order and talking points.")
add_code_block("build_cuecard.py", lines=40)

H3("build_playbook.py")
para("Generates the full team demo playbook: pitch, architecture, feature inventory, protection layer deep-dive, run-of-show, compliance, Q&A bank, roles, scoring cheat-sheet.")
add_code_block("build_playbook.py", lines=40)

H3("build_tooltable.py")
para("Generates a landscape tool table in Hinglish: which component to show, when, why, how it is built, and which judging metric it scores.")
add_code_block("build_tooltable.py", lines=45)

# 8. Benefits
H1("8. Benefits & Impact")
table(
    ["Stakeholder", "Benefit"],
    [
        ["Retail Customers", "Personalized saving/investing guidance, goal tracking, fraud protection, simplified UI for seniors/kids."],
        ["MSMEs / Businesses", "Cash-flow analysis, surplus fund management, working-capital health, collateral-free credit scoring."],
        ["Students / First-Time Investors", "Guided financial journey, simple nudges, scenario simulation, no jargon."],
        ["Families", "Shared wallets, family goals, education/retirement planning, transparent AI reasoning."],
        ["The Bank", "Better engagement, cross-selling, retention, safer digital wealth adoption, data-driven product offerings."],
        ["Regulators", "Consent-first data handling, explainable AI, audit logs, KYC gating, no misleading claims."],
    ],
    widths=[1.8, 5.2],
)

# 9. Architecture Flowchart
H1("9. Architecture Flowchart")
para("The diagram below shows how users, frontend, backend, data sources, AI providers, and admin/fraud operations connect.")

# Generate architecture image
def generate_architecture_diagram(path):
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis('off')

    def box(x, y, w, h, text, color, fontsize=9):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03", facecolor=color, edgecolor='black', linewidth=1.2)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha='center', va='center', fontsize=fontsize, weight='bold', wrap=True)

    # User
    box(0.2, 6.5, 1.8, 0.8, 'Customer / Admin\n(Mobile / Web)', '#E3F2FD')
    # Frontend
    box(3, 6.5, 2.4, 0.8, 'React 19 + Vite\nSPA (Surge/Render)', '#FFF3E0')
    # Backend
    box(6.5, 6.5, 2.4, 0.8, 'Node.js + Express\n(API server)', '#E8F5E9')
    # Database
    box(10, 6.5, 1.6, 0.8, 'SQLite / PG\nDatabase', '#F3E5F5')

    # AI
    box(6.5, 4.8, 2.4, 0.8, 'AI Orchestrator\n(Gemini/OpenAI/…)', '#E0F7FA')
    # Market data
    box(10, 4.8, 1.6, 0.8, 'Real-World Data\nRBI / Yahoo / FX', '#FFEBEE')
    # Fraud
    box(3, 4.8, 2.4, 0.8, 'Protection Layer\nRisk Score + Duress', '#FFCDD2')
    # AA
    box(0.2, 4.8, 1.8, 0.8, 'RBI Account\nAggregator', '#E8EAF6')

    # Admin
    box(3, 3.1, 2.4, 0.8, 'Admin Portal\nFraud Intelligence', '#F1F8E9')
    # Audit
    box(6.5, 3.1, 2.4, 0.8, 'Audit Logger\nImmutable Logs', '#ECEFF1')

    # Arrows
    arrow_style = dict(arrowstyle='->', color='#37474F', lw=1.5)
    ax.annotate('', xy=(3, 6.9), xytext=(2, 6.9), arrowprops=arrow_style)
    ax.annotate('', xy=(6.5, 6.9), xytext=(5.4, 6.9), arrowprops=arrow_style)
    ax.annotate('', xy=(10, 6.9), xytext=(8.9, 6.9), arrowprops=arrow_style)

    ax.annotate('', xy=(7.7, 4.8), xytext=(7.7, 6.5), arrowprops=arrow_style)
    ax.annotate('', xy=(10, 5.2), xytext=(8.9, 5.2), arrowprops=arrow_style)
    ax.annotate('', xy=(3, 5.2), xytext=(2, 5.2), arrowprops=arrow_style)
    ax.annotate('', xy=(4.2, 5.6), xytext=(4.2, 6.5), arrowprops=arrow_style)

    ax.annotate('', xy=(4.2, 3.9), xytext=(4.2, 4.8), arrowprops=arrow_style)
    ax.annotate('', xy=(6.5, 3.5), xytext=(5.4, 3.5), arrowprops=arrow_style)

    # Legend
    legend_elements = [
        mpatches.Patch(facecolor='#E3F2FD', edgecolor='black', label='User'),
        mpatches.Patch(facecolor='#FFF3E0', edgecolor='black', label='Frontend'),
        mpatches.Patch(facecolor='#E8F5E9', edgecolor='black', label='Backend'),
        mpatches.Patch(facecolor='#FFCDD2', edgecolor='black', label='Protection'),
        mpatches.Patch(facecolor='#E0F7FA', edgecolor='black', label='AI / Data'),
    ]
    ax.legend(handles=legend_elements, loc='lower center', ncol=5, bbox_to_anchor=(0.5, -0.02), fontsize=9)
    ax.set_title('SecureWealth Twin — System Architecture', fontsize=14, weight='bold', pad=10)
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()


generate_architecture_diagram(DIAGRAM)
doc.add_picture(DIAGRAM, width=Inches(6.5))

H2("9.1 Architecture Explanation")
table(
    ["Flow", "Description"],
    [
        ["User → Frontend", "React SPA collects device fingerprint and JWT after login."],
        ["Frontend → Backend", "Authenticated API calls with Bearer JWT and X-Device-Id; CORS, Helmet, rate limits, JWT validation, per-user data scoping."],
        ["Account Aggregator", "Consent-based multi-bank data pull via RBI AA flow."],
        ["Protection Flow", "High-value actions checked by client-side fraud service and server timingCheck; summed into Wealth Protection Risk Score → Allow / Warn / Block."],
        ["Data & AI", "Real-world market/RBI data ingested into SQLite; AI prompts routed through multi-provider orchestrator with caching, quota, fallback."],
        ["Admin & Fraud Ops", "Admin portal authenticates with Bearer token; Fraud Intelligence Center polls fraud APIs and receives live WebSocket alerts."],
    ],
    widths=[1.8, 5.2],
)

# 10. Compliance & Responsible AI
H1("10. Compliance & Responsible AI")
bullet("Customer Data Privacy & Consent — clear permission before using financial data; explain what is collected and why.")
bullet("Secure Handling — encryption, secure APIs, no plain-text sensitive data; safe data flow even if simulated.")
bullet("Transparent & Explainable AI — no mysterious advice; show how and why recommendations are made.")
bullet("No Misleading Claims — only suggestions or simulated outcomes, no guaranteed returns or zero-risk claims.")
bullet("Basic Financial Regulations — KYC before real investment; demo badges shown for simulation.")

# 11. Demo Flow for Judges
H1("11. Suggested Demo Flow for Judges")
table(
    ["Time", "Block", "What to Show", "Key Message"],
    [
        ["0:00-0:45", "Hook + Architecture", "Live app + architecture slide", "We grow wealth AND protect it — dual mandate."],
        ["0:45-3:00", "Wealth Twin", "Dashboard → AA → Assets → Goals → AI Recommendations", "The twin builds a full picture and gives advice WITH a reason."],
        ["3:00-6:00", "Protection Layer", "High-value action → device/timing/amount/OTP → Risk Score → Warn/Block + Duress PIN", "Every critical action passes a proportional protection layer."],
        ["6:00-7:30", "Responsible AI", "Explainable tooltip → Consent → KYC → Audit log", "Consent-first, explainable, auditable."],
        ["7:30-9:00", "Innovation / CreditBridge", "Innovation Lab or CreditBridge AI scoring", "40+ modules; lightweight by default, deep on demand."],
        ["9:00-10:00", "Scale + Close", "Tech stack + deployment + team roles", "React + Node + Docker + AA + mock CBS; ready to scale."],
    ],
    widths=[0.9, 1.4, 3.2, 2.5],
)

# 12. Appendix
H1("12. Appendix: Evaluation Metrics Mapping")
para("Map the hackathon success metrics to features in this solution:")
table(
    ["Metric", "Features That Satisfy It"],
    [
        ["Quality of Wealth Intelligence", "Wealth Twin, AI Recommendations, Goals, Portfolio, Market, Forecast"],
        ["Effectiveness of Wealth Protection", "Protection hub, Risk Meter, OTP simulation, Duress mode, Cooling Vault, Security Beast"],
        ["Simplicity & UX", "Dashboard, Command Palette, Senior/Kids mode, Toast notifications, Accessibility settings"],
        ["Use of AI / Data Analytics", "Multi-provider AI orchestrator, spending persona, BHAVISHYA, CreditBridge scoring, real-world data ingestion"],
        ["Innovation", "Innovation Lab, CreditBridge AI, Social Collateral Loan, Behavioral biometrics, Quantum lock, Sovereign vault"],
        ["Technical Design & Architecture", "Feature-folder React, Node microservice structure, SQLite/PostgreSQL adapter, Docker, CI/CD"],
        ["Scalability & Practicality", "Docker container, Render/Surge deployment, multi-provider AI fallback, offline action queue"],
    ],
    widths=[2.2, 4.8],
)

# Save
doc.save(OUTPUT)
print(f"Saved {OUTPUT} ({os.path.getsize(OUTPUT)} bytes)")
