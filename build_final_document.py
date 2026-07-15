# -*- coding: utf-8 -*-
"""
SecureWealth Twin — FINAL DOCUMENT GENERATOR
Produces: SecureWealth_Twin_FINAL_DOCUMENT.docx
A 100+ page comprehensive Word document covering:
  - Problem statement alignment
  - Why / How / What of the solution
  - Every frontend feature with purpose and tech
  - Backend routes, algorithms, services
  - Key code highlights
  - Multiple architecture flowcharts
  - Demo flow, compliance, deployment, appendices
"""

import os
import textwrap
from datetime import datetime

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = "SecureWealth_Twin_FINAL_DOCUMENT.docx"

# Brand colours
NAVY = RGBColor(0x0B, 0x2B, 0x52)
BLUE = RGBColor(0x1E, 0x5A, 0x9C)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
RED = RGBColor(0xB3, 0x1B, 0x1B)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ORANGE = RGBColor(0xD4, 0x4C, 0x00)

# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def H1(text, color=NAVY):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.font.size = Pt(22)
    r.bold = True
    p.space_after = Pt(10)
    return p


def H2(text, color=BLUE):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.font.size = Pt(16)
    r.bold = True
    p.space_after = Pt(6)
    return p


def H3(text, color=NAVY):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.font.size = Pt(12)
    r.bold = True
    p.space_after = Pt(4)
    return p


def H4(text, color=GREY):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(11)
    return p


def para(text, size=10.5, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    p.space_after = Pt(5)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.add_run(text)
    p.space_after = Pt(3)
    return p


def mini_list(items):
    for it in items:
        bullet(it)


def numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.add_run(text)
    p.space_after = Pt(3)
    return p


def code_block(code, language=''):
    if language:
        H4(f"Code snippet — {language}")
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = 'Consolas'
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    return p


def page_break():
    doc.add_page_break()


def table(headers, rows, widths=None, header_fill='0B2B52'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], header_fill)
        pr = hdr[i].paragraphs[0]
        rn = pr.add_run(h)
        rn.bold = True
        rn.font.color.rgb = WHITE
        rn.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            pr = cells[i].paragraphs[0]
            rn = pr.add_run(str(val))
            rn.font.size = Pt(9)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    return t


def feature_block(title, problem, capabilities, tech, benefit, demo_tip, code_file=None, code_snippet=None):
    H3(title)
    para(f"Problem it solves: {problem}")
    para("Key capabilities:")
    for c in capabilities:
        bullet(c)
    para(f"How it is built: {tech}")
    if code_file:
        para(f"Primary code location: {code_file}", italic=True, color=GREY)
    if code_snippet:
        code_block(code_snippet)
    para(f"User benefit: {benefit}")
    p = doc.add_paragraph()
    r = p.add_run(f"Demo talking point: {demo_tip}")
    r.italic = True
    r.font.color.rgb = GREEN
    r.font.size = Pt(10)
    p.space_after = Pt(8)


# ---------------------------------------------------------------------------
# Architecture diagram generators
# ---------------------------------------------------------------------------
def draw_box(ax, x, y, w, h, text, color, fontsize=8):
    rect = mpatches.FancyBboxPatch(
        (x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.08",
        facecolor=color, edgecolor='#263238', linewidth=1.2
    )
    ax.add_patch(rect)
    ax.text(x + w / 2, y + h / 2, text, ha='center', va='center', fontsize=fontsize, weight='bold', wrap=True)


def draw_arrow(ax, x1, y1, x2, y2, color='#263238'):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=1.5,
                                connectionstyle='arc3,rad=0'))


def save_diagram(path, title, boxes, arrows, figsize=(11, 7)):
    fig, ax = plt.subplots(figsize=figsize)
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_title(title, fontsize=14, weight='bold', color='#0B2B52', pad=10)
    for b in boxes:
        draw_box(ax, *b)
    for a in arrows:
        draw_arrow(ax, *a)
    plt.tight_layout()
    plt.savefig(path, dpi=160, bbox_inches='tight')
    plt.close(fig)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------
def cover_page():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SecureWealth Twin")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Intelligent Wealth Growth with Built-in Fraud Protection")
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FINAL COMPREHENSIVE DOCUMENT")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = GREEN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PSBs Hackathon Series — 2026")
    r.font.size = Pt(14)
    r.font.color.rgb = GREY

    doc.add_paragraph()
    para("Generated for PPT preparation. This document contains the problem-solution fit, every feature, the code that powers it, architecture flowcharts, and a demo narrative.", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(f"Document generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, color=GREY)


# ---------------------------------------------------------------------------
# Section helper
# ---------------------------------------------------------------------------
def section_intro(number, title, body):
    H1(f"{number}. {title}")
    for b in body:
        para(b)


print("Building final document...")
cover_page()
page_break()

# ---------------------------------------------------------------------------
# Table of Contents
# ---------------------------------------------------------------------------
H1("Table of Contents")
numbered("Problem Statement & Hackathon Expectations")
numbered("Solution Overview — What is SecureWealth Twin?")
numbered("Why We Built It — Problem-Solution Fit")
numbered("How We Built It — Tech Stack & Methodology")
numbered("Benefits, Impact & Success Metrics")
numbered("Architecture & Flowcharts")
numbered("Feature Encyclopedia — Wealth Intelligence")
numbered("Feature Encyclopedia — Protection & Security")
numbered("Feature Encyclopedia — Financial Operations")
numbered("Feature Encyclopedia — Credit & Lending")
numbered("Feature Encyclopedia — Innovation Lab")
numbered("Feature Encyclopedia — Compliance & Trust")
numbered("Backend Deep Dive — Routes, Algorithms & Services")
numbered("Key Code Highlights")
numbered("Python Scripts & Document Automation")
numbered("Demo Flow & Pitch Narrative")
numbered("Deployment, Scalability & Roadmap")
numbered("Responsible AI, Privacy & Compliance")
numbered("Extended Technical Reference — APIs, Personas & Research")
numbered("Appendix A — Exhaustive Component Checklist")
numbered("Appendix B — Glossary & Abbreviations")
numbered("Appendix C — Project Metadata & Credits")
page_break()

# ---------------------------------------------------------------------------
# 1. Problem Statement
# ---------------------------------------------------------------------------
section_intro("1", "Problem Statement & Hackathon Expectations", [
    "The PSBs Hackathon Series-2026, under the Project Management Office (PMO) for Cyber Security and Fraud, issued a dual challenge: help customers grow wealth intelligently, and protect that wealth from fraud, misuse, social engineering, OTP misuse, coerced transactions, and impulsive high-value actions.",
    "The problem statement, captured in 'PSB Hackathon.pdf', defines SecureWealth Twin as an AI-powered digital wealth intelligence system that acts as a virtual financial twin of a customer or business, while embedding a mandatory, lightweight cyber-security and fraud protection layer around every critical wealth action.",
    "Judges expect a working prototype that demonstrates digital wealth intelligence, a mandatory wealth protection risk score, responsible AI practices, clear consent, explainability, auditability, and a scalable technical architecture."
])

H2("1.1 Core problem areas")
table(
    ["Problem Area", "Why it matters", "What judges look for"],
    [
        ["Fragmented financial view", "Customers hold money across banks, wallets, gold, property, and investments.", "Account Aggregator integration and unified net-worth dashboard."],
        ["One-size-fits-all advice", "Generic product push ignores risk appetite, goals, and cash-flow reality.", "AI-powered, explainable, goal-based recommendations."],
        ["Fraud & social engineering", "OTP misuse, coerced transactions, phishing, and mule accounts erode trust.", "Real-time risk scoring with proportional allow/warn/block actions."],
        ["Thin-file credit exclusion", "Salaried gig workers and MSMEs lack bureau history but have rich digital footprints.", "Alternative-data credit scoring with adverse-action transparency."],
        ["Compliance & trust deficit", "Users fear opaque AI and data misuse.", "Consent-first design, decision logs, privacy center, and KYC gating."],
    ],
    widths=[1.5, 2.2, 2.4]
)

H2("1.2 Success metrics from the problem statement")
bullet("Quality of wealth insights and recommendations")
bullet("Effectiveness of cyber/fraud protection")
bullet("Simplicity and user experience")
bullet("Use of AI / data analytics")
bullet("Innovation (behavioral nudges, scenario simulation, AI coach)")
bullet("Technical design, architecture, scalability, and practicality")
page_break()

# ---------------------------------------------------------------------------
# 2. Solution Overview
# ---------------------------------------------------------------------------
section_intro("2", "Solution Overview — What is SecureWealth Twin?", [
    "SecureWealth Twin is a full-stack, AI-augmented digital banking and wealth-protection platform prototype designed for Punjab & Sind Bank (PSB) customers. It combines a React 19 + Vite 8 + TypeScript frontend with a Node.js + Express + SQLite backend, wrapped in a security-first, consent-first, explainable architecture.",
    "The product has three inseparable layers: (1) Wealth Intelligence — dashboards, goal planning, portfolio, tax, market radar, and the AI Wealth Twin; (2) Wealth Protection — Rakshak real-time risk scoring, Security Beast, fraud intelligence center, duress PIN, behavioral biometrics, and family approval gates; (3) Financial Access — CreditBridge AI for retail and MSME lending, payments hub, insurance, gold, subscriptions, and business banking.",
    "Every critical action (payment, investment change, loan application, high-value transfer) passes through the protection layer. The system computes a Wealth Protection Risk Score and applies proportional friction: Allow, Warn with cooling-off, or Block and alert guardians."
])

H2("2.1 One-line value proposition")
para("A trusted financial twin that grows your wealth and guards it — explainably, securely, and in real time.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, color=NAVY)

H2("2.2 Solution pillars")
table(
    ["Pillar", "What it does", "Flagship feature"],
    [
        ["Wealth Intelligence", "Aggregates accounts, assets, income, and goals to generate personalized insights.", "Wealth Twin + AI Recommendations"],
        ["Wealth Protection", "Scores risk on every critical action and applies proportional guardrails.", "Rakshak / Protection Layer"],
        ["Financial Access", "Democratises credit, payments, insurance, and business tools.", "CreditBridge AI"],
        ["Responsible AI", "Consent, explainability, audit logs, and no protected attributes.", "Ethics Dashboard + Decision Logs"],
        ["Operations & Scale", "Admin fraud intelligence, compliance badges, and scalable deployment.", "Fraud Intelligence Center"],
    ],
    widths=[1.4, 3.2, 1.5]
)
page_break()

# ---------------------------------------------------------------------------
# 3. Why We Built It
# ---------------------------------------------------------------------------
section_intro("3", "Why We Built It — Problem-Solution Fit", [
    "The prototype was designed bottom-up from the problem statement. Every major module maps to one or more judge expectations. This section explains the 'why' behind the three flagship layers."
])

H2("3.1 Why Wealth Intelligence?")
para("Most PSB customers see their bank account as a static ledger. We built Wealth Intelligence because growing wealth requires context: income stability, goals, risk appetite, tax exposure, market conditions, and physical assets. Without a unified view, advice becomes a product pitch. The Wealth Twin turns fragmented data into a living financial model that can simulate the future, not just report the past.")

H2("3.2 Why Wealth Protection?")
para("Trust is the currency of digital banking. PSB's mandate explicitly calls for fraud protection around every critical wealth action. We built Rakshak because fraud is no longer just 'stolen passwords'; it includes social engineering, coerced transactions, OTP fatigue, and impulsive high-value decisions. A lightweight risk score with proportional action — warn, delay, block, duress PIN — protects customers without ruining UX.")

H2("3.3 Why CreditBridge AI?")
para("Traditional credit scoring excludes gig workers, rural entrepreneurs, and MSMEs with thin bureau files but rich digital footprints. We built CreditBridge AI because Account Aggregator data, UPI history, GST filings, and cash-flow stability are strong repayment signals. A deterministic, explainable score (300-900) with adverse-action notices aligns with RBI digital lending guidelines and the DPDP Act 2023.")

H2("3.4 Why Responsible AI & Compliance?")
para("Judges explicitly demand consent, explainability, and secure data handling. We built a Privacy Center, Consent Modal, KYC gating, AI Decision Log, and Ethics Dashboard so that every automated recommendation can be audited and challenged. No protected attribute (gender, caste, religion, location) is used as an input.")

H2("3.5 Problem-to-feature mapping")
table(
    ["Problem Statement Requirement", "SecureWealth Twin Feature", "Where to find it"],
    [
        ["Account Aggregator integration", "Account Aggregator Full + Widget", "features/aa, features/assets"],
        ["Track income, goals, risk appetite", "Wealth Twin, Goals, Tax, Portfolio", "features/ai/wealthTwin, features/goals"],
        ["Study market/economic trends", "Market View, Macro Radar, Smart Triggers", "features/market"],
        ["Add property/gold/vehicles", "Physical Asset Intelligence + Vision Appraisal", "features/assets"],
        ["Explainable recommendations", "Explainable Tooltip, AI Decision Log, ELI5", "features/ai"],
        ["Risk score before critical actions", "Rakshak, PaymentGuard, TransactionGuardModal", "features/protection"],
        ["Allow / Warn / Block actions", "CoolingVaultModal, DuressMode, FamilyApprovalGate", "features/protection"],
        ["Consent before data use", "ConsentModal, PrivacyCenter", "features/compliance"],
        ["KYC gating", "KYCModal, KYCStatusCard", "features/compliance, features/dashboard"],
        ["Innovation & behavioral nudges", "Behavioral Nudges, NBA Insights, Streaks, Challenges", "features/dashboard, features/gamification"],
        ["MSME & thin-file credit", "CreditBridge AI, MSME CreditBridge", "features/credit, features/msme"],
        ["Admin fraud operations", "Fraud Intelligence Center", "features/admin, backend/routes/fraud"],
    ],
    widths=[2.4, 2.4, 1.6]
)
page_break()

# ---------------------------------------------------------------------------
# 4. How We Built It
# ---------------------------------------------------------------------------
section_intro("4", "How We Built It — Tech Stack & Methodology", [
    "The prototype was built as a modern single-page application (SPA) with a lightweight backend. We chose technologies that are easy to demo, deploy on free tiers (Render), and scale to a production environment."
])

H2("4.1 Frontend stack")
bullet("React 19 with hooks and functional components")
bullet("Vite 8 for fast builds and HMR")
bullet("TypeScript for type safety across features")
bullet("Tailwind CSS v4 for utility-first styling and dark mode")
bullet("Framer Motion for animations")
bullet("Recharts for charts")
bullet("Zustand for global state")
bullet("shadcn/ui-inspired components in client/src/shared/components/ui/")

H2("4.2 Backend stack")
bullet("Node.js + Express for REST API")
bullet("SQLite via better-sqlite3 for zero-config persistence")
bullet("JWT-based authentication and role-based admin access")
bullet("WebSocket service for real-time notifications")
bullet("Redis-backed caching with in-memory fallback")
bullet("Nodemailer for email alerts")
bullet("Razorpay integration for payments (with fallback)")

H2("4.3 AI / ML approach")
para("We deliberately avoided black-box LLMs for regulated decisions like credit and fraud. Credit scoring uses a deterministic, additive feature attribution model. Fraud detection uses rule-based signal scoring enriched by behavioral biometrics. LLMs are used only for conversational assistance, natural-language scenario queries, and report summarisation, routed through an AI orchestrator with provider fallback.")

H2("4.4 Development methodology")
bullet("Feature-first folder structure: each domain owns components, services, hooks, and types.")
bullet("Shared services for cross-cutting concerns (auth, fraud, AI, notifications).")
bullet("Mock/demo data seeded in backend for consistent judge demos.")
bullet("Python scripts automate PPT source, cue card, playbook, and tool table generation.")
bullet("Frontend and backend auto-deploy from GitHub to Render.")

H2("4.5 Repository structure")
table(
    ["Directory", "Purpose"],
    [
        ["client/src/features/", "One folder per product domain (credit, protection, dashboard, etc.)."],
        ["client/src/shared/", "Reusable UI, services, hooks, config, and types."],
        ["backend/routes/", "REST API endpoints grouped by domain."],
        ["backend/algorithms/", "Deterministic scoring, tax, GST, fraud algorithms."],
        ["backend/services/", "Database, AI provider, market data, DCF, scenario engines."],
        ["supabase/", "Schema and SQL definitions for managed Postgres fallback."],
        ["build_*.py", "Python document generators for demo and PPT preparation."],
    ],
    widths=[1.4, 4.3]
)
page_break()

# ---------------------------------------------------------------------------
# 5. Benefits & Impact
# ---------------------------------------------------------------------------
section_intro("5", "Benefits, Impact & Success Metrics", [
    "This section quantifies the value SecureWealth Twin can deliver to PSB, its customers, and regulators."
])

H2("5.1 Customer benefits")
bullet("Unified net-worth view across banks, wallets, gold, property, and investments.")
bullet("Explainable AI recommendations that improve financial literacy, not just product sales.")
bullet("Proactive fraud protection that blocks social-engineering and coerced transactions.")
bullet("Access to credit for thin-file retail and MSME borrowers using digital footprint.")
bullet("Tax, GST, and business tools that reduce advisory costs for SMEs.")

H2("5.2 Bank benefits")
bullet("Lower fraud losses through real-time risk scoring and admin case management.")
bullet("Higher customer engagement via the Wealth Twin and gamification.")
bullet("Regulatory alignment with RBI digital lending, DPDP Act 2023, and GST compliance modules.")
bullet("Data-driven cross-sell of loans, insurance, and investments.")
bullet("Scalable microservices-ready architecture.")

H2("5.3 Expected success metrics")
table(
    ["Metric", "Target / Observation", "Source / Module"],
    [
        ["Credit decision transparency", "100% of scores show per-factor impact and reason codes.", "CreditBridge AI"],
        ["Fraud intervention coverage", "Every high-value / first-time / risky action is scored.", "Rakshak / Protection"],
        ["Thin-file inclusion", "Retail & MSME scoring without mandatory bureau score.", "CreditBridge AI"],
        ["Account aggregation", "Consent-based AA-linked account fetch.", "Account Aggregator"],
        ["Tax / GST automation", "Old-vs-new regime comparison and ITC risk scan.", "Tax + GST modules"],
        ["Decision auditability", "Every AI recommendation and protection action is logged.", "Decision Log + Audit"],
    ],
    widths=[1.8, 2.6, 1.5]
)
page_break()

# ---------------------------------------------------------------------------
# 6. Architecture & Flowcharts
# ---------------------------------------------------------------------------
section_intro("6", "Architecture & Flowcharts", [
    "SecureWealth Twin is built as a layered architecture. The diagrams below show the frontend module structure, backend service stack, the security/AI decision flow, data flow for a protected transaction, and the deployment topology."
])

ASSET_DIR = "final_doc_assets"
os.makedirs(ASSET_DIR, exist_ok=True)

# Diagram 1 — Frontend architecture
save_diagram(
    os.path.join(ASSET_DIR, "frontend_arch.png"),
    "Frontend Architecture — React + Vite + TypeScript",
    [
        (0.2, 5.8, 1.4, 0.7, "User / Browser", "#BBDEFB"),
        (2.2, 5.8, 1.4, 0.7, "React Router", "#C8E6C9"),
        (4.2, 5.8, 1.4, 0.7, "Zustand Store", "#C8E6C9"),
        (6.2, 5.8, 1.4, 0.7, "Shared UI / Hooks", "#C8E6C9"),
        (8.2, 5.8, 1.4, 0.7, "Axios / Fetch", "#C8E6C9"),
        (0.5, 4.3, 1.7, 0.7, "Dashboard", "#FFE0B2"),
        (2.5, 4.3, 1.7, 0.7, "Wealth Twin", "#FFE0B2"),
        (4.5, 4.3, 1.7, 0.7, "Protection", "#FFE0B2"),
        (6.5, 4.3, 1.7, 0.7, "CreditBridge", "#FFE0B2"),
        (8.5, 4.3, 1.7, 0.7, "Innovation Lab", "#FFE0B2"),
        (0.5, 2.8, 1.7, 0.7, "AA / Assets", "#FFE0B2"),
        (2.5, 2.8, 1.7, 0.7, "Payments", "#FFE0B2"),
        (4.5, 2.8, 1.7, 0.7, "Tax / GST", "#FFE0B2"),
        (6.5, 2.8, 1.7, 0.7, "Family / Goals", "#FFE0B2"),
        (8.5, 2.8, 1.7, 0.7, "Market / Business", "#FFE0B2"),
        (3.5, 1.3, 4.0, 0.7, "Backend API (REST + WebSocket)", "#FFCDD2"),
    ],
    [
        (0.9, 5.8, 2.2, 5.8),
        (2.9, 5.8, 4.2, 5.8),
        (4.9, 5.8, 6.2, 5.8),
        (6.9, 5.8, 8.2, 5.8),
        (8.9, 5.8, 9.6, 5.8), (9.6, 5.8, 9.6, 1.65), (9.6, 1.65, 7.5, 1.65),
        (1.3, 5.1, 1.3, 4.3),
        (3.3, 5.1, 3.3, 4.3),
        (5.3, 5.1, 5.3, 4.3),
        (7.3, 5.1, 7.3, 4.3),
        (9.3, 5.1, 9.3, 4.3),
        (1.3, 3.5, 1.3, 2.8),
        (3.3, 3.5, 3.3, 2.8),
        (5.3, 3.5, 5.3, 2.8),
        (7.3, 3.5, 7.3, 2.8),
        (9.3, 3.5, 9.3, 2.8),
        (5.5, 2.1, 5.5, 1.65),
    ],
)

H2("6.1 Frontend architecture")
para("The frontend is organised by feature domains. Each domain owns its components, services, hooks, and types. Shared UI primitives, services, and configuration live in client/src/shared/ so that any feature can reuse authentication, fraud detection, notifications, and AI orchestration.")
doc.add_picture(os.path.join(ASSET_DIR, "frontend_arch.png"), width=Inches(6.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Diagram 2 — Backend architecture
save_diagram(
    os.path.join(ASSET_DIR, "backend_arch.png"),
    "Backend Architecture — Node.js + Express + SQLite",
    [
        (0.2, 5.8, 1.4, 0.7, "Client", "#BBDEFB"),
        (2.2, 5.8, 1.4, 0.7, "Express Router", "#C8E6C9"),
        (4.2, 5.8, 1.4, 0.7, "JWT Auth", "#C8E6C9"),
        (6.2, 5.8, 1.4, 0.7, "Rate Limiter", "#C8E6C9"),
        (8.2, 5.8, 1.4, 0.7, "Validator", "#C8E6C9"),
        (0.5, 4.3, 1.6, 0.6, "auth.js", "#FFE0B2"),
        (2.2, 4.3, 1.6, 0.6, "banking.js", "#FFE0B2"),
        (3.9, 4.3, 1.6, 0.6, "fraud.js", "#FFE0B2"),
        (5.6, 4.3, 1.6, 0.6, "msme.js", "#FFE0B2"),
        (7.3, 4.3, 1.6, 0.6, "gst.js", "#FFE0B2"),
        (9.0, 4.3, 1.5, 0.6, "ai.js", "#FFE0B2"),
        (1.5, 3.0, 2.0, 0.6, "Algorithms", "#E1BEE7"),
        (4.0, 3.0, 2.0, 0.6, "Services", "#E1BEE7"),
        (6.5, 3.0, 2.0, 0.6, "Database", "#E1BEE7"),
        (1.5, 1.5, 2.0, 0.6, "WebSocket", "#FFCDD2"),
        (4.0, 1.5, 2.0, 0.6, "Redis Cache", "#FFCDD2"),
        (6.5, 1.5, 2.0, 0.6, "Email / Razorpay", "#FFCDD2"),
    ],
    [
        (0.9, 5.8, 2.2, 5.8),
        (2.9, 5.8, 4.2, 5.8),
        (4.9, 5.8, 6.2, 5.8),
        (6.9, 5.8, 8.2, 5.8),
        (8.9, 5.8, 9.6, 5.8),
        (1.3, 5.1, 1.3, 4.3),
        (3.0, 5.1, 3.0, 4.3),
        (4.7, 5.1, 4.7, 4.3),
        (6.4, 5.1, 6.4, 4.3),
        (8.1, 5.1, 8.1, 4.3),
        (9.75, 5.1, 9.75, 4.3),
        (2.0, 3.7, 2.0, 3.3),
        (4.5, 3.7, 4.5, 3.3),
        (7.0, 3.7, 7.0, 3.3),
        (2.0, 2.4, 2.0, 1.8),
        (4.5, 2.4, 4.5, 1.8),
        (7.0, 2.4, 7.0, 1.8),
    ],
)

H2("6.2 Backend architecture")
para("The backend exposes domain-specific REST routes under /api/v1. Each route uses services and deterministic algorithms. WebSocket pushes real-time alerts; Redis caches market data and AI responses; SQLite is the default zero-config store.")
doc.add_picture(os.path.join(ASSET_DIR, "backend_arch.png"), width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Diagram 3 — Security / AI decision flow
save_diagram(
    os.path.join(ASSET_DIR, "security_flow.png"),
    "Security & AI Decision Flow — Rakshak Protection Layer",
    [
        (0.2, 5.8, 1.6, 0.7, "User Action", "#BBDEFB"),
        (2.4, 5.8, 1.6, 0.7, "Device Fingerprint", "#C8E6C9"),
        (4.6, 5.8, 1.6, 0.7, "Behavioral Biometrics", "#C8E6C9"),
        (6.8, 5.8, 1.6, 0.7, "Transaction History", "#C8E6C9"),
        (9.0, 5.8, 1.3, 0.7, "Signals", "#FFF9C4"),
        (3.8, 4.3, 3.4, 0.7, "Fraud Detection Engine", "#FFE0B2"),
        (3.8, 2.8, 1.0, 0.7, "Allow", "#C8E6C9"),
        (5.0, 2.8, 1.0, 0.7, "Warn", "#FFF59D"),
        (6.2, 2.8, 1.0, 0.7, "Block", "#FFCDD2"),
        (3.5, 1.3, 4.0, 0.7, "Audit Log + Guardian Alert", "#E1BEE7"),
    ],
    [
        (0.9, 5.8, 2.4, 5.8),
        (2.9, 5.8, 4.6, 5.8),
        (5.1, 5.8, 6.8, 5.8),
        (7.3, 5.8, 9.0, 5.8),
        (9.65, 5.8, 9.65, 4.65), (9.65, 4.65, 7.2, 4.65),
        (5.5, 4.3, 5.5, 3.5),
        (4.3, 2.8, 4.3, 2.1),
        (5.5, 2.8, 5.5, 2.1),
        (6.7, 2.8, 6.7, 2.1),
        (5.5, 2.1, 5.5, 1.65),
    ],
)

H2("6.3 Security / AI decision flow")
para("Every critical action is intercepted by the protection layer. The engine combines device trust, behavioral biometrics, transaction history, and user-declared signals into a Wealth Protection Risk Score. The score maps to Allow, Warn (with cooling-off or family approval), or Block (with guardian alert and audit log).")
doc.add_picture(os.path.join(ASSET_DIR, "security_flow.png"), width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Diagram 4 — Data flow
save_diagram(
    os.path.join(ASSET_DIR, "data_flow.png"),
    "Data Flow — AA Consent to Protected Action",
    [
        (0.2, 5.8, 1.6, 0.7, "User Consent", "#BBDEFB"),
        (2.4, 5.8, 1.6, 0.7, "AA Gateway", "#C8E6C9"),
        (4.6, 5.8, 1.6, 0.7, "Data Ingestion", "#C8E6C9"),
        (6.8, 5.8, 1.6, 0.7, "Wealth Twin", "#FFE0B2"),
        (9.0, 5.8, 1.3, 0.7, "Insights", "#FFF9C4"),
        (6.8, 4.3, 1.6, 0.7, "Protection", "#FFCDD2"),
        (6.8, 2.8, 1.6, 0.7, "Action", "#C8E6C9"),
        (6.8, 1.3, 1.6, 0.7, "Audit", "#E1BEE7"),
    ],
    [
        (0.9, 5.8, 2.4, 5.8),
        (3.1, 5.8, 4.6, 5.8),
        (5.3, 5.8, 6.8, 5.8),
        (7.6, 5.8, 9.0, 5.8),
        (7.6, 5.1, 7.6, 4.3),
        (7.6, 3.6, 7.6, 2.8),
        (7.6, 2.1, 7.6, 1.3),
    ],
)

H2("6.4 End-to-end data flow")
para("The journey starts with explicit user consent through the Account Aggregator gateway. Consented data is ingested, normalised, and fed into the Wealth Twin. AI insights and recommendations are generated, but every outbound action is gated by the protection layer and written to an immutable audit log.")
doc.add_picture(os.path.join(ASSET_DIR, "data_flow.png"), width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Diagram 5 — Deployment
save_diagram(
    os.path.join(ASSET_DIR, "deployment.png"),
    "Deployment Topology — Render + GitHub CI/CD",
    [
        (0.2, 5.8, 1.6, 0.7, "GitHub Repo", "#BBDEFB"),
        (2.4, 5.8, 1.6, 0.7, "GitHub Actions", "#C8E6C9"),
        (4.6, 5.8, 1.6, 0.7, "Render CDN", "#FFE0B2"),
        (6.8, 5.8, 1.6, 0.7, "Render Web", "#FFE0B2"),
        (9.0, 5.8, 1.3, 0.7, "SQLite", "#E1BEE7"),
        (4.6, 4.3, 1.6, 0.7, "Browser", "#BBDEFB"),
        (6.8, 4.3, 1.6, 0.7, "Node API", "#C8E6C9"),
        (9.0, 4.3, 1.3, 0.7, "Redis", "#E1BEE7"),
    ],
    [
        (0.9, 5.8, 2.4, 5.8),
        (3.1, 5.8, 4.6, 5.8),
        (5.3, 5.8, 6.8, 5.8),
        (7.5, 5.8, 9.0, 5.8),
        (5.4, 5.1, 5.4, 4.3),
        (7.6, 5.1, 7.6, 4.3),
        (9.65, 5.1, 9.65, 4.65), (9.65, 4.65, 9.65, 4.3),
    ],
)

H2("6.5 Deployment topology")
para("The frontend is a static Vite build served by Render's CDN. The backend runs as a Node service on Render. SQLite is used for the prototype; production can migrate to Postgres (supabase/schema.sql is included) or a managed database. Redis is optional with in-memory fallback.")
doc.add_picture(os.path.join(ASSET_DIR, "deployment.png"), width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
page_break()

# ---------------------------------------------------------------------------
# 7. Feature Encyclopedia — Wealth Intelligence
# ---------------------------------------------------------------------------
section_intro("7", "Feature Encyclopedia — Wealth Intelligence", [
    "Wealth Intelligence turns raw account data into actionable, goal-oriented insight. This section covers the dashboard, Wealth Twin, goals, portfolio, tax, market, forecast, assets, family, business, and special modes."
])

H2("7.1 Dashboard & Financial Pulse")

feature_block(
    "Dashboard View",
    "Users need a single screen that summarises their entire financial life rather than opening multiple apps.",
    [
        "Net-worth card, income vs spend, recent transactions, and quick actions.",
        "KYC status, device status, and predictive shield badge.",
        "Stock ticker, market intelligence hero, and monthly narrative.",
    ],
    "React component in features/dashboard/DashboardView.tsx with Zustand for global state and Recharts for mini charts.",
    "One-glance financial health and fastest path to any action.",
    "Start the demo here: 'This is your financial home.'",
    code_file="client/src/features/dashboard/DashboardView.tsx",
)

feature_block(
    "Financial Pulse",
    "A single score is easier to communicate than a table of metrics.",
    [
        "Derives an overall health score from cash-flow, savings, debt, and goal progress.",
        "Colour-coded pulse with trend indicator.",
    ],
    "Computed in client-side services from AA and local transaction data; rendered by FinancialPulse.tsx.",
    "Motivates users to improve financial habits through a simple score.",
    "Show the pulse moving up after a good financial decision.",
)

feature_block(
    "Monthly Narrative",
    "Numbers are cold; stories are memorable.",
    [
        "Auto-generates a plain-English summary of the month's income, spends, savings, and anomalies.",
        "Powered by deterministic templates and optional LLM enrichment.",
    ],
    "MonthlyNarrative.tsx consumes transaction aggregates from the banking store.",
    "Users understand context without reading statements.",
    "'You saved 18% more than last month — great progress.'",
)

feature_block(
    "Behavioral Nudges",
    "Humans are predictably irrational with money.",
    [
        "Contextual nudges to pause before impulsive spends, top up goals, or avoid duplicate subscriptions.",
        "Rules engine mapped to behavioural finance principles.",
    ],
    "BehavioralNudges.tsx driven by nbaService.ts.",
    "Improves savings rate and reduces regretful spends.",
    "'You usually spend ₹3k on food delivery on weekends — set a limit?'",
)

H2("7.2 Wealth Twin")
para("The Wealth Twin is the centrepiece of the solution — a living model of the user's financial future. It aggregates net worth, goals, portfolio, retirement, tax, and what-if scenarios into tabbed views.")

feature_block(
    "Wealth Twin Overview Tab",
    "Users cannot plan what they cannot see.",
    [
        "Net-worth snapshot across liquid, physical, and investment assets.",
        "Asset allocation pie chart and liability view.",
    ],
    "OverviewTab.tsx, WealthTwinContext.tsx, and useWealthTwinData.ts orchestrate data from AA, assets, and local store.",
    "A unified, always-current financial picture.",
    "'Here is your entire balance sheet in one place.'",
)

feature_block(
    "Goals Tab",
    "Goals fail when they are not quantified.",
    [
        "Add, track, and visualise goals (emergency fund, vacation, education, home).",
        "Dynamic compass showing progress and projected completion date.",
        "Conflict intelligence warns when goals compete for the same rupee.",
    ],
    "GoalsTab.tsx, GoalTracker.tsx, GoalConflictIntelligence.tsx, AddGoalModal.tsx.",
    "Goal achievement becomes measurable and motivating.",
    "'Your emergency fund will be ready by March 2027 if you save ₹5,000/month.'",
)

feature_block(
    "Rebalance Tab",
    "Portfolios drift away from target allocation over time.",
    [
        "Compares current vs target asset allocation.",
        "Suggests buy/sell actions to rebalance.",
    ],
    "RebalanceTab.tsx computes drift from portfolio holdings.",
    "Maintains risk profile and improves long-term returns.",
    "'Your equity allocation is 12% above target — consider moving ₹40k to debt.'",
)

feature_block(
    "Retirement Tab",
    "Retirement planning feels distant and complex.",
    [
        "Projects retirement corpus based on current savings, expected return, and expenses.",
        "Shows monthly SIP gap.",
    ],
    "RetirementTab.tsx uses deterministic compound-interest projections.",
    "Users start retirement planning early.",
    "'Increase your SIP by ₹2,000/month to retire at 55.'",
)

feature_block(
    "Tax Tab",
    "Tax optimisation is a major source of hidden savings.",
    [
        "Old vs new regime comparison, HRA calculator, Section 80C tracker, and deadline calendar.",
        "Backend tax optimizer recommends the better regime and missing deductions.",
    ],
    "TaxTab.tsx + backend/algorithms/taxOptimizer.js + backend/routes/tax.js.",
    "Users minimise tax outgo legally and simply.",
    "'Switch to the old regime to save ₹18,420 this year.'",
)

feature_block(
    "What-If Tab",
    "People want to explore consequences before committing.",
    [
        "Sliders for income change, expense shock, market correction, and job loss.",
        "Instant projection of net worth and goal impact.",
    ],
    "WhatIfTab.tsx and client/src/shared/services/scenarioEngine.ts.",
    "Builds confidence in big financial decisions.",
    "'What if you lose your job for 3 months? Your emergency fund covers it.'",
)

H2("7.3 AI Recommendations & Explainability")

feature_block(
    "AI Recommendations View",
    "Generic cross-sell harms trust.",
    [
        "Next-best-action cards for SIP top-up, debt reduction, tax saving, insurance gap, and more.",
        "Ranked by impact, urgency, and goal alignment.",
    ],
    "AIRecommendationsView.tsx + RecommendationCard.tsx + nbaService.ts.",
    "Users receive timely, relevant, and explainable advice.",
    "'These three actions can improve your financial health the most this month.'",
)

feature_block(
    "Explainable Tooltip",
    "Opaque AI creates regulatory and trust risk.",
    [
        "Every AI recommendation shows why it was made.",
        "Traceable factors: cash-flow, goals, risk profile, market data.",
    ],
    "ExplainableTooltip.tsx + ELI5Tooltip.tsx.",
    "Users trust recommendations because they understand them.",
    "Hover over the recommendation to see the reasoning.",
)

feature_block(
    "AI Decision Log",
    "Regulators and users require proof that decisions are auditable.",
    [
        "Timestamped log of every AI-driven recommendation, protection action, and credit decision.",
        "Exportable for compliance review.",
    ],
    "AIDecisionLog.tsx writes to blockchainService.ts and auditLogger.ts.",
    "Full accountability and dispute resolution.",
    "'Every AI decision is logged and can be reviewed.'",
)

feature_block(
    "Financial Twin Chat",
    "Some users prefer conversation over dashboards.",
    [
        "Natural-language chat about balances, goals, spending, and recommendations.",
        "Routes to deterministic answers first; LLM fallback for open-ended queries.",
    ],
    "FinancialTwinChat.tsx + WealthChat.tsx + aiOrchestrator.ts.",
    "Lowers the barrier to financial literacy.",
    "Ask 'How am I doing on savings?' and get an instant answer.",
)

H2("7.4 Goals, Assets & Family")

feature_block(
    "Goal Tracker & Celebration",
    "Progress needs reward.",
    [
        "Visual goal progress bars, milestones, and celebration animations.",
        "Boosts manager lets users one-time top-up goals.",
    ],
    "GoalTracker.tsx, GoalCelebration.tsx, BoostsManager.tsx.",
    "Positive reinforcement drives savings discipline.",
    "Trigger a confetti celebration when a goal is hit.",
)

feature_block(
    "Account Aggregator",
    "A bank cannot see accounts held elsewhere.",
    [
        "Consent-based AA account linking and fetch.",
        "Animated fetch state and fallback demo data when backend is unavailable.",
    ],
    "AccountAggregatorFull.tsx, AAFetchAnimation.tsx, AACallbackHandler.tsx.",
    "True net-worth visibility across all banks.",
    "'With your consent, we pull all accounts into one dashboard.'",
)

feature_block(
    "Physical Asset Intelligence",
    "Net worth is more than bank balances.",
    [
        "Manual entry and AI-vision appraisal for gold, property, vehicles.",
        "Adds physical assets to the Wealth Twin balance sheet.",
    ],
    "PhysicalAssetIntelligence.tsx, ManualAssetForm.tsx, VisionAppraisalModal.tsx.",
    "Complete, accurate net-worth picture.",
    "'Snap a photo of your gold jewellery to estimate value.'",
)

feature_block(
    "Family Dashboard",
    "Household finance is a team sport.",
    [
        "Shared family view with member-level visibility controls.",
        "Family approval gate for large transfers.",
    ],
    "FamilyDashboard.tsx + guardianService.ts.",
    "Families plan together and protect jointly.",
    "'Your spouse can see savings goals, not individual transactions.'",
)

H2("7.5 Market, Forecast & Business")

feature_block(
    "Market View & Macro Radar",
    "Retail investors lack macro context.",
    [
        "Live quotes, curated news, macro signal tower, and smart triggers.",
        "Global macro radar flags rate changes, inflation, and currency moves.",
    ],
    "MarketView.tsx, GlobalMacroRadar.tsx, MacroSignalTower.tsx, useMacroFeed.ts.",
    "Users invest with macro awareness.",
    "'Inflation is rising — consider shorter-duration debt.'",
)

feature_block(
    "Forecast & Scenario Simulator",
    "Planning requires forward visibility.",
    [
        "Project net worth, cash-flow, and goal completion under base/bull/bear scenarios.",
        "Monte Carlo simulation shows probability bands.",
    ],
    "ForecastView.tsx, ScenarioSimulator.tsx, MonteCarloSimulator.tsx, backend/services/scenario-engine.js.",
    "Users stress-test plans before committing money.",
    "'There is a 78% chance you will reach your retirement target.'",
)

feature_block(
    "Business / SME Mode",
    "MSMEs need working capital visibility, not just personal finance.",
    [
        "Cash-flow timeline, working capital health, surplus fund advisor.",
        "Integrates with GST and CreditBridge MSME scoring.",
    ],
    "BusinessMode.tsx, SMEDashboard.tsx, CashFlowTimeline.tsx, WorkingCapitalHealth.tsx, SurplusFundAdvisor.tsx.",
    "Small business owners manage liquidity and credit in one place.",
    "'Here is your 12-month cash-flow forecast and working-capital score.'",
)

feature_block(
    "NRI Mode",
    "NRIs have distinct tax, remittance, and investment needs.",
    [
        "NRI-specific dashboard cards for NRE/NRO, remittance, and FEMA-aware nudges.",
    ],
    "NRIMode.tsx.",
    "Relevant experience for the NRI customer segment.",
    "'Track your India investments and remittances here.'",
)

feature_block(
    "Kids & Senior Modes",
    "Financial literacy and accessibility vary by age.",
    [
        "Kids Mode: simplified interface with chores, savings jars, and quizzes.",
        "Senior Mode: large fonts, high contrast, simplified actions, voice support.",
    ],
    "KidsMode.tsx, SeniorMode.tsx, AccessibilitySettings.tsx.",
    "Inclusion across age groups and abilities.",
    "'This is a kid-friendly view to learn saving habits.'",
)
page_break()

# ---------------------------------------------------------------------------
# 8. Feature Encyclopedia — Protection & Security
# ---------------------------------------------------------------------------
section_intro("8", "Feature Encyclopedia — Protection & Security", [
    "This is the mandatory layer required by the hackathon. It includes Rakshak transaction protection, the Security Beast command centre, the Fraud Intelligence Center for admins, duress and coerced modes, behavioral biometrics, and device trust."
])

H2("8.1 Rakshak / Wealth Protection Layer")

feature_block(
    "Rakshak Intervention Chat",
    "High-risk moments need human-like friction, not just a popup.",
    [
        "Conversational intervention when a transaction looks suspicious.",
        "Asks context questions and offers pause, cancel, or guardian alert.",
    ],
    "RakshakInterventionChat.tsx + backend/routes/ai.js /rakshak-intervention endpoint.",
    "Reduces impulsive and coerced transactions.",
    "'Can you confirm why you are sending ₹2 lakh at 2 AM to a new payee?'",
)

feature_block(
    "Transaction Guard Modal",
    "Risk needs a clear, proportionate action.",
    [
        "Displays Wealth Protection Risk Score and reason codes.",
        "Buttons for Allow, Warn/Cool-off, Block, and Duress PIN.",
    ],
    "TransactionGuardModal.tsx + PaymentGuard.tsx.",
    "User knows exactly why an action is restricted.",
    "'Score 82/100 — high risk. Block or wait 24 hours.'",
)

feature_block(
    "Cooling Vault Modal",
    "Cooling-off periods prevent impulse.",
    [
        "Locks high-risk actions behind a configurable delay (e.g., 30 min / 24 hr).",
        "Explains the delay and offers safer alternatives.",
    ],
    "CoolingVaultModal.tsx.",
    "Protects against scams and emotional decisions.",
    "'This investment will be held in the cooling vault until tomorrow 10 AM.'",
)

feature_block(
    "Duress PIN & Coerced Mode",
    "Victims of coercion may be forced to authenticate.",
    [
        "Secret duress PIN that fakes success while silently locking the account and alerting guardians.",
        "Coerced Mode banner visible only to the user under threat.",
    ],
    "DuressPinSetup.tsx, DuressMode.tsx, CoercedModeBanner.tsx, duressService.ts.",
    "Protects users during robbery or domestic coercion.",
    "'Enter your duress PIN. It looks like success, but help is on the way.'",
)

feature_block(
    "Family Approval Gate",
    "Large household decisions benefit from shared consent.",
    [
        "Configurable threshold above which a family member must approve.",
        "Push notification to guardian and in-app approval UI.",
    ],
    "FamilyApprovalGate.tsx + guardianService.ts.",
    "Prevents unilateral large transfers.",
    "'Transfers above ₹50,000 require your spouse's approval.'",
)

feature_block(
    "Family Safe Word",
    "Users in distress may need a discreet signal.",
    [
        "A pre-agreed word that, when entered, triggers a silent alert.",
        "Works inside chat or transaction notes.",
    ],
    "FamilySafeWord.tsx.",
    "Discreet SOS during a supervised transaction.",
    "'Type the safe word and the app alerts your family without the fraudster knowing.'",
)

feature_block(
    "Fraud Detection Engine",
    "Real-time risk scoring from transaction patterns.",
    [
        "Analyzes recent transactions for high amounts, rushed debits, blocked retries, and first-time payees.",
        "Returns signals and a 0-100 risk score.",
    ],
    "client/src/shared/services/fraudDetectionService.ts.",
    "Protection layer gets objective, auditable signals.",
    "'The engine flagged 5 recent debits and a new payee — score 75.'",
    code_file="client/src/shared/services/fraudDetectionService.ts",
    code_snippet='''const HIGH_AMOUNT_THRESHOLD = 50000;
const VERY_HIGH_AMOUNT_THRESHOLD = 200000;

export function analyzeTransactions(transactions: TransactionLike[]): FraudAnalysis {
  const sorted = [...transactions].sort(
    (a, b) => new Date(b.date).getTime() - new Date(a.date).getTime()
  );
  const recent = sorted.slice(0, 30);
  let riskScore = 0;
  // High-value transaction
  const highValueTx = recent.find((t) => t.amount >= HIGH_AMOUNT_THRESHOLD);
  if (highValueTx) {
    riskScore += highValueTx.amount >= VERY_HIGH_THRESHOLD ? 35 : 20;
  }
  // Rushed action
  const debitCount = recent.filter((t) => t.type === 'debit').length;
  if (debitCount >= 5) riskScore += 15;
  return { signals, riskScore: Math.min(100, riskScore), reasons };
}''',
)

feature_block(
    "Behavioral Biometrics",
    "Passwords can be stolen; behaviour is harder to fake.",
    [
        "Captures keystroke dynamics, mouse speed, and dwell time.",
        "Detects deviation from baseline profile and raises anomaly level.",
    ],
    "client/src/shared/services/behavioralBiometricsService.ts.",
    "Adds passive, continuous authentication.",
    "'Your typing rhythm changed — please verify with biometrics.'",
    code_file="client/src/shared/services/behavioralBiometricsService.ts",
    code_snippet='''export function calculateDeviation(metrics, profile) {
  if (!profile || profile.sampleCount < 10) return 0;
  const interKeyDev = computeDeviation(metrics.interKeyIntervals, [profile.avgInterKeyMs]);
  const dwellDev = computeDeviation(metrics.dwellTimes, [profile.avgDwellMs]);
  const mouseDev = computeDeviation(metrics.mouseSpeeds, [profile.avgMouseSpeedPxPerSec]);
  return (interKeyDev * 0.45 + dwellDev * 0.35 + mouseDev * 0.2);
}''',
)

feature_block(
    "Device Fingerprint & Trust",
    "New or untrusted devices are a top fraud signal.",
    [
        "Generates a device fingerprint from browser and OS signals.",
        "Trust score decays over time and flags unknown devices.",
    ],
    "client/src/shared/services/fingerprintService.ts.",
    "Blocks account takeover from stolen credentials on new devices.",
    "'This login is from an unknown device — approve it?'",
)

feature_block(
    "Location Verifier",
    "Impossible travel and location spoofing are common attack vectors.",
    [
        "Checks transaction location against trusted locations.",
        "Flags VPN/proxy usage when possible.",
    ],
    "LocationVerifier.tsx.",
    "Adds a geography dimension to risk scoring.",
    "'Transaction initiated 1,200 km from your home location.'",
)

feature_block(
    "URL Safety Checker",
    "Phishing links are a leading cause of fraud.",
    [
        "Heuristic checks plus DNS-over-HTTPS and TLS probe.",
        "Warns before opening suspicious links from SMS/email.",
    ],
    "URLSafetyChecker.tsx + urlSafetyService.ts.",
    "Cuts phishing success rate.",
    "Paste a link; the app tells you if it is safe.",
)

feature_block(
    "Scam Caller ID",
    "Phone-based social engineering is rising.",
    [
        "Checks incoming call numbers against known fraud patterns.",
        "Displays a risk badge during a call.",
    ],
    "ScamCallerID.tsx.",
    "Warns users before they share OTPs or details.",
    "'This number matches reported scam patterns.'",
)

H2("8.2 Security Beast")

feature_block(
    "Security Beast Dashboard",
    "Security features scattered across apps are ignored.",
    [
        "Unified command centre showing security score, passkeys, device trust, decoy account, dead man's switch, and threat intel.",
    ],
    "SecurityBeastView.tsx + SecurityScoreDashboard.tsx.",
    "Users can see and act on their security posture in one place.",
    "'Your security score is 72/100. Enable passkeys to reach 90.'",
)

feature_block(
    "Passkey / WebAuthn Auth",
    "Passwords are weak and phishable.",
    [
        "Platform authenticator registration and login.",
        "Falls back to password + OTP.",
    ],
    "PasskeyAuth.tsx + passkeyService.ts + webauthnService.ts.",
    "Strong, phishing-resistant authentication.",
    "'Use your fingerprint instead of a password.'",
)

feature_block(
    "Decoy Account",
    "Coercion can force users to reveal balances.",
    [
        "A decoy view with limited, pre-set data that opens under duress.",
        "Protects real wealth details during robbery.",
    ],
    "DecoyAccountView.tsx.",
    "Safety in physical threat scenarios.",
    "'Enter decoy PIN to show a fake low-balance account.'",
)

feature_block(
    "Dead Man's Switch",
    "Digital assets need inheritance planning.",
    [
        "Configures trusted contacts who can request account access after inactivity.",
        "Time-locked release with optional legal document upload.",
    ],
    "DeadMansSwitch.tsx + DigitalInheritance.tsx.",
    "Peace of mind for estate planning.",
    "'If inactive for 90 days, your nominee can initiate access.'",
)

feature_block(
    "Post-Quantum Crypto & Secure Enclave",
    "Future-proofing against quantum and side-channel threats.",
    [
        "Showcases post-quantum key concepts and secure-enclave/TPM attestation.",
    ],
    "PostQuantumCrypto.tsx, SecureEnclaveCheck.tsx, TpmAttestation.tsx.",
    "Demonstrates forward-looking security architecture.",
    "'Your data is protected by quantum-resistant design patterns.'",
)

H2("8.3 Fraud Intelligence Center (Admin)")

feature_block(
    "Fraud Intelligence Center",
    "Bank fraud teams need a central war room.",
    [
        "Case explorer, investigator notes, status workflow, priority, and category tags.",
        "Live feed, heatmap, map view, correlation clusters, and mule-trace graph.",
        "Export to Excel/CSV/PDF.",
    ],
    "FraudIntelligenceCenter.tsx + FraudCaseExplorer.tsx + FraudTraceGraph.tsx + backend/routes/fraud.js.",
    "Faster fraud investigation and reporting.",
    "'Here is the full fraud case, its money-flow hops, and related cases.'",
    code_file="backend/routes/fraud.js",
    code_snippet='''const VALID_STATUSES = ['open','investigating','escalated','closed','false_positive'];
const VALID_PRIORITIES = ['low','medium','high','critical'];
const VALID_CATEGORIES = ['account_takeover','mule_transfer','card_fraud','phishing','insider','identity_theft','velocity'];

router.get('/cases', adminApiAuth, (req, res) => { /* filter/paginate */ });
router.get('/cases/:id/hops', adminApiAuth, (req, res) => { /* multi-hop trace */ });
router.get('/correlations', adminApiAuth, (req, res) => { /* IP/destination clusters */ });
router.get('/export/cases', adminApiAuth, async (req, res) => { /* Excel/CSV export */ });''',
)

feature_block(
    "Fraud Rules Panel",
    "Static rules need tuning without a release.",
    [
        "Admin UI to enable/disable fraud detection rules.",
        "Rule hit-rate and false-positive metrics.",
    ],
    "FraudRulesPanel.tsx.",
    "Operational agility for fraud teams.",
    "'Turn on the new UPI-velocity rule for the festive season.'",
)

feature_block(
    "Threat Intel & Money Mule Graph",
    "Fraud is a network problem.",
    [
        "Threat-intelligence feed and multi-hop beneficiary graph for mule tracing.",
    ],
    "ThreatIntel.tsx, MoneyMuleGraph.tsx, FraudCorrelationPanel.tsx.",
    "Visualises fraud rings and shared infrastructure.",
    "'These three accounts share a common mule destination.'",
)
page_break()

# ---------------------------------------------------------------------------
# 9. Feature Encyclopedia — Financial Operations
# ---------------------------------------------------------------------------
section_intro("9", "Feature Encyclopedia — Financial Operations", [
    "Payments, subscriptions, insurance, gold, tax, transactions, and gamification form the operating layer of SecureWealth Twin."
])

H2("9.1 Payments Hub")

feature_block(
    "UPI Payment Simulator",
    "UPI is India's dominant payment rail.",
    [
        "Simulated UPI collect/request/pay flow with QR scanner.",
        "Animated transaction toast and receipt.",
    ],
    "UPIPaymentSimulator.tsx, QrScannerSimulator.tsx, AnimatedTransactionToast.tsx.",
    "Safe sandbox to demonstrate payments without real money.",
    "'Pay ₹500 to the chai stall via UPI.'",
)

feature_block(
    "Voice Payment",
    "Accessibility and speed.",
    [
        "Voice-activated payment confirmation with spoken amount and payee.",
    ],
    "VoicePayment.tsx + voiceService.ts.",
    "Hands-free payments for drivers and seniors.",
    "'Pay Ramesh 200 rupees.'",
)

feature_block(
    "Bill Splitter & Group Jar",
    "Social finance is a natural engagement driver.",
    [
        "Split bills among friends with individual shares and reminders.",
        "Group savings jar for trips, events, or gifts.",
    ],
    "BillSplitter.tsx, GroupJar.tsx.",
    "Reduces payment awkwardness and builds group savings.",
    "'Split the dinner bill of ₹3,600 among four people.'",
)

feature_block(
    "Payment Streaks & Rewards",
    "Habits stick when they are rewarded.",
    [
        "Consecutive payment streaks with fire animation.",
        "Spin wheel, cashback piggy, and rewards dashboard.",
    ],
    "StreakTracker.tsx, StreakFire.tsx, SpinWheel.tsx, CashbackPiggy.tsx, RewardsDashboardCard.tsx.",
    "Increases daily engagement and digital payment adoption.",
    "'You have a 12-day UPI streak — spin to win cashback.'",
)

feature_block(
    "Cashback Wallet",
    "Users want visibility into earned rewards.",
    [
        "Tracks cashback earned across transactions.",
        "Auto-sweep suggestions for idle cashback.",
    ],
    "CashbackWallet.tsx + cashbackEngine.ts.",
    "Turns rewards into productive capital.",
    "'You have ₹450 cashback — invest it in your emergency goal?'",
)

feature_block(
    "Recurring Payments",
    "Subscription and bill fatigue is real.",
    [
        "Lists all recurring debits with upcoming dates and amount changes.",
        "Smart duplicate detection flags overlapping subscriptions.",
    ],
    "RecurringPayments.tsx + subscriptionPredictor.ts.",
    "Prevents missed payments and unused subscriptions.",
    "'You are paying for two music apps. Cancel one to save ₹1,200/year.'",
)

feature_block(
    "Subscription Tracker",
    "Indians are accumulating OTT, cloud, and fitness subscriptions.",
    [
        "Tracks monthly/annual subscriptions, renewal dates, and spend trend.",
        "CultFit tracker and admin agent for offers.",
    ],
    "SubscriptionTracker.tsx, CultFitTracker.tsx, AdminAgent.tsx.",
    "Visibility into silent wealth drains.",
    "'Your subscriptions total ₹3,400/month — 12% of discretionary spend.'",
)

H2("9.2 Tax, GST & Business Compliance")

feature_block(
    "Tax Suite",
    "Tax planning is intimidating and error-prone.",
    [
        "Income tax calculator, old vs new regime comparison, Section 80C tracker, HRA, and deadline calendar.",
    ],
    "TaxView.tsx + backend/routes/tax.js + backend/algorithms/taxOptimizer.js.",
    "Maximises legal tax savings.",
    "'You can save ₹24,000 by investing ₹50,000 more in ELSS.'",
    code_file="backend/algorithms/taxOptimizer.js",
    code_snippet='''function optimizeTaxes(profile) {
  const results = { profile, oldRegime: {}, newRegime: {}, optimal: {}, recommendations: [] };
  results.oldRegime = calculateOldRegimeTax(profile);
  results.newRegime = calculateNewRegimeTax(profile);
  results.optimal = results.oldRegime.totalTax < results.newRegime.totalTax
    ? { ...results.oldRegime, regime: 'Old Regime' }
    : { ...results.newRegime, regime: 'New Regime' };
  results.recommendations = generateRecommendations(profile, results.optimal);
  return results;
}''',
)

feature_block(
    "GST Intelligence Suite",
    "MSMEs struggle with GST compliance and fake invoices.",
    [
        "GSTIN validation, ITC risk scanner, shell-company detector, GST rate verifier, and missing-ITC recovery predictor.",
    ],
    "backend/routes/gst.js + backend/algorithms/gstinValidator.js, itcRiskScanner.js, shellCompanyDetector.js, taxRateErrorDetector.js, missingITCRecovery.js.",
    "Protects MSMEs from GST penalties and fake suppliers.",
    "'Three of your suppliers show shell-company indicators.'",
    code_file="backend/algorithms/shellCompanyDetector.js",
    code_snippet='''riskIndicators: {
  NEW_REGISTRATION: { weight: 25, label: 'New Registration (<6 months)' },
  NO_RETURN_FILING: { weight: 35, label: 'No Return Filing' },
  CIRCULAR_TRANSACTIONS: { weight: 40, label: 'Circular Transactions' },
  CANCELLED_STATUS: { weight: 50, label: 'Cancelled GSTIN' },
  BLACKLISTED: { weight: 100, label: 'Blacklisted (Fake Invoice)' }
}''',
)

feature_block(
    "Account Statement & Audit Log",
    "Trust requires transparency.",
    [
        "Filterable account statement with AI categorisation.",
        "User-facing audit log of logins, consents, and protection actions.",
    ],
    "AccountStatement.tsx, AuditLog.tsx, AICategorization.tsx, SecurityLog.tsx.",
    "Dispute resolution and self-service transparency.",
    "'Here is every action taken on your account this month.'",
)

H2("9.3 Investments, Insurance & Gold")

feature_block(
    "Portfolio View",
    "Investments need risk and ESG context.",
    [
        "Holdings view with allocation, ESG score, and performance chart.",
    ],
    "PortfolioView.tsx, ESGScore.tsx.",
    "Holistic investment tracking.",
    "'Your portfolio ESG score is 72 — above benchmark.'",
)

feature_block(
    "Digital Gold",
    "Gold is a culturally preferred savings vehicle.",
    [
        "Buy/sell digital gold in small denominations.",
        "Links to goals and Wealth Twin.",
    ],
    "DigitalGold.tsx.",
    "Easy, secure gold savings.",
    "'Buy ₹500 of digital gold for your daughter's wedding goal.'",
)

feature_block(
    "Parametric Insurance",
    "Insurance claims are slow and opaque.",
    [
        "Explores parametric products where payout triggers are objective (rainfall, flight delay, hospitalisation).",
    ],
    "ParametricInsurance.tsx.",
    "Faster, transparent protection.",
    "'If rainfall in your district drops below 50%, payout is automatic.'",
)

feature_block(
    "Calculators",
    "Users want quick answers without advice.",
    [
        "Rent-vs-buy, EMI, SIP, and tax calculators.",
    ],
    "CalculatorsView.tsx, RentVsBuyCalculator.tsx.",
    "Self-service financial decision support.",
    "'Should you rent or buy that house? Compare in 30 seconds.'",
)

H2("9.4 Transactions & Gamification")

feature_block(
    "Transactions View",
    "Statements are dense and hard to read.",
    [
        "AI categorisation, smart duplicate detection, transaction tagger, emotion check-in, and receipt scan.",
    ],
    "TransactionsView.tsx, AICategorization.tsx, TransactionTagger.tsx, ScanReceipt.tsx, EmotionCheckin.tsx, SmartDuplicateDetection.tsx.",
    "Users understand spending patterns and reduce duplicates.",
    "'You spent ₹8,200 on dining this month — 20% above average.'",
)

feature_block(
    "Gamification",
    "Finance is boring unless it is a game.",
    [
        "Badges, streaks, challenges, fantasy portfolio league, and investment quiz.",
    ],
    "ChallengesView.tsx, BadgeStreak.tsx, FantasyLeague.tsx, FantasyPortfolio.tsx, InvestmentQuiz.tsx.",
    "Builds literacy and engagement through play.",
    "'Join the fantasy league to learn investing without real money.'",
)

feature_block(
    "Gig Income Smoother",
    "Gig workers have volatile income.",
    [
        "Detects irregular income and suggests a tax/EMI smoothing fund.",
    ],
    "GigIncomeSmoother.tsx.",
    "Stabilises cash-flow for non-salaried users.",
    "'Your income varies 35% monthly. Build a ₹60k smoothing buffer.'",
)
page_break()

# ---------------------------------------------------------------------------
# 10. Feature Encyclopedia — Credit & Lending
# ---------------------------------------------------------------------------
section_intro("10", "Feature Encyclopedia — Credit & Lending", [
    "Credit and lending are core to financial inclusion. SecureWealth Twin includes both a client-side CreditBridge AI engine and backend MSME scoring, plus a loans hub and social-collateral concept."
])

H2("10.1 CreditBridge AI (Retail + MSME)")

feature_block(
    "CreditBridge AI Scoring",
    "Thin-file borrowers are rejected by traditional bureau-only models.",
    [
        "Deterministic 300-900 credit score for retail and MSME.",
        "Retail factors: income stability, EMI burden, credit utilisation, payment history, banking vintage, UPI volume/consistency, savings rate, employment tenure, bureau score.",
        "MSME factors: vintage, turnover, GST filings/growth, digital payment share, bank balance, receivables cycle, sector, existing loans, Udyam, collateral, women-led.",
    ],
    "client/src/features/credit/services/creditBridgeEngine.ts + client/src/features/credit/components/CreditBridgeAI.tsx.",
    "Inclusive, auditable credit decisions.",
    "'Even without a CIBIL score, your cash-flow can qualify you.'",
    code_file="client/src/features/credit/services/creditBridgeEngine.ts",
    code_snippet='''export const RISK_BANDS: RiskBand[] = [
  { min: 300, max: 549, label: 'High Risk', color: '#dc2626', interest: '24% - 36%', approval: 'Rejected' },
  { min: 550, max: 649, label: 'Sub-Prime', color: '#ea580c', interest: '18% - 24%', approval: 'Conditional' },
  { min: 650, max: 749, label: 'Near-Prime', color: '#ca8a04', interest: '14% - 18%', approval: 'Approved' },
  { min: 750, max: 900, label: 'Prime', color: '#16a34a', interest: '10% - 14%', approval: 'Approved' },
];

export function calculateRetailScore(inputs: RetailInputs): ScoreResult {
  let score = 300;
  const incomeScore = Math.min(180, (inputs.monthlyIncome / 50000) * 30);
  score += incomeScore;
  const emiRatio = inputs.existingEmis / Math.max(inputs.monthlyIncome, 1);
  const emiImpact = clamp(60 - emiRatio * 300, -120, 60);
  score += emiImpact;
  // payment history, vintage, UPI, savings, tenure, bureau bonus ...
  score = clamp(Math.round(score), 300, 900);
  return { score, band: getBand(score), factors, maxLoan, mode: 'retail' };
}''',
)

feature_block(
    "Ethics Dashboard & Accountability",
    "Automated credit decisions must be explainable and challengeable.",
    [
        "Accountability scorecard across 6 dimensions.",
        "Consent ledger, human-review queue, and adverse-action notice generator.",
    ],
    "client/src/features/credit/services/ethicsEngine.ts.",
    "Builds regulatory trust and customer rights.",
    "'Here is exactly why you were declined and how to improve.'",
    code_file="client/src/features/credit/services/ethicsEngine.ts",
    code_snippet='''export const GOVERNANCE_INFO = {
  modelVersion: 'CB-AI-v2.1.0',
  approvedBy: 'PSB CreditBridge Ethics Committee',
  dataSources: [
    'Bank statement cash-flow (AA with consent)',
    'UPI transaction history (consent-based)',
    'GST filings (MSME, public + consent)',
    'Self-reported business / employment data',
    'Optional CIBIL bureau score',
  ],
  regulatoryFrameworks: [
    'RBI Master Direction on Digital Lending, 2022',
    'DPDP Act 2023 — consent & right to correction',
    'RBI Fair Practices Code',
    'EAA Framework (Saxena et al., 2025)',
  ],
};

export function generateAdverseActionNotice(result: ScoreResult): AdverseActionNotice {
  return {
    decision: 'We are unable to approve credit at this time.',
    score: result.score,
    band: result.band.label,
    reasonCodes: generateReasonCodes(result),
    applicantRights: [
      'Right to a clear explanation of the factors affecting your score.',
      'Right to request a human review of an automated decision.',
      'Right to correct inaccurate data used in the assessment.',
      'Right to know what data sources were used and withdraw consent.',
    ],
  };
}''',
)

feature_block(
    "Lender Marketplace",
    "Approved borrowers still shop for the best rate.",
    [
        "Product catalog from SBI, HDFC, Bajaj, Lendingkart, Kinara, Incred.",
        "Smart match scoring, eligibility filters, and EMI calculator.",
    ],
    "client/src/features/credit/services/lenderMarketplace.ts.",
    "Transparent comparison and informed borrowing.",
    "'You are eligible for 4 offers. Here is the cheapest EMI.'",
    code_file="client/src/features/credit/services/lenderMarketplace.ts",
    code_snippet='''export function calculateEMI(principal: number, annualRatePercent: number, tenureMonths: number): number {
  if (principal <= 0 || tenureMonths <= 0 || annualRatePercent <= 0) return 0;
  const r = annualRatePercent / 12 / 100;
  const emi = (principal * r * Math.pow(1 + r, tenureMonths)) / (Math.pow(1 + r, tenureMonths) - 1);
  return Math.round(emi);
}

export function matchProducts(result: ScoreResult, inputs): MatchedOffer[] {
  return LENDER_PRODUCTS.map((product) => {
    if (result.score < product.minScore) return { product, matchScore: 0, eligible: false };
    let matchScore = Math.min(30, (result.score - product.minScore) / 2);
    // turnover, collateral, sector, women-led boosts ...
    return { product, matchScore, eligible: true, estimatedRate, estimatedMaxAmount: eligibleCap, reason };
  }).sort((a, b) => b.matchScore - a.matchScore);
}''',
)

H2("10.2 MSME & Loans Hub")

feature_block(
    "MSME CreditBridge",
    "MSMEs need fast, collateral-friendly credit decisions.",
    [
        "Apply, preview score, view offers, accept, and track disbursement.",
        "Admin portfolio analytics and bias audit.",
    ],
    "features/msme/MSMEcreditbridgeView.tsx + backend/routes/msme.js + backend/algorithms/msmeCreditScore.js.",
    "Working capital for small businesses.",
    "'Your MSME score is 760. Three lenders matched.'",
    code_file="backend/algorithms/msmeCreditScore.js",
    code_snippet='''const FACTOR_WEIGHTS = {
  gstCompliance: 0.25,
  cashFlowStability: 0.30,
  transactionVolume: 0.20,
  digitalAdoption: 0.15,
  creditHistory: 0.10,
};

function assessApplication(input) {
  const factors = [
    scoreFactor('GST Compliance', FACTOR_WEIGHTS.gstCompliance, gstCompliance, 100, null, '✅'),
    scoreFactor('Cash Flow Stability', FACTOR_WEIGHTS.cashFlowStability, cashFlowStability, 100, null, '✅'),
    scoreFactor('Transaction Volume', FACTOR_WEIGHTS.transactionVolume, transactionVolume, 100, null, '⚠️'),
  ];
  const score = Math.min(1000, Math.max(0, factors.reduce((sum, f) => sum + parseInt(f.impact, 10), 0)));
  const fraudSignals = generateFraudSignals(input);
  const decision = determineDecision(determineCategory(score), fraudSignals);
  const offers = buildOffers(input, score, decision);
  return { score, category, decision, factors, fraudSignals, offers };
}''',
)

feature_block(
    "Loans Hub",
    "Loans are confusing with hidden costs.",
    [
        "Loan research showcase, impact simulator, and EMI comparison.",
        "Social-collateral loan concept for community-backed borrowing.",
    ],
    "LoansHub.tsx, LoanResearchShowcase.tsx, LoanImpactSimulator.tsx, SocialCollateralLoan.tsx.",
    "Borrowers understand total cost and impact.",
    "'This loan will delay your home goal by 4 months.'",
)

feature_block(
    "Loan Center & Recurring Payments",
    "Existing loans need management too.",
    [
        "Loan centre shows outstanding principal, EMI calendar, and prepayment impact.",
        "Recurring payments manage auto-debits.",
    ],
    "LoanCenter.tsx, RecurringPayments.tsx.",
    "Avoid missed EMIs and plan prepayments.",
    "'Prepay ₹50,000 to save ₹1.2 lakh interest.'",
)
page_break()

# ---------------------------------------------------------------------------
# 11. Feature Encyclopedia — Innovation Lab
# ---------------------------------------------------------------------------
section_intro("11", "Feature Encyclopedia — Innovation Lab", [
    "The Innovation Lab contains forward-looking modules that demonstrate future capabilities: predictive engines, emotional analytics, generational wealth, and crisis simulation."
])

H2("11.1 BHAVISHYA Engine")

feature_block(
    "BHAVISHYA Engine",
    "Future planning should be predictive, not reactive.",
    [
        "Predicts life events (marriage, education, medical) and macro shocks.",
        "Recommends proactive savings buffers.",
    ],
    "BhavishyaEngine.tsx, LifeEventPredictor.tsx, CrisisPredictor.tsx.",
    "Users prepare before crises happen.",
    "'You may face a medical expense in the next 18 months — build a ₹2L buffer.'",
)

feature_block(
    "Future Self Simulator",
    "People discount future consequences.",
    [
        "Shows an avatar of the user's future financial self based on current habits.",
    ],
    "FutureSelfSimulator.tsx.",
    "Motivates long-term behaviour change.",
    "'At this savings rate, your 60-year-old self will have ₹2.4 Cr.'",
)

feature_block(
    "Time Machine",
    "Past and future net-worth visualisation.",
    [
        "Slider to see net worth in past years and projected future.",
    ],
    "TimeMachine.tsx, TemporalWealth.tsx.",
    "Contextualises progress over a lifetime.",
    "'Five years ago your net worth was ₹4 lakh; in 10 years it could be ₹1.2 Cr.'",
)

H2("11.2 Emotional & Behavioural Analytics")

feature_block(
    "Emotional Heatmap & Resonance",
    "Spending is emotional.",
    [
        "Tags transactions with user mood and surfaces emotional spending patterns.",
        "Resonance view links spending to wellbeing.",
    ],
    "EmotionalHeatmap.tsx, EmotionCheckin.tsx, EmotionalResonance.tsx, MoodMeter.tsx.",
    "Raises self-awareness about impulse spending.",
    "'You tend to shop online when stressed.'",
)

feature_block(
    "Neuro Friction Widget",
    "Complex UI causes bad decisions.",
    [
        "Identifies screens and flows with high cognitive load.",
        "Suggests simplifications.",
    ],
    "NeuroFrictionWidget.tsx.",
    "Improves UX and reduces financial mistakes.",
    "'This investment screen has 14 fields — simplify to 6.'",
)

feature_block(
    "Festival-Aware Engine",
    "Indian saving and spending are seasonal.",
    [
        "Detects festivals and suggests savings plans, SIP pauses, or expense budgets.",
    ],
    "FestivalAwareEngine.tsx.",
    "Aligns advice with cultural cash-flow cycles.",
    "'Diwali is 3 months away — save ₹8,000/month for gifts.'",
)

H2("11.3 Generational & Community Wealth")

feature_block(
    "Generational Wealth & Digital Inheritance",
    "Wealth transfer is poorly digitised.",
    [
        "Plans inheritance with nominee mapping and time-locked document release.",
    ],
    "GenerationalWealth.tsx, DigitalInheritance.tsx, DeadMansSwitch.tsx.",
    "Smooth, secure inter-generational wealth transfer.",
    "'Your FDs, gold, and insurance are mapped to nominees.'",
)

feature_block(
    "Community DNA & Collective Immune System",
    "Fraud protection is stronger as a community.",
    [
        "Anonymised community fraud pattern sharing.",
        "Collective immune system flags new scam variants quickly.",
    ],
    "CommunityDNA.tsx, CollectiveImmuneSystem.tsx.",
    "Network-effect defence against scams.",
    "'Users like you reported this payee as suspicious.'",
)

H2("11.4 Simulation & Visualisation")

feature_block(
    "Monte Carlo & Macro Shock Simulators",
    "Point forecasts are misleading.",
    [
        "Monte Carlo shows probability distributions of future wealth.",
        "Macro shock simulator tests job loss, market crash, and inflation spikes.",
    ],
    "MonteCarloSimulator.tsx, MacroShockSimulator.tsx, LifeShockSimulator.tsx, StressTestSimulator.tsx.",
    "Users understand range of outcomes and build resilience.",
    "'In a 2008-style crash, your portfolio could fall 34%.'",
)

feature_block(
    "Neural Network Viz & Particle Constellation",
    "AI should be visual, not magical.",
    [
        "Animated visualisations of how AI models connect signals to decisions.",
    ],
    "NeuralNetworkViz.tsx, ParticleConstellation.tsx, FinancialDNAHelix.tsx.",
    "Demystifies AI for end users.",
    "'Each dot is a financial signal; the lines show how they influence your score.'",
)

feature_block(
    "Wealth Weather & Prosperity Score",
    "Abstract risk needs intuitive metaphors.",
    [
        "Weather metaphor for portfolio risk (sunny/stormy).",
        "Prosperity score combines wealth, security, and wellbeing.",
    ],
    "WealthWeather.tsx, ProsperityScore.tsx, PreparednessScore.tsx, ChakraBalance.tsx.",
    "Simple mental models for complex risk.",
    "'Your wealth weather is partly cloudy — reduce equity exposure.'",
)

feature_block(
    "Dream Visualizer",
    "Goals are abstract until visualised.",
    [
        "Turns savings goals into aspirational imagery and timelines.",
    ],
    "DreamVisualizer.tsx.",
    "Emotional connection to goals.",
    "'This is what your dream home fund looks like in 2029.'",
)
page_break()

# ---------------------------------------------------------------------------
# 12. Feature Encyclopedia — Compliance & Trust
# ---------------------------------------------------------------------------
section_intro("12", "Feature Encyclopedia — Compliance & Trust", [
    "Compliance is not a feature bolted on at the end; it is woven into the architecture. This section covers KYC, consent, privacy, explainability, and audit."
])

feature_block(
    "KYC Modal & Status",
    "Regulated actions require verified identity.",
    [
        "PAN/Aadhaar submission flow and status tracking.",
        "KYC gating for high-risk actions.",
    ],
    "KYCModal.tsx, KYCStatusCard.tsx, backend/routes/kyc.js.",
    "Meets RBI KYC requirements and protects minors.",
    "'Complete KYC to unlock investments above ₹50,000.'",
)

feature_block(
    "Consent Modal & Privacy Center",
    "DPDP Act 2023 makes consent central.",
    [
        "Granular consent capture before using financial data.",
        "Privacy centre shows what is collected, why, and how to withdraw.",
    ],
    "ConsentModal.tsx, PrivacyCenter.tsx, PrivacyView.tsx, PrivacyAuditPanel.tsx.",
    "Trust through transparency and user control.",
    "'You can withdraw AA consent at any time.'",
)

feature_block(
    "Compliance Badges & Bar",
    "Users need visible proof of trust.",
    [
        "Badges for encryption, consent, audit, and regulatory alignment.",
        "Top compliance bar summarises status.",
    ],
    "ComplianceBadges.tsx, ComplianceBar.tsx.",
    "Reassures users and judges.",
    "'This platform is built for RBI digital lending and DPDP compliance.'",
)

feature_block(
    "Blockchain Audit",
    "Audit logs can be tampered with if stored normally.",
    [
        "SHA-256 chained audit log for protection actions and AI decisions.",
    ],
    "client/src/shared/services/blockchainService.ts, BlockchainAudit.tsx.",
    "Immutable evidence for disputes and regulators.",
    "'Every decision is hashed and chained like a blockchain.'",
)

feature_block(
    "AI Decision Log & Explainability",
    "Automated decisions must be auditable.",
    [
        "Timestamped, factor-level log for credit, fraud, and recommendations.",
        "ELI5 tooltips make logs readable.",
    ],
    "AIDecisionLog.tsx, ExplainableTooltip.tsx, ELI5Tooltip.tsx.",
    "Right to explanation and human review.",
    "'Here is the exact logic that approved your loan.'",
)

feature_block(
    "Judge Tour & Demo Mode",
    "Hackathon demos must be smooth and self-contained.",
    [
        "Guided tour for judges, demo personas, cinematic intro, and notification demo.",
    ],
    "JudgeTour.tsx, DemoMode.tsx, DemoPersonas.tsx, CinematicIntro.tsx, NotificationDemo.tsx.",
    "Controlled demo environment with consistent data.",
    "'Switch to persona Ramesh to see the MSME journey.'",
)
page_break()

# ---------------------------------------------------------------------------
# 13. Backend Deep Dive
# ---------------------------------------------------------------------------
section_intro("13", "Backend Deep Dive — Routes, Algorithms & Services", [
    "The backend supports every frontend feature with deterministic algorithms, data persistence, and integrations. This section maps the backend surface area."
])

H2("13.1 REST API Routes")
table(
    ["Route file", "Responsibilities"],
    [
        ["admin.js", "Admin login, user management, stats, audit logs, IP geolocation, fraud actions."],
        ["ai.js", "AI orchestration: /ask, /summarize, /chat, /explain-cell, /rakshak-intervention, /execute-agent-action, /anomalies."],
        ["analytics.js", "Usage tracking, patent analytics, executive dashboard."],
        ["auth.js", "JWT registration/login/refresh, demo-login, face register/verify, profile."],
        ["banking.js", "Accounts, transactions, beneficiaries, cards, bills, subscriptions, goals, assets, loans, statements, dashboard."],
        ["business.js", "SME dashboard, cash-flow timeline, surplus advisor, working-capital health."],
        ["charts.js", "Football-field, waterfall, tornado, Monte Carlo, sensitivity matrix data."],
        ["documents.js", "GST invoice and generic report generator."],
        ["export.js", "Excel DCF workbook export with formulas and integrity checks."],
        ["extract.js", "PDF annual-report extraction via AI with quotas and SHA-256 cache."],
        ["financial-model.js", "5-year projections, DCF valuation, multiples, AI insights."],
        ["fraud.js", "Fraud cases, hops, accounts, notes, rules, stats, live feed, correlations, mule trace, export."],
        ["gallery.js", "Public model gallery with runtime DCF."],
        ["gst.js", "GSTIN validation, ITC risk, shell-company detection, rate verification, missing-ITC recovery."],
        ["kyc.js", "KYC status and PAN/Aadhaar submission."],
        ["market-data.js", "Quotes, historicals, comparables, benchmarks, macro signals, news."],
        ["msme.js", "MSME CreditBridge apply, score preview, offers, accept, admin portfolio, bias audit."],
        ["nlp-query.js", "Natural-language what-if query parser and executor."],
        ["realData.js", "RBI macro, instrument prices, forex, external datasets."],
        ["scenarios.js", "Base/upside/downside/stress/custom scenarios and LBO."],
        ["screener.js", "Screener.in search, company data, DCF inputs, peer comparison."],
        ["tax.js", "Multi-regime tax, HRA, optimisation."],
    ],
    widths=[1.2, 4.8]
)

H2("13.2 Deterministic Algorithms")
table(
    ["Algorithm", "Purpose"],
    [
        ["gstinValidator.js", "GSTIN format, checksum, state/entity risk decoding (PAT-001)."],
        ["itcRiskScanner.js", "Flags Input Tax Credit risks and mismatches (PAT-002)."],
        ["missingITCRecovery.js", "Predicts recoverable missing ITC and follow-up urgency (PAT-006)."],
        ["msmeCreditScore.js", "Deterministic MSME credit score using GST, cash-flow, digital adoption, fraud signals."],
        ["shellCompanyDetector.js", "Supplier risk profiling for shell-company indicators (PAT-003)."],
        ["taxOptimizer.js", "Old vs new regime comparison with recommendations (PAT-004)."],
        ["taxRateErrorDetector.js", "Verifies GST rates against HSN/SAC database (PAT-005)."],
    ],
    widths=[1.8, 4.0]
)

H2("13.3 Core Services")
table(
    ["Service", "Purpose"],
    [
        ["database.js", "SQLite schema and CRUD helpers for users, banking, fraud, MSME, market."],
        ["ai-provider.js", "Multi-provider AI dispatcher with Gemini default, BYOK fallback, retries, cost estimation."],
        ["anomaly-detector.js", "Deterministic model sanity checks (WACC bounds, balance-sheet balance)."],
        ["cacheService.js", "Redis-backed caching with in-memory fallback and TTL."],
        ["chart-engine.js", "Generates football-field, waterfall, tornado, Monte Carlo, sensitivity matrix."],
        ["dataIngestion.js", "Ingests RBI macro, market prices, forex, external datasets."],
        ["dcf-engine.js", "Runtime FCFF DCF with Gordon-growth terminal value."],
        ["demoData.js", "Pre-seeded demo user (demo-001) with accounts, assets, goals, transactions."],
        ["emailService.js", "Nodemailer/SMTP wrapper for welcome and tax reminders."],
        ["market-data.js", "Yahoo Finance / Alpha Vantage / mock fallback for quotes and historicals."],
        ["nlp-query.js", "Parses plain-English queries into scenario adjustments."],
        ["paymentService.js", "Razorpay order creation/verification with fallback mode."],
        ["scenario-engine.js", "Base/bull/bear/stress/custom scenarios and LBO."],
        ["screener-service.js", "Screener.in scraper with caching and DCF input builder."],
        ["websocket.js", "JWT-authenticated WebSocket for real-time notifications."],
    ],
    widths=[1.8, 4.0]
)

H2("13.4 Database Schema Highlights")
para("The default store is SQLite via better-sqlite3. Tables include users, accounts, transactions, beneficiaries, cards, bills, subscriptions, goals, assets, loans, fraud_cases, fraud_hops, msme_applications, market_quotes, and audit_logs. The supabase/schema.sql file provides a Postgres-ready migration path.")
page_break()

# ---------------------------------------------------------------------------
# 14. Key Code Highlights
# ---------------------------------------------------------------------------
section_intro("14", "Key Code Highlights", [
    "This section contains selected code snippets that demonstrate how critical features are implemented. Snippets are intentionally concise; full files are in the repository."
])

H2("14.1 Retail Credit Score Engine")
para("The engine is deterministic, additive, and fully auditable. Every factor exposes its impact and reason.")
code_block('''// client/src/features/credit/services/creditBridgeEngine.ts
export interface ScoreFactor {
  name: string;
  impact: number;
  weight: string;
  reason: string;
}

const clamp = (v: number, min: number, max: number) => Math.max(min, Math.min(max, v));

export function calculateRetailScore(inputs: RetailInputs): ScoreResult {
  const factors: ScoreFactor[] = [];
  let score = 300;

  const incomeScore = Math.min(180, (inputs.monthlyIncome / 50000) * 30);
  factors.push({
    name: 'Monthly Income',
    impact: Math.round(incomeScore),
    weight: '20%',
    reason: `₹${inputs.monthlyIncome.toLocaleString('en-IN')} monthly income`,
  });
  score += incomeScore;

  const emiRatio = inputs.existingEmis / Math.max(inputs.monthlyIncome, 1);
  const emiImpact = clamp(60 - emiRatio * 300, -120, 60);
  factors.push({
    name: 'EMI Burden',
    impact: Math.round(emiImpact),
    weight: 'High',
    reason: `EMI-to-income ratio ${(emiRatio * 100).toFixed(1)}%`,
  });
  score += emiImpact;

  // ... credit utilisation, payment history, vintage, UPI, savings, tenure, bureau bonus
  score = clamp(Math.round(score), 300, 900);
  return { score, band: getBand(score), factors, maxLoan: estimateMaxLoan(score, inputs), mode: 'retail' };
}''', "TypeScript")

H2("14.2 MSME Credit Score Engine")
para("Backend MSME scoring uses a weighted-factor approach with fraud signals and ELI5 explanation generation.")
code_block('''// backend/algorithms/msmeCreditScore.js
const FACTOR_WEIGHTS = {
  gstCompliance: 0.25,
  cashFlowStability: 0.30,
  transactionVolume: 0.20,
  digitalAdoption: 0.15,
  creditHistory: 0.10,
};

function generateFraudSignals(input) {
  const signals = [];
  if (input.gstin && input.gstin.length !== 15) {
    signals.push({ type: 'gst', message: 'GSTIN format appears invalid.', severity: 'high' });
  }
  if (input.annualTurnover > 0 && input.requestedAmount > input.annualTurnover * 0.5) {
    signals.push({ type: 'amount', message: 'Requested amount exceeds 50% of stated annual turnover.', severity: 'medium' });
  }
  if (input.employees && input.employees < 1 && input.annualTurnover > 5000000) {
    signals.push({ type: 'identity', message: 'High turnover with zero reported employees.', severity: 'medium' });
  }
  return signals;
}

function assessApplication(input) {
  const factors = [
    scoreFactor('GST Compliance', FACTOR_WEIGHTS.gstCompliance, gstCompliance, 100),
    scoreFactor('Cash Flow Stability', FACTOR_WEIGHTS.cashFlowStability, cashFlowStability, 100),
    scoreFactor('Transaction Volume', FACTOR_WEIGHTS.transactionVolume, transactionVolume, 100),
    scoreFactor('Digital Adoption', FACTOR_WEIGHTS.digitalAdoption, digitalAdoption, 100),
    scoreFactor('Credit History', FACTOR_WEIGHTS.creditHistory, creditHistory, 100),
  ];
  const score = Math.min(1000, Math.max(0, factors.reduce((sum, f) => sum + parseInt(f.impact, 10), 0)));
  const fraudSignals = generateFraudSignals(input);
  const decision = determineDecision(determineCategory(score), fraudSignals);
  const offers = buildOffers(input, score, decision);
  return { score, category, decision, factors, fraudSignals, offers, auditTrail: { engineVersion: '1.0.0', framework: 'TEAA', scoredAt: new Date().toISOString() } };
}''', "JavaScript")

H2("14.3 Fraud Detection Signals")
para("Fraud detection operates on transaction history rather than opaque models.")
code_block('''// client/src/shared/services/fraudDetectionService.ts
const HIGH_AMOUNT_THRESHOLD = 50000;
const VERY_HIGH_AMOUNT_THRESHOLD = 200000;

export function analyzeTransactions(transactions: TransactionLike[]): FraudAnalysis {
  const sorted = [...transactions].sort(
    (a, b) => new Date(b.date).getTime() - new Date(a.date).getTime()
  );
  const recent = sorted.slice(0, 30);
  const signals: RiskSignals = { newDevice: false, rushedAction: false, unusualAmount: false, otpRetries: false, firstTimeInvest: false, abnormalBehavior: false };
  const reasons: string[] = [];
  let riskScore = 0;

  const highValueTx = recent.find((t) => t.amount >= HIGH_AMOUNT_THRESHOLD);
  if (highValueTx) {
    signals.unusualAmount = true;
    reasons.push(`High-value transaction ₹${highValueTx.amount.toLocaleString()} detected`);
    riskScore += highValueTx.amount >= VERY_HIGH_AMOUNT_THRESHOLD ? 35 : 20;
  }

  const debitCount = recent.filter((t) => t.type === 'debit').length;
  if (debitCount >= 5) {
    signals.rushedAction = true;
    reasons.push(`${debitCount} recent debits indicate rushed activity`);
    riskScore += 15;
  }

  const blockedOrDelayed = recent.filter(
    (t) => t.status === 'BLOCKED' || t.status === 'DELAYED' || t.riskLevel === 'HIGH'
  );
  if (blockedOrDelayed.length > 0) {
    signals.otpRetries = true;
    reasons.push(`${blockedOrDelayed.length} recent blocked/delayed/high-risk transactions`);
    riskScore += 20;
  }

  riskScore = Math.min(100, riskScore);
  return { signals, riskScore, reasons };
}''', "TypeScript")

H2("14.4 Tax Optimisation")
para("Tax optimisation compares regimes and recommends deductions.")
code_block('''// backend/algorithms/taxOptimizer.js
const TAX_SLABS_NEW = [
  { limit: 300000, rate: 0 },
  { limit: 700000, rate: 0.05 },
  { limit: 1000000, rate: 0.10 },
  { limit: 1200000, rate: 0.15 },
  { limit: 1500000, rate: 0.20 },
  { limit: Infinity, rate: 0.30 }
];

function optimizeTaxes(profile) {
  const results = { profile, oldRegime: {}, newRegime: {}, optimal: {}, recommendations: [], savings: 0, confidence: 0 };
  results.oldRegime = calculateOldRegimeTax(profile);
  results.newRegime = calculateNewRegimeTax(profile);
  results.optimal = results.oldRegime.totalTax < results.newRegime.totalTax
    ? { ...results.oldRegime, regime: 'Old Regime' }
    : { ...results.newRegime, regime: 'New Regime' };
  results.recommendations = generateRecommendations(profile, results.optimal);
  results.savings = calculatePotentialSavings(profile, results.recommendations);
  results.confidence = calculateConfidenceScore(profile, results.recommendations);
  return results;
}''', "JavaScript")

H2("14.5 Shell Company Detection")
para("Supplier risk profiling protects MSMEs from fake invoices.")
code_block('''// backend/algorithms/shellCompanyDetector.js
riskIndicators: {
  NEW_REGISTRATION: { weight: 25, label: 'New Registration (<6 months)' },
  HIGH_VALUE_LOW_FREQ: { weight: 30, label: 'High Value, Low Frequency' },
  NO_RETURN_FILING: { weight: 35, label: 'No Return Filing' },
  CIRCULAR_TRANSACTIONS: { weight: 40, label: 'Circular Transactions' },
  ROUND_FIGURE_INVOICES: { weight: 10, label: 'Round Figure Invoices' },
  MONTH_END_CONCENTRATION: { weight: 15, label: 'Month-End Invoice Spike' },
  CANCELLED_STATUS: { weight: 50, label: 'Cancelled GSTIN' },
  BLACKLISTED: { weight: 100, label: 'Blacklisted (Fake Invoice)' }
},

checkRiskIndicators(supplier) {
  const indicators = [];
  const roundFigureCount = supplier.invoices.filter(inv => {
    const amount = parseFloat(inv.taxable_value || inv.amount || 0);
    return amount % 10000 === 0 && amount > 50000;
  }).length;
  if (roundFigureCount > supplier.invoices.length * 0.5) {
    indicators.push({ ...this.riskIndicators.ROUND_FIGURE_INVOICES, count: roundFigureCount });
  }
  const avgInvoiceValue = supplier.totalValue / supplier.invoices.length;
  if (avgInvoiceValue > 500000 && supplier.invoices.length < 5) {
    indicators.push({ ...this.riskIndicators.HIGH_VALUE_LOW_FREQ, avgValue: avgInvoiceValue });
  }
  return indicators;
}''', "JavaScript")

H2("14.6 Fraud Intelligence Center API")
para("Admin fraud operations expose full CRUD, correlation, and export endpoints.")
code_block('''// backend/routes/fraud.js
const VALID_STATUSES = ['open','investigating','escalated','closed','false_positive'];
const VALID_PRIORITIES = ['low','medium','high','critical'];
const VALID_CATEGORIES = ['account_takeover','mule_transfer','card_fraud','phishing','insider','identity_theft','velocity'];
const VALID_ACTIONS = ['acknowledge','investigate','escalate','close','false_positive'];

router.get('/cases', adminApiAuth, (req, res) => { /* filter/paginate */ });
router.post('/cases', adminApiAuth, validateBody(fraudSchemas.createCase), (req, res) => { /* create */ });
router.get('/cases/:id', adminApiAuth, (req, res) => { /* full case */ });
router.patch('/cases/:id', adminApiAuth, validateBody(fraudSchemas.updateCase), (req, res) => { /* update */ });
router.get('/cases/:id/hops', adminApiAuth, (req, res) => { /* multi-hop trace */ });
router.get('/correlations', adminApiAuth, (req, res) => { /* IP/destination/beneficiary clusters */ });
router.get('/export/cases', adminApiAuth, async (req, res) => { /* Excel/CSV export via exceljs */ });''', "JavaScript")
page_break()

# ---------------------------------------------------------------------------
# 15. Python Scripts & Document Automation
# ---------------------------------------------------------------------------
section_intro("15", "Python Scripts & Document Automation", [
    "Several Python scripts automate the creation of demo and presentation material. They live in the project root and use python-docx, matplotlib, and sometimes pandas/openpyxl."
])

H2("15.1 Why Python?")
para("Python was chosen for document automation because python-docx is the simplest way to programmatically generate rich Word documents with headings, tables, bullet lists, and embedded images. Matplotlib generates architecture diagrams. This very document is produced by a Python script so that any future iteration can be regenerated in seconds.")

H2("15.2 Script Inventory")
table(
    ["Script", "Output", "Purpose"],
    [
        ["build_final_document.py", "SecureWealth_Twin_FINAL_DOCUMENT.docx", "This 100+ page comprehensive final document."],
        ["build_comprehensive_doc.py", "SecureWealth_Twin_Comprehensive_PPT_Source.docx", "Slide-by-slide PPT source with diagrams and demo tips."],
        ["build_ppt_source_doc.py", "SecureWealth_Twin_PPT_Source.docx", "Standard PPT source document."],
        ["build_cuecard.py", "SecureWealth_Twin_Cue_Card.docx", "One-page printable 10-minute demo cue card."],
        ["build_playbook.py", "SecureWealth_Twin_Demo_Playbook.docx", "Detailed team demo playbook with talking points."],
        ["build_tooltable.py", "SecureWealth_Twin_Tool_Table.docx", "Hinglish 10-minute demo tool/priority table."],
    ],
    widths=[1.6, 2.4, 2.0]
)

H2("15.3 How build_final_document.py works")
para("The script follows a builder pattern:")
numbered("Imports python-docx, matplotlib, and standard library modules.")
numbered("Defines helper functions: H1/H2/H3 headings, paragraphs, bullets, tables, code blocks, and page breaks.")
numbered("Generates architecture diagrams with matplotlib and saves PNG files in final_doc_assets/.")
numbered("Iterates over feature dictionaries and calls feature_block() for every module.")
numbered("Appends backend inventories, code highlights, demo flow, compliance notes, and exhaustive appendices.")
numbered("Saves the final .docx file.")

code_block('''# build_final_document.py (core pattern)
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

OUTPUT = "SecureWealth_Twin_FINAL_DOCUMENT.docx"
doc = Document()

# Helper: styled heading
def H1(text, color=NAVY):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.font.size = Pt(22)
    r.bold = True

# Helper: feature block with problem, capabilities, tech, benefit, demo tip
def feature_block(title, problem, capabilities, tech, benefit, demo_tip):
    H3(title)
    para(f"Problem it solves: {problem}")
    para("Key capabilities:")
    for c in capabilities: bullet(c)
    para(f"How it is built: {tech}")
    para(f"User benefit: {benefit}")
    para(f"Demo talking point: {demo_tip}")

# Generate architecture diagram
def save_diagram(path, title, boxes, arrows):
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 11); ax.set_ylim(0, 7); ax.axis('off')
    ax.set_title(title, fontsize=14, weight='bold', color='#0B2B52')
    for x, y, w, h, text, color in boxes:
        rect = mpatches.FancyBboxPatch((x,y),w,h, boxstyle="round,pad=0.02", facecolor=color, edgecolor='black')
        ax.add_patch(rect)
        ax.text(x+w/2, y+h/2, text, ha='center', va='center', fontsize=8, weight='bold')
    for x1,y1,x2,y2 in arrows:
        ax.annotate('', xy=(x2,y2), xytext=(x1,y1), arrowprops=dict(arrowstyle='->', color='#263238', lw=1.5))
    plt.savefig(path, dpi=160, bbox_inches='tight')
    plt.close(fig)

# Add diagram to document
doc.add_picture('final_doc_assets/frontend_arch.png', width=Inches(6.5))

# Save
doc.save(OUTPUT)''', "Python")

H2("15.4 Re-running the document")
para("To regenerate this document after code changes, run: python build_final_document.py. The script scans no source files at runtime; it contains curated content. For fully auto-generated inventories, pair it with the explore agent or a static file scanner.")
page_break()

# ---------------------------------------------------------------------------
# 16. Demo Flow & Pitch Narrative
# ---------------------------------------------------------------------------
section_intro("16", "Demo Flow & Pitch Narrative", [
    "A 10-minute demo should cover the problem, the three pillars, and a memorable close. This section provides a suggested flow with timing and talking points."
])

H2("16.1 Suggested 10-minute demo flow")
table(
    ["Time", "Section", "What to show", "Talking point"],
    [
        ["0:00-0:30", "Hook", "Cinematic intro + problem statement", "Banks must grow wealth AND protect it."],
        ["0:30-2:00", "Wealth Intelligence", "Dashboard + Wealth Twin + Goals", "One view of your entire financial life."],
        ["2:00-3:30", "Protection", "Rakshak high-risk transaction", "Every critical action gets a risk score."],
        ["3:30-5:00", "Credit Inclusion", "CreditBridge AI retail or MSME", "Credit for thin-file users with explainable reasons."],
        ["5:00-6:00", "Innovation", "BHAVISHYA or Future Self", "Predictive, not reactive."],
        ["6:00-7:00", "Admin / Fraud", "Fraud Intelligence Center", "Bank-side war room for investigation."],
        ["7:00-8:00", "Compliance", "Privacy Center + Decision Log", "Consent-first, explainable, auditable."],
        ["8:00-9:00", "Tech / Scale", "Architecture diagram", "React + Node + SQLite/Postgres + Render."],
        ["9:00-10:00", "Close", "Impact metrics", "Grow wealth, guard trust, include everyone."],
    ],
    widths=[0.8, 1.3, 2.0, 2.4]
)

H2("16.2 Elevator pitch")
para("SecureWealth Twin is an AI-powered financial twin for Punjab & Sind Bank customers. It grows wealth through personalised, explainable recommendations and protects every critical action with a real-time fraud layer. It also expands financial inclusion via CreditBridge AI, which scores thin-file retail and MSME borrowers using Account Aggregator, UPI, and GST data — all without protected attributes. The result: trust, inclusion, and smarter money.")

H2("16.3 Demo personas")
table(
    ["Persona", "Journey to demo"],
    [
        ["Ramesh, Salaried", "Dashboard → Wealth Twin → Tax optimizer → Investment quiz."],
        ["Priya, Gig Worker", "Gig income smoother → CreditBridge AI retail → UPI payment."],
        ["Vikram, MSME Owner", "Business mode → GST intelligence → MSME CreditBridge → lender offers."],
        ["Senior Citizen", "Senior mode → Security Beast → Voice payment."],
        ["Fraud Analyst", "Admin login → Fraud Intelligence Center → case trace → export."],
    ],
    widths=[1.3, 4.5]
)
page_break()

# ---------------------------------------------------------------------------
# 17. Deployment, Scalability & Roadmap
# ---------------------------------------------------------------------------
section_intro("17", "Deployment, Scalability & Roadmap", [
    "The prototype is live and auto-deployed. This section explains how it is hosted today and how it can scale."
])

H2("17.1 Current deployment")
table(
    ["Component", "URL / Location", "Technology"],
    [
        ["Frontend", "https://psb-securewealth-frontend.onrender.com/", "React 19 + Vite 8 + Render static site"],
        ["Backend", "https://psb-securewealth-backend.onrender.com/", "Node.js + Express + Render web service"],
        ["Repository", "https://github.com/financeforumglrc/psb-securewealth-frontend.git", "GitHub"],
        ["Database", "SQLite (default)", "better-sqlite3"],
        ["Postgres schema", "supabase/schema.sql", "Supabase migration ready"],
    ],
    widths=[1.3, 3.2, 1.8]
)

H2("17.2 CI/CD pipeline")
para("GitHub Actions workflows in .github/workflows/ run lint/build checks on pull requests and auto-deploy the frontend to Render on pushes to main. The backend service restarts automatically on code changes.")

H2("17.3 Scalability levers")
bullet("Move from SQLite to Postgres or a managed cloud database for concurrent users.")
bullet("Containerise backend with Docker and orchestrate with Kubernetes or Render private services.")
bullet("Add Redis cluster for session, cache, and real-time WebSocket state.")
bullet("Scale AI provider routing with cost-aware load balancing and response caching.")
bullet("Integrate real Account Aggregator (Setu / Perfios) for production consents.")
bullet("Connect to CBS (Core Banking System) APIs for live balances and transactions.")

H2("17.4 Production roadmap")
table(
    ["Phase", "Timeline", "Deliverables"],
    [
        ["Pilot", "0-3 months", "50k PSB customers, AA integration, CreditBridge retail launch."],
        ["Scale", "3-9 months", "MSME CreditBridge, Fraud Intelligence Center, WebSocket alerts."],
        ["Enhance", "9-18 months", "Voice payments, advanced behavioural biometrics, open API marketplace."],
        ["Enterprise", "18+ months", "White-label for other PSBs, regulator dashboard, AI model registry."],
    ],
    widths=[1.0, 1.3, 3.3]
)
page_break()

# ---------------------------------------------------------------------------
# 18. Responsible AI, Privacy & Compliance
# ---------------------------------------------------------------------------
section_intro("18", "Responsible AI, Privacy & Compliance", [
    "SecureWealth Twin is designed to satisfy RBI, DPDP Act 2023, and the hackathon's responsible-AI expectations."
])

H2("18.1 No protected attributes")
para("CreditBridge AI and the AI recommendation engine never use gender, caste, religion, or precise location as inputs. This directly addresses the Saxena et al. (2025) finding that algorithmic credit systems can discriminate against women and marginalised groups.")

H2("18.2 Consent-first data use")
para("Account Aggregator consent is explicit, granular, time-bound, and revocable. The Privacy Center shows every data source and its purpose. Consent ledger entries are stored in the audit log.")

H2("18.3 Explainability & right to human review")
para("Every automated credit and protection decision exposes the factors, weights, and reason codes. Users can request human review from the Ethics Dashboard. Adverse-action notices list applicant rights.")

H2("18.4 Security & encryption")
bullet("JWT tokens with HttpOnly cookies in production.")
bullet("SHA-256 chained audit logs for tamper evidence.")
bullet("Passkey / WebAuthn support for phishing-resistant auth.")
bullet("Duress PIN and decoy account for physical-coercion scenarios.")
bullet("Device fingerprint and behavioural biometrics for continuous trust.")

H2("18.5 Regulatory alignment")
table(
    ["Regulation / Guideline", "How SecureWealth Twin aligns"],
    [
        ["RBI Master Direction on Digital Lending, 2022", "No automatic credit increase without consent; adverse-action notice; key fact statement."],
        ["DPDP Act 2023", "Explicit consent, data minimisation, right to correction, grievance mechanism."],
        ["RBI Fair Practices Code", "Transparent pricing, no misleading zero-risk claims, grievance redress."],
        ["EAA Framework (Saxena et al., 2025)", "Explainable AI, oversight, and right to challenge embedded in CreditBridge AI."],
    ],
    widths=[2.0, 3.6]
)
page_break()

# ---------------------------------------------------------------------------
# 19. Extended Technical Reference — APIs, Personas & Research
# ---------------------------------------------------------------------------
H1("19. Extended Technical Reference — APIs, Personas & Research")
para("This section provides additional technical depth for judges who want to see API contracts, user personas, research grounding, and risk thinking. It also includes extra code snippets that demonstrate the breadth of the implementation.")

H2("19.1 Key API Contract Summary")
para("The backend exposes a consistent REST API under /api/v1. Below are representative endpoints for each domain.")

H3("Authentication")
table(
    ["Endpoint", "Method", "Purpose"],
    [
        ["/api/v1/auth/register", "POST", "Create a new user account."],
        ["/api/v1/auth/login", "POST", "Login and receive JWT tokens."],
        ["/api/v1/auth/demo-login", "POST", "Login as a synthetic demo persona."],
        ["/api/v1/auth/refresh", "POST", "Refresh access token."],
        ["/api/v1/auth/me", "GET", "Get current user profile."],
        ["/api/v1/auth/face-register", "POST", "Register face descriptor for face login."],
        ["/api/v1/auth/face-verify", "POST", "Verify face descriptor for face login."],
    ],
    widths=[2.2, 0.8, 2.8]
)

H3("Banking & Wealth")
table(
    ["Endpoint", "Method", "Purpose"],
    [
        ["/api/v1/banking/dashboard", "GET", "Aggregated dashboard summary."],
        ["/api/v1/banking/accounts", "GET/POST", "List/create accounts."],
        ["/api/v1/banking/transactions", "GET/POST", "List/create transactions."],
        ["/api/v1/banking/goals", "GET/POST", "Goal CRUD."],
        ["/api/v1/banking/assets", "GET/POST", "Asset CRUD."],
        ["/api/v1/banking/loans", "GET/POST", "Loan CRUD and EMI schedule."],
        ["/api/v1/banking/recurring", "GET/POST", "Recurring payments."],
    ],
    widths=[2.4, 0.8, 2.6]
)

H3("Credit & MSME")
table(
    ["Endpoint", "Method", "Purpose"],
    [
        ["/api/v1/msme/apply", "POST", "Submit MSME loan application."],
        ["/api/v1/msme/score-preview", "POST", "Preview MSME score without application."],
        ["/api/v1/msme/offers", "GET", "Get matched lender offers."],
        ["/api/v1/msme/accept", "POST", "Accept an offer."],
        ["/api/v1/msme/admin/portfolio", "GET", "Admin portfolio analytics."],
        ["/api/v1/msme/admin/bias-audit", "GET", "Bias audit report."],
    ],
    widths=[2.6, 0.8, 2.4]
)

H3("Fraud & Admin")
table(
    ["Endpoint", "Method", "Purpose"],
    [
        ["/api/v1/fraud/cases", "GET/POST", "List/create fraud cases."],
        ["/api/v1/fraud/cases/:id/hops", "GET", "Multi-hop money trace."],
        ["/api/v1/fraud/correlations", "GET", "Cluster related cases."],
        ["/api/v1/fraud/stats/summary", "GET", "Dashboard statistics."],
        ["/api/v1/fraud/export/cases", "GET", "Excel/CSV export."],
        ["/api/v1/admin/stats", "GET", "Admin platform statistics."],
        ["/api/v1/admin/audit-logs", "GET", "Admin audit log with IP geolocation."],
    ],
    widths=[2.4, 0.8, 2.6]
)

H3("AI, Tax & GST")
table(
    ["Endpoint", "Method", "Purpose"],
    [
        ["/api/v1/ai/ask", "POST", "General AI Q&A."],
        ["/api/v1/ai/rakshak-intervention", "POST", "Generate intervention message."],
        ["/api/v1/ai/execute-agent-action", "POST", "Execute autonomous agent action."],
        ["/api/v1/tax/calculate", "POST", "Calculate and optimise tax."],
        ["/api/v1/gst/validate-gstin", "POST", "Validate GSTIN."],
        ["/api/v1/gst/itc-risk", "POST", "Scan ITC risk."],
        ["/api/v1/gst/shell-company", "POST", "Detect shell-company indicators."],
    ],
    widths=[2.4, 0.8, 2.6]
)
page_break()

H2("19.2 User Personas")

H3("Persona 1 — Ramesh, 32, Salaried IT Professional")
para("Background: Earns ₹90,000/month, has a home loan EMI, invests via SIPs, and wants to optimise taxes. Pain points: multiple apps, missed 80C deadlines, impulsive online spends.")
para("Features used: Dashboard, Wealth Twin, Tax Optimizer, Behavioral Nudges, Goal Tracker, UPI Payments.")
para("Outcome: Saves ₹24,000 tax, reduces impulsive spends by 15%, and reaches emergency fund goal 4 months early.")

H3("Persona 2 — Priya, 27, Gig Designer")
para("Background: Irregular income ₹40k-80k/month, no CIBIL score, needs credit for a laptop. Pain points: rejected by traditional lenders, income volatility.")
para("Features used: Gig Income Smoother, CreditBridge AI Retail, UPI Payment Streaks.")
para("Outcome: Qualifies for a ₹1.5 lakh personal loan based on UPI cash-flow and banking vintage, with explainable reason codes.")

H3("Persona 3 — Vikram, 45, MSME Owner")
para("Background: Runs a GST-registered manufacturing unit, turnover ₹1.2 Cr, needs working capital. Pain points: slow loan approvals, GST compliance anxiety, fake supplier invoices.")
para("Features used: Business Mode, GST Intelligence Suite, MSME CreditBridge, Lender Marketplace.")
para("Outcome: Gets MSME loan offer in minutes, flags two risky suppliers, and recovers ₹78,000 missing ITC.")

H3("Persona 4 — Sunita, 68, Retired Senior")
para("Background: Pension income, limited digital literacy, worried about fraud calls. Pain points: small fonts, complex menus, scam calls.")
para("Features used: Senior Mode, Security Beast, Scam Caller ID, Voice Payment, Family Approval Gate.")
para("Outcome: Uses voice commands safely; a suspicious call is flagged before any OTP is shared.")

H3("Persona 5 — Ankit, Fraud Analyst at PSB")
para("Background: Monitors fraud alerts, investigates mule networks, files regulatory reports. Pain points: scattered data, manual Excel reports.")
para("Features used: Fraud Intelligence Center, Fraud Trace Graph, Correlation Panel, Export Panel.")
para("Outcome: Identifies a fraud ring in minutes and exports a case bundle for law enforcement.")
page_break()

H2("19.3 Research & Regulatory Grounding")
para("The design is grounded in published research and Indian regulation.")
bullet("Saxena et al. (2025) found that AI credit scoring can achieve 85% repayment accuracy with 73% device-data usage, but also showed a 2.4x higher rejection rate for women and a 45-60% overall rejection rate. CreditBridge AI responds by excluding protected attributes and exposing every weight.")
bullet("The EAA (Explainability, Accountability, Auditability) framework informs the AI Decision Log, Ethics Dashboard, and adverse-action notices.")
bullet("RBI Master Direction on Digital Lending, 2022 mandates consent, key fact statements, grievance mechanisms, and no automatic credit limit increases. These are reflected in CreditBridge AI and the Consent Modal.")
bullet("The DPDP Act 2023 requires lawful purpose, data minimisation, and user rights. The Privacy Center and consent ledger operationalise these requirements.")
bullet("RBI Fair Practices Code requires transparent pricing and non-discriminatory practices. The lender marketplace shows estimated rates and eligibility clearly.")

H2("19.4 Competitive Differentiation")
table(
    ["Dimension", "Traditional Banking App", "Pure Fintech", "SecureWealth Twin"],
    [
        ["Wealth view", "Single bank only", "Aggregated but product-push", "AA + physical assets + goals + future simulation"],
        ["Fraud protection", "OTP + static limits", "Rule-based alerts", "Real-time risk score + duress PIN + biometrics + family gate"],
        ["Credit inclusion", "Bureau-only", "Alternative data, often opaque", "Deterministic explainable score, adverse-action notice"],
        ["Explainability", "Minimal", "Black-box models", "Factor-level impact, decision logs, ELI5"],
        ["Compliance UX", "Hidden", "Mixed", "Consent centre, privacy audit, compliance badges"],
    ],
    widths=[1.4, 1.7, 1.6, 1.6]
)
page_break()

H2("19.5 Risk Register & Mitigations")
table(
    ["Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["Model bias in credit scoring", "Medium", "High", "No protected attributes; factor transparency; bias audit endpoint."],
        ["False positives in fraud blocking", "Medium", "Medium", "Graduated allow/warn/block; human review queue; user override."],
        ["Data breach", "Low", "High", "JWT, passkeys, encryption at rest/transit, audit logs."],
        ["AA consent fatigue", "Medium", "Low", "Granular, time-bound consents; one-click renew/withdraw."],
        ["Backend unavailability", "Low", "Medium", "Offline queue, mock/demo data fallbacks, SQLite to Postgres migration path."],
    ],
    widths=[2.0, 0.9, 0.8, 2.2]
)

H2("19.6 Additional Code Snippets")

H3("AI Orchestrator with provider fallback")
code_block('''// client/src/shared/services/aiOrchestrator.ts
export async function orchestrate(request: AIRequest): Promise<AIResponse> {
  const providers = getEnabledProviders();
  const strategy = request.strategy || 'fastest-first';
  if (strategy === 'fallback') {
    for (const provider of providers) {
      try { return await callProvider(provider, request); }
      catch (e) { logFallback(provider, e); }
    }
  }
  if (strategy === 'ensemble') {
    const answers = await Promise.allSettled(providers.map(p => callProvider(p, request)));
    return aggregateAnswers(answers);
  }
  return raceProviders(providers, request);
}''', "TypeScript")

H3("Duress PIN logic")
code_block('''// client/src/shared/services/duressService.ts
export function evaluatePin(entered: string, realPin: string, duressPin: string): PinResult {
  if (entered === duressPin) {
    return { success: true, duress: true, action: 'LOCK_AND_ALERT' };
  }
  if (entered === realPin) {
    return { success: true, duress: false, action: 'ALLOW' };
  }
  return { success: false, duress: false, action: 'RETRY' };
}''', "TypeScript")

H3("Device fingerprint trust score")
code_block('''// client/src/shared/services/fingerprintService.ts
export function computeTrustScore(fp: DeviceFingerprint, history: FingerprintHistory): number {
  let score = 100;
  if (!history.known) score -= 30;
  if (fp.vpn || fp.tor) score -= 25;
  if (fp.os === 'uncommon') score -= 10;
  const daysSinceSeen = daysBetween(history.lastSeen, new Date());
  if (daysSinceSeen > 30) score -= 15;
  return Math.max(0, score);
}''', "TypeScript")

H3("Blockchain audit chain")
code_block('''// client/src/shared/services/blockchainService.ts
export function appendAuditRecord(record: AuditRecord): AuditBlock {
  const previous = chain[chain.length - 1];
  const block = {
    index: chain.length,
    timestamp: Date.now(),
    record,
    previousHash: previous.hash,
    hash: sha256(previous.hash + JSON.stringify(record) + Date.now()),
  };
  chain.push(block);
  return block;
}''', "TypeScript")

H2("19.7 Anticipated Judge Q&A")
para("The following questions and answers prepare the team for common hackathon judging queries.")

H3("Q1: How is this different from a normal banking app?")
para("A: Traditional banking apps show one bank's ledger and static limits. SecureWealth Twin aggregates AA data, physical assets, and goals; provides AI-driven explainable advice; and embeds a real-time protection layer with duress PIN, behavioural biometrics, and family gates. It also includes CreditBridge AI for thin-file borrowers.")

H3("Q2: Is the credit score really explainable?")
para("A: Yes. The score is deterministic. Every factor exposes its impact, weight, and reason. Users see a waterfall of contributions and receive adverse-action notices with reason codes and rights.")

H3("Q3: How do you prevent algorithmic bias?")
para("A: We do not use gender, caste, religion, or location as inputs. We provide a bias-audit endpoint for MSME decisions and a human-review queue for contested outcomes.")

H3("Q4: What about data privacy?")
para("A: Consent is captured before any AA fetch or AI use. The Privacy Center shows every data source and purpose. Users can withdraw consent and request deletion.")

H3("Q5: How does the fraud layer avoid false positives?")
para("A: The engine uses graduated actions — Allow, Warn with cooling-off, or Block — rather than hard declines. Each decision exposes the signals. Users and family guardians can override or confirm.")

H3("Q6: Is the backend production-ready?")
para("A: The prototype uses SQLite and Render for speed of demo. The architecture is stateless and ready for Postgres, Redis, and Kubernetes. A Supabase schema migration is included.")

H3("Q7: How does the AI orchestrator handle provider failure?")
para("A: It supports fastest-first, fallback, and ensemble strategies. If the primary provider fails, the request routes to the next configured provider.")

H3("Q8: What is the MSME CreditBridge advantage?")
para("A: MSMEs are scored using GST compliance, cash-flow stability, transaction volume, digital adoption, and credit history — not just a bureau score. Fraud signals and ELI5 explanations are included.")

H3("Q9: Can the bank operations team use this today?")
para("A: The Fraud Intelligence Center provides case management, multi-hop tracing, correlation clusters, rules management, and Excel export — sufficient for a pilot fraud operations team.")

H3("Q10: How scalable is the architecture?")
para("A: Frontend is a static build. Backend routes are stateless. Database can be migrated to Postgres. Redis can be clustered. AI calls are cached and cost-routed. The design supports horizontal scaling.")

H2("19.8 References")
numbered("Saxena, A., Sharma, R., Mohanty, S. N., & Sharma, A. (2025). Algorithmic Accountability in AI-Driven Credit Scoring. AI For Sustainable Progress.")
numbered("Reserve Bank of India. (2022). Master Direction on Digital Lending.")
numbered("Ministry of Electronics and Information Technology, Government of India. (2023). Digital Personal Data Protection Act.")
numbered("Reserve Bank of India. Fair Practices Code for Lenders.")
numbered("Account Aggregator ecosystem specifications (Sahamati / ReBIT).")
page_break()

# ---------------------------------------------------------------------------
# 20. Appendix A — Exhaustive Component Checklist
# ---------------------------------------------------------------------------
H1("Appendix A — Exhaustive Component Checklist")
para("This appendix lists every major component, service, route, and script discovered in the codebase. It is provided so the PPT author can verify that no module has been omitted.")

H2("A.1 Frontend Features")

def dir_list(title, files):
    H3(title)
    for f in files:
        bullet(f)

dir_list("Account Aggregator (features/aa)", [
    "AACallbackHandler.tsx — Handles AA OAuth/consent callback flow.",
    "AAFetchAnimation.tsx — Animated loading state for linked-account fetch.",
    "AccountAggregatorFull.tsx — Full consent flow and linked-account dashboard.",
])

dir_list("Accessibility (features/accessibility)", [
    "AccessibilitySettings.tsx — Font size, contrast, reduced-motion preferences.",
])

dir_list("Admin / Fraud Ops (features/admin)", [
    "AdminActivityTab.tsx", "AdminDashboard.tsx", "AdminLoginArchitecture.tsx", "AlertHistoryTab.tsx",
    "AlertToast.tsx", "DemoTour.tsx", "FraudCaseDetail.tsx", "FraudCaseExplorer.tsx",
    "FraudCorrelationPanel.tsx", "FraudExportPanel.tsx", "FraudHeatmap.tsx", "FraudIntelligenceCenter.tsx",
    "FraudMapView.tsx", "FraudRiskExplainer.tsx", "FraudRulesPanel.tsx", "FraudTimeline.tsx",
    "FraudTraceGraph.tsx", "SystemHealthTab.tsx", "useFraudCases.ts", "fraudDataGenerator.ts",
    "fraudTypes.ts", "permissions.ts", "fraudService.ts",
])

dir_list("AI / Wealth Twin (features/ai)", [
    "AIDecisionLog.tsx", "AIRecommendationsView.tsx", "AgenticActionCard.tsx", "BehavioralEngine.tsx",
    "ELI5Tooltip.tsx", "ExplainableTooltip.tsx", "FinancialLiteracyCards.tsx", "FinancialTwinChat.tsx",
    "RecommendationCard.tsx", "SmartActionOrchestrator.tsx", "WealthChat.tsx", "WealthTwinView.tsx",
    "wealthTwin/ExplainablePanel.tsx", "wealthTwin/GoalsTab.tsx", "wealthTwin/OverviewTab.tsx",
    "wealthTwin/RebalanceTab.tsx", "wealthTwin/RetirementTab.tsx", "wealthTwin/TaxTab.tsx",
    "wealthTwin/TwinTabs.tsx", "wealthTwin/WealthTwinContext.tsx", "wealthTwin/WhatIfTab.tsx",
    "wealthTwin/useWealthTwinData.ts", "wealthTwin/utils.ts",
])

dir_list("Architecture (features/architecture)", [
    "FeaturesUniverse.tsx", "PerformanceMetrics.tsx", "SystemArchitecture.tsx",
])

dir_list("Assets (features/assets)", [
    "AccountAggregatorWidget.tsx", "LinkAccountModal.tsx", "ManualAssetForm.tsx",
    "PhysicalAssetIntelligence.tsx", "VisionAppraisalModal.tsx",
])

dir_list("Auth (features/auth)", [
    "BiometricAuth.tsx", "CreateAccountModal.tsx", "CreateAccountModal.test.tsx",
    "FaceLoginModal.tsx", "LoginPage.tsx", "LoginPortal.tsx",
])

dir_list("Banking (features/banking)", [
    "AccountStatement.tsx", "AuditLog.tsx", "LoanCenter.tsx", "RecurringPayments.tsx",
])

dir_list("Bills (features/bills)", [
    "BillCalendar.tsx",
])

dir_list("Business / SME (features/business)", [
    "BusinessMode.tsx", "CashFlowTimeline.tsx", "SMEDashboard.tsx",
    "SurplusFundAdvisor.tsx", "WorkingCapitalHealth.tsx",
])

dir_list("Calculators (features/calculators)", [
    "CalculatorsView.tsx", "RentVsBuyCalculator.tsx",
])

dir_list("Challenges (features/challenges)", [
    "ChallengesView.tsx",
])

dir_list("Compliance (features/compliance)", [
    "ComplianceBadges.tsx", "ComplianceBar.tsx", "ConsentModal.tsx",
    "DemoControls.tsx", "KYCModal.tsx", "PrivacyCenter.tsx",
])

dir_list("Credit / CreditBridge AI (features/credit)", [
    "CreditBridgeAI.tsx", "CreditHealth.tsx", "DemoCreditCard.tsx",
    "services/creditBridgeEngine.ts", "services/ethicsEngine.ts", "services/lenderMarketplace.ts",
])

dir_list("Dashboard (features/dashboard)", [
    "AdaptiveInsight.tsx", "BehavioralNudges.tsx", "DashboardView.tsx", "DashboardWidget.tsx",
    "DeviceStatusCard.tsx", "FinancialPulse.tsx", "FinancialWeather.tsx", "KYCStatusCard.tsx",
    "MarketIntelligenceHero.tsx", "MonthlyNarrative.tsx", "NBAInsights.tsx", "NetWorthCard.tsx",
    "NotificationCenter.tsx", "PhysicalAssetsPromo.tsx", "PredictiveShieldBadge.tsx", "QuickActions.tsx",
    "StockTicker.tsx", "WealthBenchmark.tsx", "WealthDNA.tsx", "WealthTwinHero.tsx",
])

dir_list("Demo (features/demo)", [
    "BlockchainAudit.tsx", "CinematicIntro.tsx", "CursorSpotlight.tsx", "DemoAssistant.tsx",
    "DemoMode.tsx", "DemoPersonas.tsx", "DemoShowcase.tsx", "FeatureUniverse.tsx",
    "FraudRadar.tsx", "JudgeTour.tsx", "NotificationDemo.tsx",
])

dir_list("Family (features/family)", [
    "FamilyDashboard.tsx",
])

dir_list("Forecast (features/forecast)", [
    "ForecastView.tsx", "ScenarioSimulator.tsx", "WhatIfSimulator.tsx",
])

dir_list("Gamification (features/gamification)", [
    "BadgeStreak.tsx", "FantasyLeague.tsx", "FantasyPortfolio.tsx",
])

dir_list("Goals (features/goals)", [
    "AddGoalModal.tsx", "BoostCard.tsx", "BoostsManager.tsx", "DynamicCompass.tsx",
    "GoalCelebration.tsx", "GoalConflictIntelligence.tsx", "GoalConflictModal.tsx", "GoalTracker.tsx",
])

dir_list("Gold (features/gold)", [
    "DigitalGold.tsx",
])

dir_list("Income (features/income)", [
    "GigIncomeSmoother.tsx",
])

dir_list("Innovation Lab (features/innovation)", [
    "AIFutureTwin.tsx", "AIInsightsAggregator.tsx", "AutoInstrumentGenerator.tsx", "AutonomousAgent.tsx",
    "BhavishyaEngine.tsx", "ChakraBalance.tsx", "CollectiveImmuneSystem.tsx", "CommunityDNA.tsx",
    "CrisisPredictor.tsx", "DigitalInheritance.tsx", "DreamVisualizer.tsx", "EmotionalHeatmap.tsx",
    "EmotionalResonance.tsx", "FestivalAwareEngine.tsx", "FinancialDNAHelix.tsx", "FutureSelfSimulator.tsx",
    "GenerationalWealth.tsx", "InnovationLabView.tsx", "InnovationOverview.tsx", "LifeEventPredictor.tsx",
    "LifeShockSimulator.tsx", "MacroShockSimulator.tsx", "MarketIntelligence.tsx", "MonteCarloSimulator.tsx",
    "NeuralNetworkViz.tsx", "NeuroFrictionWidget.tsx", "ParticleConstellation.tsx", "PreparednessScore.tsx",
    "ProsperityScore.tsx", "QuantumLock.tsx", "SovereignVault.tsx", "TemporalWealth.tsx", "TimeMachine.tsx",
    "WealthWeather.tsx",
])

dir_list("Insurance (features/insurance)", [
    "ParametricInsurance.tsx",
])

dir_list("Kids (features/kids)", [
    "KidsMode.tsx",
])

dir_list("Loans (features/loans)", [
    "LoanImpactSimulator.tsx", "LoanResearchShowcase.tsx", "LoansHub.tsx", "SocialCollateralLoan.tsx",
])

dir_list("Location (features/location)", [
    "LocationVerifier.tsx",
])

dir_list("Market (features/market)", [
    "GlobalMacroRadar.tsx", "MacroSignalTower.tsx", "MarketNewsFeed.tsx",
    "MarketStrategist.tsx", "MarketView.tsx", "SmartTriggers.tsx", "useMacroFeed.ts",
])

dir_list("MSME (features/msme)", [
    "MSMEcreditbridgeView.tsx", "services/msmeScoreEngine.ts",
])

dir_list("NRI (features/nri)", [
    "NRIMode.tsx",
])

dir_list("Onboarding (features/onboarding)", [
    "OnboardingWizard.tsx",
])

dir_list("Payments (features/payments)", [
    "AdReward.tsx", "AnimatedTransactionToast.tsx", "BillSplitter.tsx", "CashbackPiggy.tsx",
    "CashbackWallet.tsx", "FloatingPayButton.tsx", "GroupJar.tsx", "MPINInput.tsx", "MoodMeter.tsx",
    "PaymentHub.tsx", "PaymentRequests.tsx", "PaymentsPage.tsx", "QrScannerSimulator.tsx",
    "ReferralSection.tsx", "RewardsDashboardCard.tsx", "SpinWheel.tsx", "StreakFire.tsx",
    "StreakTracker.tsx", "UPIPaymentSimulator.tsx", "VoicePayment.tsx",
])

dir_list("Pitch (features/pitch)", [
    "PitchDeckView.tsx", "PitchMode.tsx",
])

dir_list("Portfolio (features/portfolio)", [
    "ESGScore.tsx", "PortfolioView.tsx",
])

dir_list("Privacy (features/privacy)", [
    "PrivacyAuditPanel.tsx", "PrivacyView.tsx",
])

dir_list("Profile (features/profile)", [
    "ProfileSettings.tsx",
])

dir_list("Protection / Rakshak (features/protection)", [
    "RakshakInterventionChat.tsx", "BehavioralBiometrics.tsx", "CoercedModeBanner.tsx",
    "CoolingVaultModal.tsx", "CounterfactualPanel.tsx", "DuressMode.tsx", "DuressPinSetup.tsx",
    "FamilyApprovalGate.tsx", "FamilySafeWord.tsx", "FraudDetectionEngine.tsx", "LockdownOverlay.tsx",
    "MoneyMuleGraph.tsx", "OTPSimulation.tsx", "OTPSimulation.test.tsx", "PanicButton.tsx",
    "PaymentGuard.tsx", "ProtectionModal.tsx", "ProtectionView.tsx", "RiskMeter.tsx",
    "ScamCallerID.tsx", "SecureCheckout.tsx", "SecurityLog.tsx", "StressTestSimulator.tsx",
    "ThreatIntel.tsx", "TransactionGuardModal.tsx", "URLSafetyChecker.tsx",
])

dir_list("PSB (features/psb)", [
    "AccessibleFooter.tsx", "PSBLogo.tsx", "PSBSchemesCard.tsx", "QuickPayCard.tsx",
    "RecentTransactionsTable.tsx", "SecurityHealthWidget.tsx", "WelcomeBanner.tsx",
])

dir_list("Quiz (features/quiz)", [
    "InvestmentQuiz.tsx",
])

dir_list("Reports (features/report)", [
    "FinancialReport.tsx", "ReportGeneratorModal.tsx",
])

dir_list("Salary (features/salary)", [
    "AddSalaryModal.tsx",
])

dir_list("Security Beast (features/security)", [
    "AntiScamShield.tsx", "BlockchainAudit.tsx", "DeadMansSwitch.tsx", "DecentralizedId.tsx",
    "DecoyAccountView.tsx", "DeviceFingerprintPanel.tsx", "DuressTrigger.tsx", "EbpfMonitor.tsx",
    "GhostMode.tsx", "HoneytokenManager.tsx", "PasskeyAuth.tsx", "PostQuantumCrypto.tsx",
    "SecureEnclaveCheck.tsx", "SecurityBeastView.tsx", "SecurityScoreDashboard.tsx",
    "TpmAttestation.tsx", "TransactionTrap.tsx", "VoiceCommandBar.tsx",
])

dir_list("Senior (features/senior)", [
    "SeniorMode.tsx",
])

dir_list("Subscriptions (features/subscriptions)", [
    "AdminAgent.tsx", "CultFitTracker.tsx", "SubscriptionTracker.tsx",
])

dir_list("Tax (features/tax)", [
    "OldVsNewRegime.tsx", "Section80CTracker.tsx", "TaxCalculator.tsx",
    "TaxDeadlineCalendar.tsx", "TaxView.tsx",
])

dir_list("Transactions (features/transactions)", [
    "AICategorization.tsx", "EmotionCheckin.tsx", "ScanReceipt.tsx", "SmartDuplicateDetection.tsx",
    "TransactionComparison.tsx", "TransactionDetailModal.tsx", "TransactionTagger.tsx", "TransactionsView.tsx",
])

dir_list("Values (features/values)", [
    "ValuesAlignment.tsx",
])

H2("A.2 Authenticated App Views")
para("The main shell switches among these views:")
mini_list([
    "dashboard", "wealth-twin", "ai-recommendations", "goals", "portfolio", "family", "assets",
    "market", "forecast", "protection", "privacy", "tax", "calculators", "transactions", "features",
    "architecture", "bills", "credit-health", "creditbridge-ai", "notification-demo", "digital-gold",
    "challenges", "kids-mode", "subscriptions", "accessibility", "nri-mode", "business-mode",
    "values-alignment", "fantasy-league", "boosts", "security-beast", "bhavishya", "innovation-lab",
    "pitch-deck", "payments", "loan-center", "msme-creditbridge", "loans-hub", "loan-research",
    "loan-impact", "social-collateral-loan", "recurring-payments", "account-statement", "audit-log", "profile"
])

H2("A.3 Backend API Routes")
mini_list([
    "admin.js", "ai.js", "analytics.js", "auth.js", "banking.js", "business.js", "charts.js",
    "documents.js", "export.js", "extract.js", "financial-model.js", "fraud.js", "gallery.js",
    "gst.js", "kyc.js", "market-data.js", "msme.js", "nlp-query.js", "realData.js", "scenarios.js",
    "screener.js", "tax.js"
])

H2("A.4 Backend Algorithms")
mini_list([
    "gstinValidator.js", "itcRiskScanner.js", "missingITCRecovery.js", "msmeCreditScore.js",
    "shellCompanyDetector.js", "taxOptimizer.js", "taxRateErrorDetector.js"
])

H2("A.5 Backend Services")
mini_list([
    "ai-provider.js", "anomaly-detector.js", "cacheService.js", "chart-engine.js", "database.js",
    "dataIngestion.js", "dcf-engine.js", "demoData.js", "emailService.js", "market-data.js",
    "nlp-query.js", "paymentService.js", "scenario-engine.js", "screener-service.js", "websocket.js"
])

H2("A.6 Key Shared Services (client/src/shared/services/)")
mini_list([
    "aiOrchestrator.ts", "aiRouter.ts", "aiPrompts.ts", "aiConfig.ts",
    "fraudDetectionService.ts", "fingerprintService.ts", "behavioralBiometricsService.ts",
    "duressService.ts", "guardianService.ts", "blockchainService.ts", "auditLogger.ts",
    "nbaService.ts", "twinService.ts", "scenarioEngine.ts", "notificationService.ts",
    "cashbackEngine.ts", "streakService.ts", "leagueService.ts", "subscriptionPredictor.ts",
    "privacyAuditService.ts", "urlSafetyService.ts", "sentimentService.ts", "newsService.ts",
    "voiceService.ts", "passkeyService.ts", "webauthnService.ts", "postQuantumService.ts",
    "attestationService.ts", "browserThreatService.ts", "didService.ts", "offlineQueue.ts"
])

H2("A.7 Python / Document Automation Scripts")
mini_list([
    "build_final_document.py — this 100+ page comprehensive final document.",
    "build_comprehensive_doc.py — exhaustive PPT source with diagrams.",
    "build_ppt_source_doc.py — standard PPT source document.",
    "build_cuecard.py — one-page 10-minute demo cue card.",
    "build_playbook.py — detailed team demo playbook.",
    "build_tooltable.py — Hinglish demo tool/priority table.",
])

page_break()

# ---------------------------------------------------------------------------
# 21. Appendix B — Glossary & Abbreviations
# ---------------------------------------------------------------------------
H1("Appendix B — Glossary & Abbreviations")
table(
    ["Term", "Meaning"],
    [
        ["AA", "Account Aggregator — consent-based financial data sharing framework (RBI)."],
        ["AI", "Artificial Intelligence."],
        ["CIBIL", "Credit Information Bureau (India) Limited."],
        ["CBS", "Core Banking System."],
        ["DCF", "Discounted Cash Flow valuation method."],
        ["DPDP Act", "Digital Personal Data Protection Act, 2023."],
        ["EAA", "Explainability, Accountability, Auditability framework."],
        ["ELI5", "Explain Like I am 5 — simple explanation layer."],
        ["EMI", "Equated Monthly Instalment."],
        ["ESG", "Environmental, Social, Governance investing criteria."],
        ["GST", "Goods and Services Tax."],
        ["GSTIN", "GST Identification Number."],
        ["HRA", "House Rent Allowance."],
        ["ITC", "Input Tax Credit under GST."],
        ["KYC", "Know Your Customer."],
        ["MSME", "Micro, Small and Medium Enterprises."],
        ["NBA", "Next-Best-Action recommendation."],
        ["OTP", "One-Time Password."],
        ["PAT", "PSB Algorithm Toolkit — naming prefix for backend algorithms."],
        ["PSB", "Punjab & Sind Bank / Public Sector Bank."],
        ["RBI", "Reserve Bank of India."],
        ["SHAP", "SHapley Additive exPlanations for model interpretability."],
        ["SIP", "Systematic Investment Plan."],
        ["SME", "Small and Medium Enterprise."],
        ["TEAA", "Transparency, Equity, Accountability, Auditability framework."],
        ["UPI", "Unified Payments Interface."],
        ["WebAuthn", "Web Authentication API for passkeys."],
    ],
    widths=[1.0, 4.8]
)
page_break()

# ---------------------------------------------------------------------------
# 22. Appendix C — Project Metadata & Credits
# ---------------------------------------------------------------------------
H1("Appendix C — Project Metadata & Credits")
para("This document was generated automatically from the SecureWealth Twin codebase. The following metadata is accurate as of generation time.")

table(
    ["Field", "Value"],
    [
        ["Project name", "SecureWealth Twin"],
        ["Event", "PSBs Hackathon Series-2026"],
        ["Domain", "Project Management Office (PMO) — Cyber Security and Fraud"],
        ["Problem title", "SecureWealth Twin — Intelligent Wealth Growth with Built-in Fraud Protection"],
        ["Frontend stack", "React 19 + Vite 8 + TypeScript + Tailwind CSS v4 + Framer Motion + Recharts + Zustand"],
        ["Backend stack", "Node.js + Express + SQLite (better-sqlite3) + Redis fallback + WebSocket"],
        ["Live frontend", "https://psb-securewealth-frontend.onrender.com/"],
        ["Live backend", "https://psb-securewealth-backend.onrender.com/"],
        ["Repository", "https://github.com/financeforumglrc/psb-securewealth-frontend.git"],
        ["Document generator", "build_final_document.py"],
        ["Output file", OUTPUT],
    ],
    widths=[1.5, 4.2]
)

H2("C.1 How to use this document for the PPT")
numbered("Use Sections 1-5 for the problem, solution, and impact slides.")
numbered("Use Section 6 architecture diagrams as full-slide visuals.")
numbered("Pick 8-10 flagship features from Sections 7-12 for detailed product slides.")
numbered("Use Section 13-14 for technical depth and code-story slides.")
numbered("Use Section 16 for the demo script and speaker notes.")
numbered("Use Appendices for proof of completeness and Q&A backup.")

H2("C.2 End of document")
para("This concludes the SecureWealth Twin Final Comprehensive Document. All features, architecture, code rationale, and demo guidance are included to enable rapid PPT creation. For updates, rerun python build_final_document.py.")

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUTPUT)
print(f"Saved {OUTPUT} ({os.path.getsize(OUTPUT)} bytes)")
