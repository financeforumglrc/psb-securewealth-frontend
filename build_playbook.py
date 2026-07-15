# -*- coding: utf-8 -*-
"""Generates the PSB SecureWealth Twin — Team Demo Playbook (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0B, 0x2B, 0x52)
BLUE = RGBColor(0x1E, 0x5A, 0x9C)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
RED = RGBColor(0xB3, 0x1B, 0x1B)
AMBER = RGBColor(0xB5, 0x6A, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# base style
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(10.5)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def H1(text, color=NAVY):
    p = doc.add_heading(level=1)
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(17); r.bold = True
    return p

def H2(text, color=BLUE):
    p = doc.add_heading(level=2)
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(13.5); r.bold = True
    return p

def H3(text, color=NAVY):
    p = doc.add_heading(level=3)
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(11.5); r.bold = True
    return p

def para(text, size=10.5, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, bold_lead=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + 0.25*level)
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style='List Number')
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def table(headers, rows, widths=None, header_fill='0B2B52'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], header_fill)
        pr = hdr[i].paragraphs[0]; run = pr.add_run(h)
        run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            pr = cells[i].paragraphs[0]
            run = pr.add_run(str(val)); run.font.size = Pt(9.5)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def rule():
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'1E5A9C')
    pbdr.append(bottom); pPr.append(pbdr)

# ============================ COVER ============================
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SecureWealth Twin'); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = NAVY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Intelligent Wealth Growth with Built-in Fraud Protection'); r.font.size = Pt(15); r.font.color.rgb = BLUE; r.italic = True
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TEAM DEMO PLAYBOOK & TECHNICAL BRIEFING'); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RED
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PSBs Hackathon Series 2026  •  Domain: Cyber Security & Fraud'); r.font.size = Pt(11); r.font.color.rgb = GREY
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Read this fully before judging. Every teammate must be able to explain\nany screen, any feature, and how it is coded.'); r.font.size = Pt(11); r.italic = True; r.font.color.rgb = GREY
doc.add_page_break()

# ============================ TOC-ish ============================
H1('What is in this document')
para('This playbook is the single source of truth for the whole team. It covers, in order:')
sections = [
 ('1. Executive Summary', 'The 30-second and 2-minute version of what we built.'),
 ('2. Problem Statement -> Solution Map', 'Proof that every required item is covered.'),
 ('3. System Architecture & Tech Stack', 'How the pieces fit; what runs where.'),
 ('4. How the Code is Organized', 'Folder structure so anyone can find anything.'),
 ('5. Complete Feature Inventory', 'Every module on the website + how it is coded.'),
 ('6. The Protection Layer (CORE)', 'The mandatory cyber-security layer in depth.'),
 ('7. The 10-Minute Demo Run-of-Show', 'Exact tools to click, in order, with timings.'),
 ('8. Compliance & Responsible AI', 'Consent, explainability, KYC, no false claims.'),
 ('9. Judge Q&A Bank', 'Likely questions + the answer each teammate gives.'),
 ('10. Do-NOT-Show List', 'What to avoid on stage and why.'),
 ('11. Team Roles for the Demo', 'Who speaks to what.'),
 ('12. Scoring Cheat-Sheet', 'One line per metric to hit full marks.'),
]
for title, desc in sections:
    bullet(desc, bold_lead=title + ' — ')
doc.add_page_break()

# ============================ 1. EXEC SUMMARY ============================
H1('1. Executive Summary')
H3('The 30-second pitch (memorize this)')
para('"SecureWealth Twin is an AI-powered digital financial twin that helps customers grow wealth '
     'intelligently, while a MANDATORY, lightweight cyber-protection layer sits in front of every '
     'critical money action. The twin learns spending, goals and risk appetite, pulls a full financial '
     'picture via Account Aggregator, and gives explainable advice. Before any high-value action executes, '
     'our protection engine scores the risk from real signals — new device, rushed action, unusual amount, '
     'OTP misuse, coercion — and proportionately Allows, Warns with a cooling-off, or Blocks the action. '
     'Wealth is not only created, but safeguarded."', italic=True)
H3('The dual mandate (the heart of the problem)')
table(['Bank\'s Responsibility', 'Our Answer'],
      [['Grow wealth intelligently', 'Wealth Twin: behaviour learning, goals, net worth, AA, market-aware advice, future simulation'],
       ['Protect wealth from fraud & misuse', 'Mandatory Protection Layer: risk score -> Allow / Warn / Block on every critical action']],
      widths=[2.6, 4.2])
H3('Why we win')
bullet('We implemented the protection signals EXACTLY as the problem statement\'s Implementation Guide lists them (device, <10s timing, amount-vs-history, OTP retries, first-time action, behaviour).', bold_lead='Precision: ')
bullet('A fully deployed, working full-stack app (not slides) — frontend on Surge, backend on Render, Dockerised.', bold_lead='Completeness: ')
bullet('Every recommendation shows a plain-English reason; consent + KYC + privacy are built in.', bold_lead='Responsible AI: ')
bullet('Duress mode, behavioural biometrics, cooling-vault and scenario simulation go beyond basic budgeting tools.', bold_lead='Innovation: ')
doc.add_page_break()

# ============================ 2. PS MAP ============================
H1('2. Problem Statement -> Solution Map')
para('Each row is something the PDF explicitly asks for, and where we deliver it. Use this table to answer '
     '"does your solution do X?" — the answer is always yes, here:', space_after=8)
table(['Problem Statement Requirement', 'Where We Deliver It (screen / module)'],
 [['1. Learn from spending, saving, investment habits', 'Dashboard, useSpendingPersona, fraudDetectionService (real history analysis)'],
  ['2. Understand income, goals, risk appetite', 'GoalTracker, Wealth Twin profile, useRecommendationEngine'],
  ['3. Market/economic trends -> strategic advice (sell gold, shift to FD)', 'MarketView, MarketIntelligenceHero, GlobalMacroRadar, StockTicker'],
  ['4. Account Aggregator integration', 'AccountAggregatorFull + AAFetchAnimation (consent-based pull)'],
  ['5. Add property, gold, vehicles -> full net worth', 'ManualAssetForm, PhysicalAssetIntelligence, NetWorthCard'],
  ['6. Timely suggestions (portfolio, save, rebalance, tax)', 'AIRecommendationsView, NBAInsights, RecommendationCard'],
  ['7. Built-in risk checks BEFORE execution', 'Protection Layer: timingCheck + RiskMeter + TransactionGuardModal'],
  ['Wealth Protection Risk Score (Low/Med/High)', 'RiskMeter + fraudDetectionService scoring'],
  ['Proportionate action: Allow / Warn / Block', 'TransactionGuardModal, CoolingVaultModal, LockdownOverlay'],
  ['Predict future financial scenarios', 'ForecastView, MonteCarloSimulator, FutureSelfSimulator'],
  ['Simple, engaging UX + gamification', 'Clean dashboard, BadgeStreak, ChallengesView, FantasyLeague'],
  ['Responsible AI: privacy, consent, explainability', 'ConsentModal, PrivacyCenter, ExplainableTooltip, AIDecisionLog'],
  ['Corporate / business (optional)', 'BusinessMode (cash-flow, liquidity, surplus funds)'],
  ['KYC before real investment', 'KYCModal + backend /kyc routes'],
 ], widths=[3.5, 3.3])
doc.add_page_break()

# ============================ 3. ARCH ============================
H1('3. System Architecture & Tech Stack')
H3('The flow (mirrors the PDF architecture diagram)')
para('Customer (Web App)  ->  Wealth Intelligence Twin  ->  Wealth Protection Check  ->  '
     'Final Cyber-Risk Decision (Allow / Warn / Block)  ->  Action / Simulation  ->  Audit Log',
     bold=True, color=NAVY)
para('When a judge asks "what is your architecture?", draw or point to exactly this chain. Our code is '
     'organised around it.', italic=True, color=GREY)
H3('Technology stack')
table(['Layer', 'Technology', 'Why'],
 [['Frontend (UI)', 'React 18 + TypeScript + Vite', 'Fast, component-based, type-safe; suggested in PDF'],
  ['Styling', 'Tailwind CSS', 'Clean, responsive dashboard quickly'],
  ['State management', 'React Context (Auth, Security, NBA, Rewards)', 'Lightweight global state without heavy libraries'],
  ['Backend (API)', 'Node.js + Express', 'API-first, suggested in PDF; easy bank integration'],
  ['Database', 'SQLite (better-sqlite3) + optional Supabase', 'Zero-config persistence for the prototype'],
  ['Auth', 'JWT (HS256) + bcrypt password hashing', 'Stateless, industry-standard sessions'],
  ['AI', 'Multi-provider orchestrator (Gemini, Groq, OpenAI, Anthropic...)', 'Fallback + cost-aware routing, never a single point of failure'],
  ['Security middleware', 'helmet, express-rate-limit, CORS allow-list', 'Standard hardening on every request'],
  ['Deployment', 'Render (backend) + Surge (frontend) + Docker', 'Cloud-native; "local simulation using Docker" per PDF'],
 ], widths=[1.6, 2.8, 2.4])
H3('Key engineering facts to quote')
bullet('Authentication uses JWT signed with HS256 (algorithm pinned to prevent confusion attacks) and passwords hashed with bcrypt at cost factor 12.', bold_lead='Auth: ')
bullet('Every banking query is scoped to the logged-in user id (no user can read another user\'s data).', bold_lead='Data isolation: ')
bullet('Money movement uses an atomic SQLite transaction (executeTransfer) — balance check + debit + credit + record happen together or not at all.', bold_lead='Integrity: ')
bullet('The AI orchestrator has a circuit breaker — if a provider fails repeatedly it is skipped for 2 minutes and the next provider is used automatically.', bold_lead='Resilience: ')
doc.add_page_break()

# ============================ 4. CODE ORG ============================
H1('4. How the Code is Organized')
para('We use a feature-folder architecture. Each feature owns its components, so any teammate can find and '
     'explain code fast. If a judge says "show me the code for X", go to features/X.', space_after=8)
H3('Frontend  (client/src/)')
table(['Folder', 'What lives here'],
 [['app/', 'App.tsx, AuthenticatedApp.tsx — top-level shell and routing'],
  ['features/<name>/components/', 'All UI for a feature (e.g. protection/, dashboard/, ai/, payments/)'],
  ['shared/context/', 'AuthContext, SecurityContext, NBAContext, RewardsContext (global state)'],
  ['shared/services/', '50+ logic services: fraudDetectionService, duressService, totpService, aiOrchestrator, fingerprintService...'],
  ['shared/hooks/', 'Reusable logic: useProtectionEngine, useRecommendationEngine, useForecastEngine, useBehavioralBiometrics'],
  ['shared/lib/', 'backendApi.ts (calls Node backend), protectionApi.ts'],
  ['shared/data/', 'mockBankingData, userProfiles, aaBanks (demo data)'],
 ], widths=[2.2, 4.6])
H3('Backend  (backend/)')
table(['Folder', 'What lives here'],
 [['server.js', 'Express app: security middleware, rate limits, route mounting'],
  ['routes/', 'API endpoints: auth, banking, kyc, ai, market-data, tax, gst, charts...'],
  ['middleware/', 'auth.js (JWT), timingCheck.js (the <10s rule), errorHandler.js, auditLogger.js'],
  ['services/', 'database.js (SQLite), ai-provider.js, market-data.js, paymentService.js'],
  ['algorithms/', 'Financial algorithms (tax, GST tools) — not all used in this demo'],
 ], widths=[2.2, 4.6])
doc.add_page_break()

# ============================ 5. FEATURE INVENTORY ============================
H1('5. Complete Feature Inventory')
para('Grouped by area. For each major feature: what it does (say this to judges) and how it is coded '
     '(say this if they probe). Bold names are the actual components in our repo.', space_after=8)

H2('A. Wealth Intelligence Twin')
H3('Dashboard & Net Worth')
bullet('Clean landing screen showing balances, net worth, financial pulse, and the day\'s top insights.', bold_lead='What: ')
bullet('DashboardView composes NetWorthCard, FinancialPulse, FinancialWeather, NBAInsights. Net worth = bank balances + manually added assets (property/gold/vehicles), summed client-side.', bold_lead='How: ')
H3('Account Aggregator (AA)')
bullet('User consents, and we pull a consolidated financial picture from other banks — exactly PDF requirement #4.', bold_lead='What: ')
bullet('AccountAggregatorFull drives a consent flow; AAFetchAnimation shows the simulated fetch. Bank list from shared/data/aaBanks.', bold_lead='How: ')
H3('Physical Assets')
bullet('Add property, gold, vehicles; the twin folds them into net worth and risk view — PDF requirement #5.', bold_lead='What: ')
bullet('ManualAssetForm captures assets; PhysicalAssetIntelligence values/score them; persisted via backend /banking/assets (per-user).', bold_lead='How: ')
H3('Goals & Recommendations')
bullet('Goal-based planning ("Save Rs.500 more to reach your goal"), SIP/deposit/tax suggestions with reasons.', bold_lead='What: ')
bullet('GoalTracker + AddGoalModal store goals; useRecommendationEngine + AIRecommendationsView generate next-best-actions; each card carries an explanation.', bold_lead='How: ')
H3('Market Intelligence')
bullet('Market/economic context drives strategic advice (e.g., "shift to FD") — PDF requirement #3.', bold_lead='What: ')
bullet('MarketView, MarketIntelligenceHero, GlobalMacroRadar, StockTicker; backend /market route fetches quotes; useLivePrices hook on the client.', bold_lead='How: ')
H3('Future Simulation')
bullet('Projects future wealth and goal outcomes; runs scenario / Monte-Carlo style what-ifs.', bold_lead='What: ')
bullet('ForecastView + useForecastEngine; MonteCarloSimulator and FutureSelfSimulator run client-side simulations and chart the distribution.', bold_lead='How: ')

H2('B. The Mandatory Protection Layer (detailed in Section 6)')
bullet('Risk scoring engine + transaction guard + cooling vault + duress mode + behavioural biometrics. This is the core; full detail is in Section 6.')

H2('C. Conversational & Explainable AI')
bullet('WealthChat / FinancialTwinChat — chat with your twin in plain language.', bold_lead='What: ')
bullet('aiOrchestrator routes the prompt across providers (fallback / fastest-first / ensemble) with a circuit breaker; aiPrompts builds the context from the user profile.', bold_lead='How: ')
bullet('ExplainableTooltip, ELI5Tooltip, AIDecisionLog show WHY a recommendation was made — satisfies "Transparent & Explainable AI".', bold_lead='Explainability: ')

H2('D. Compliance & Privacy')
bullet('ConsentModal (explicit permission before using data), PrivacyCenter + PrivacyAuditPanel (what we collect & why), KYCModal (KYC before investment), ComplianceBar/Badges.', bold_lead='What: ')

H2('E. Engagement / UX')
bullet('BadgeStreak, ChallengesView, FantasyLeague (gamification), SeniorMode / KidsMode / AccessibilitySettings (inclusive UX), voice features.', bold_lead='What: ')

H2('F. Business Mode (optional in PDF)')
bullet('BusinessMode analyses cash-flow, liquidity and surplus-fund options for SME customers — bonus coverage of the corporate segment.', bold_lead='What: ')

H2('G. Innovation Lab (flash for 30s only)')
bullet('BhavishyaEngine & CrisisPredictor (predict shocks), LifeEventPredictor, WealthDNA / FinancialDNAHelix (behavioural fingerprint visual). Shown briefly to demonstrate innovation, not demoed deeply.', bold_lead='What: ')
doc.add_page_break()

# ============================ 6. PROTECTION LAYER ============================
H1('6. The Protection Layer  (THE CORE — know this cold)', RED)
para('The domain is "Cyber Security & Fraud". This layer is the single most important scored component. '
     'Every teammate must be able to explain it. It maps 1:1 to the PDF\'s "Cyber-Protection Layer — '
     'Implementation Guide".', bold=True)
H3('The six risk signals (we built every one the PDF lists)')
table(['PDF Protection Aspect', 'Our Implementation', 'Example Output'],
 [['1. Device Trust (trusted vs new)', 'fingerprintService creates a device id; first device trusted, new device raises score', '"New device detected - risk increased"'],
  ['2. Login & Session (rushed <10s)', 'timingCheck.js middleware: high-value action within 10s of login adds risk points', '"Action taken unusually fast"'],
  ['3. Amount vs history', 'fraudDetectionService flags amount > 2-3x the user\'s normal pattern', '"Amount higher than your normal pattern"'],
  ['4. OTP usage pattern', 'OTPSimulation + backend /otp; multiple retries raise risk', '"Multiple OTP attempts detected"'],
  ['5. First-time / new action', 'Boolean check: first SIP / new investment type -> moderate risk', '"First-time investment - extra review"'],
  ['6. Behaviour consistency', 'behavioralBiometricsService + counters track retry / cancel loops', '"Unusual retry pattern observed"'],
 ], widths=[2.1, 3.0, 1.7])
H3('The scoring engine (say this exactly)')
para('"Each signal carries a weight. We sum the weights into a Wealth Protection Risk Score, then apply '
     'thresholds: Low = Allow, Medium = Warn with a cooling-off message, High = Block / delay. '
     'In code this is fraudDetectionService.analyzeTransactions() on the client and timingCheck on the '
     'server, surfaced through the RiskMeter component."', italic=True)
table(['Risk Score', 'Decision', 'UI shown'],
 [['Low', 'Allow the action', 'Proceeds normally'],
  ['Medium', 'Warn + cooling-off', 'CoolingVaultModal / TransactionGuardModal'],
  ['High', 'Block / delay', 'LockdownOverlay / blocked confirmation']],
 widths=[1.6, 2.6, 2.6])
H3('Standout protections (our innovation in the security layer)')
bullet('A separate secret PIN. If forced/coerced, the user enters the duress PIN: the app shows a FAKE success, silently logs an alert, and locks the account for 24h. Directly answers the PDF\'s "coerced transactions" risk. Code: duressService + DuressMode + CoercedModeBanner.', bold_lead='Duress Mode: ')
bullet('Time-based one-time code on high-value transfers; entering a trap code freezes the account. Code: totpService + TransactionTrap.', bold_lead='Transaction Trap / TOTP: ')
bullet('A high-risk action can route into a cooling vault — a deliberate delay so impulsive/large actions get a second chance. Answers "impulsive high-value actions". Code: CoolingVaultModal.', bold_lead='Cooling Vault: ')
bullet('Every protected action writes an audit record (who, what, amount, device, time). Code: auditLogger + AuditLog screen.', bold_lead='Audit Trail: ')
para('NOTE for the team: in this prototype these protections are simulated client-side (the PDF explicitly '
     'allows "simulation/demo only" and "no real device monitoring needed"). If asked, say: "It is a working '
     'simulation of the protection logic; in production the same scoring runs server-side in front of the '
     'core banking transaction API."', italic=True, color=AMBER)
doc.add_page_break()

# ============================ 7. DEMO RUN OF SHOW ============================
H1('7. The 10-Minute Demo Run-of-Show', RED)
para('Follow this order exactly. Spend the most time on the Protection Layer. One person drives the screen, '
     'one person narrates (see Section 11). Times are cumulative.', bold=True)

def demo_block(time_range, title, color=NAVY):
    H2(f'{time_range}  —  {title}', color)

demo_block('0:00 - 0:45', 'Hook + Architecture')
numbered('Open the deployed app live (proves it is real).')
numbered('Show ONE architecture slide = the PDF flow (Customer -> Twin -> Protection -> Decision -> Action -> Audit).')
numbered('Say the 30-second pitch from Section 1.')
para('Scores: Architecture (Metric 6), Practicality (Metric 7).', italic=True, color=GREY)

demo_block('0:45 - 3:00', 'Wealth Intelligence Twin')
numbered('DashboardView + NetWorthCard — clean dashboard, full net worth.')
numbered('AccountAggregatorFull + AAFetchAnimation — consent, then pull other-bank data (PDF #4).')
numbered('ManualAssetForm / PhysicalAssetIntelligence — add gold/property/vehicle into net worth (PDF #5).')
numbered('GoalTracker — a goal with "save Rs.X more" guidance (PDF Outcome #3).')
numbered('AIRecommendationsView / NBAInsights — a recommendation WITH its reason shown.')
para('Scores: Wealth Intelligence (Metric 1), AI/Analytics (Metric 4), Explainability (Compliance).', italic=True, color=GREY)

demo_block('3:00 - 6:30', 'PROTECTION LAYER (spend the most time here)', RED)
numbered('Start a HIGH-VALUE action right after login (e.g., large transfer / new SIP).')
numbered('DeviceStatusCard / DeviceFingerprintPanel — "new device" signal.')
numbered('Point out the <10s timing flag (timingCheck) — "action taken unusually fast".')
numbered('RiskMeter shows amount-vs-history flag and assembles the score.')
numbered('OTPSimulation — show an OTP retry raising risk.')
numbered('Wealth Protection Risk Score appears: Low / Medium / High.')
numbered('TransactionGuardModal / CoolingVaultModal — Warn + cooling-off for Medium.')
numbered('DuressMode — enter duress PIN -> fake success + silent alert + lockdown (the "wow").')
numbered('AuditLog — show the action was recorded.')
para('Scores: Effectiveness of Wealth Protection (Metric 2 — the core), Innovation (Metric 5).', italic=True, color=GREY)

demo_block('6:30 - 8:00', 'Responsible AI + Compliance')
numbered('ExplainableTooltip / AIDecisionLog — "why this recommendation".')
numbered('ConsentModal / PrivacyCenter — consent + data transparency.')
numbered('KYCModal — KYC before investment.')
numbered('Point to the "Simulation / Demo only" badge.')
para('Scores: all four Compliance requirements, Metric 4.', italic=True, color=GREY)

demo_block('8:00 - 9:00', 'Innovation flash (60s, do not go deep)')
numbered('Open InnovationLabView; flash BhavishyaEngine / CrisisPredictor, FutureSelfSimulator, BusinessMode.')
numbered('Say: "40+ more modules — lightweight by default, deep on demand."')
para('Scores: Innovation (Metric 5), Simplicity framing (Metric 3).', italic=True, color=GREY)

demo_block('9:00 - 10:00', 'Scale + Close')
numbered('One line on architecture & scale: React + Node microservice + Docker + AA + mock CBS APIs.')
numbered('Close: "Wealth created AND safeguarded — exactly the dual mandate."')
para('Scores: Scalability & Practicality (Metric 7).', italic=True, color=GREY)
doc.add_page_break()

# ============================ 8. COMPLIANCE ============================
H1('8. Compliance & Responsible AI')
table(['PDF Compliance Requirement', 'How we satisfy it'],
 [['Customer data privacy & consent', 'ConsentModal before any data use; PrivacyCenter explains what & why'],
  ['Secure handling (encryption, no plaintext sensitive data)', 'JWT + bcrypt; per-user data isolation; sensitive PII masked / marked demo-only'],
  ['Transparent & explainable AI', 'ExplainableTooltip, ELI5Tooltip, AIDecisionLog show the reasoning'],
  ['No guaranteed returns / misleading claims', 'We show "suggestions / simulated outcomes", never promised profit'],
  ['Basic financial regulations (KYC)', 'KYCModal + /kyc routes; "for simulation/demo only" labelling']],
 widths=[3.2, 3.6])
para('Team reminder: if anyone asks about data security, lead with "consent-first, explainable, and we never '
     'store real sensitive data in plain text in the demo." Do NOT over-claim ("military grade", "unbreakable", '
     '"patented"). The PDF penalises misleading claims.', italic=True, color=AMBER)
doc.add_page_break()

# ============================ 9. Q&A ============================
H1('9. Judge Q&A Bank')
para('Likely questions and the crisp answer to give. Practise these out loud.', space_after=8)
qa = [
 ('Is this a real working app or mock-ups?',
  'Fully working and deployed — frontend on Surge, backend (Node + Express + SQLite) on Render. The protection signals run as a working simulation, which the problem statement explicitly allows.'),
 ('How exactly is the risk score calculated?',
  'Each signal (new device, <10s timing, amount vs history, OTP retries, first-time action, behaviour) has a weight. We sum them into a score and threshold it: Low=Allow, Medium=Warn/cooling-off, High=Block. Server side it is timingCheck.js; client side fraudDetectionService.analyzeTransactions().'),
 ('What stops a user reading another user\'s data?',
  'Every banking query is scoped to the authenticated user id from the JWT. There is no endpoint that returns data without the user filter; we verified ownership on every account/goal/transaction route.'),
 ('How do you handle authentication?',
  'JWT signed with HS256 (algorithm pinned), 7-day access token + refresh token stored server-side for revocation, passwords hashed with bcrypt cost 12, login/register rate-limited.'),
 ('What is the duress / coercion feature?',
  'A secret second PIN. If a customer is forced to transact, they enter the duress PIN: the app fakes a success screen, silently raises an alert, and locks the account 24h. It directly addresses the "coerced transactions" risk in the brief.'),
 ('How does the AI work / what if the AI provider is down?',
  'A multi-provider orchestrator routes the prompt with fallback and a circuit breaker — if one provider fails repeatedly it is skipped for 2 minutes and the next is used. No single point of failure.'),
 ('How is this explainable / not a black box?',
  'Every recommendation and every protection decision carries a plain-English reason shown in the UI (ExplainableTooltip / AIDecisionLog). Rules are transparent weighted signals, not a hidden model.'),
 ('How would this integrate with a real bank?',
  'API-first: Account Aggregator for data-in, mock CBS APIs for execution, and the protection layer is a stateless check that drops in front of any transaction endpoint. Dockerised for cloud deployment.'),
 ('How does it scale to millions of users?',
  'Stateless JWT auth and a stateless protection microservice scale horizontally; SQLite is the prototype store and swaps to Postgres/managed DB in production. AI calls are rate-limited and cached.'),
 ('What is genuinely innovative here vs a budgeting app?',
  'Duress mode, behavioural-biometric signals, cooling-vault for impulsive actions, and scenario/Monte-Carlo future simulation — protection and prediction, not just tracking.'),
 ('Where is the data stored and is it secure?',
  'Prototype data is in SQLite on the server; sensitive identifiers are masked or marked demo-only. We do not store real Aadhaar/card data in plain text in the demo.'),
 ('Can you show the code for the protection check?',
  'Yes — backend/middleware/timingCheck.js for the <10s rule and shared/services/fraudDetectionService.ts for the scoring. (Have these two files open in a tab.)'),
]
for q, a in qa:
    p = doc.add_paragraph(); r = p.add_run('Q: ' + q); r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10.5)
    p2 = doc.add_paragraph(); r2 = p2.add_run('A: ' + a); r2.font.size = Pt(10.5); p2.paragraph_format.space_after = Pt(8)
doc.add_page_break()

# ============================ 10. DO NOT SHOW ============================
H1('10. Do-NOT-Show List', RED)
para('Opening these live can break the demo or hurt the score. Avoid them on stage.', bold=True)
table(['Avoid', 'Why'],
 [['Backend scenario / NLP what-if engine', 'A runtime import bug can throw an error live'],
  ['Exact tax-saving rupee figures (Tax view)', 'A tax-rule edge case can show a wrong number; speak qualitatively instead'],
  ['GST / ITC / DCF / LBO financial-modelling tools', 'Not in THIS problem statement; signals "bolted-on" and wastes time'],
  ['Any "47 patents / military-grade / post-quantum / unbreakable" wording', 'Over-claiming is penalised; the PDF forbids misleading claims'],
  ['Opening 40 features one by one', 'Breaks the "Simplicity" metric; judges value a clean, focused flow']],
 widths=[3.0, 3.8])
para('If a judge specifically asks about one of these, answer briefly and honestly ("that module exists but '
     'is outside this problem\'s scope / is a work-in-progress"), then steer back to the protection layer.',
     italic=True, color=GREY)
doc.add_page_break()

# ============================ 11. TEAM ROLES ============================
H1('11. Team Roles for the Demo')
para('Assign these before you walk in. Fill names in the brackets.', space_after=8)
table(['Role', 'Owns', 'Name'],
 [['Driver', 'Controls the screen; clicks in the Section-7 order; never improvises navigation', '[ ____ ]'],
  ['Narrator / Lead', 'Says the pitch, frames each section, handles the close', '[ ____ ]'],
  ['Protection expert', 'Answers all Section-6 / security questions; has timingCheck.js + fraudDetectionService.ts open', '[ ____ ]'],
  ['AI & data expert', 'Answers AI orchestrator, explainability, recommendation questions', '[ ____ ]'],
  ['Architecture & scale', 'Answers stack, integration, scalability, compliance questions', '[ ____ ]']],
 widths=[1.7, 4.0, 1.1])
para('Everyone reads Sections 6, 7 and 9. The protection layer is the core — no teammate should be unable to '
     'explain it.', bold=True, color=RED)
doc.add_page_break()

# ============================ 12. CHEAT SHEET ============================
H1('12. Scoring Cheat-Sheet (the 7 metrics)')
para('One line to consciously hit each official success metric during the demo.', space_after=8)
table(['Metric', 'Hit it by...'],
 [['1. Quality of Wealth Intelligence', 'Showing a recommendation WITH a clear reason + goal guidance'],
  ['2. Effectiveness of Wealth Protection', 'The full risk-score -> Allow/Warn/Block flow + duress mode (spend most time)'],
  ['3. Simplicity & UX', 'Clean default dashboard; "lightweight by default, deep on demand"'],
  ['4. Use of AI / Data Analytics', 'Multi-provider orchestrator + signals computed from real history'],
  ['5. Innovation', 'Duress mode, cooling vault, behavioural biometrics, scenario simulation'],
  ['6. Technical Design & Architecture', 'Point to the PDF-matching flow diagram + feature-folder modularity'],
  ['7. Scalability & Practicality', 'Stateless auth + Docker + AA / mock CBS integration story']],
 widths=[2.8, 4.0])
rule()
para('Win condition: demo the protection layer flawlessly, stay clean and focused, pitch straight to these '
     'seven lines, and make no false claims.', bold=True, color=GREEN)

doc.save(r'G:/PSB/DS_Financial/frontend-render/SecureWealth_Twin_Demo_Playbook.docx')
print('SAVED: SecureWealth_Twin_Demo_Playbook.docx')
