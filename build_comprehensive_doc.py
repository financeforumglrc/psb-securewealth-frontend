# -*- coding: utf-8 -*-
"""Comprehensive PPT source Word document for SecureWealth Twin."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

OUTPUT = "SecureWealth_Twin_Comprehensive_PPT_Source.docx"
NAVY = RGBColor(0x0B, 0x2B, 0x52)
BLUE = RGBColor(0x1E, 0x5A, 0x9C)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
RED = RGBColor(0xB3, 0x1B, 0x1B)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(10.5)
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def H1(text, color=NAVY):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.font.size = Pt(20)
    r.bold = True
    return p


def H2(text, color=BLUE):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.font.size = Pt(15)
    r.bold = True
    return p


def H3(text, color=NAVY):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.font.size = Pt(12)
    r.bold = True
    return p


def para(text, size=10.5, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.add_run(text)
    return p


def mini_list(items):
    for it in items:
        bullet(it)


def table(headers, rows, widths=None, header_fill='0B2B52'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], header_fill)
        pr = hdr[i].paragraphs[0]
        rn = pr.add_run(h)
        rn.bold = True
        rn.font.color.rgb = WHITE
        rn.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            pr = cells[i].paragraphs[0]
            rn = pr.add_run(str(val))
            rn.font.size = Pt(9)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    return t


def feature_block(title, problem, capabilities, tech, benefit, demo_tip):
    H3(title)
    para("Problem it solves: " + problem)
    para("Key capabilities:")
    for c in capabilities:
        bullet(c)
    para("How it is built: " + tech)
    para("User benefit: " + benefit)
    p = doc.add_paragraph()
    r = p.add_run("Demo talking point: " + demo_tip)
    r.italic = True
    r.font.color.rgb = GREEN
    doc.add_paragraph()


def add_code_block(path, lines=35):
    if not os.path.exists(path):
        doc.add_paragraph("(file not found)")
        return
    with open(path, 'r', encoding='utf-8') as f:
        code = ''.join(f.readlines()[:lines])
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(code)
    r.font.name = 'Courier New'
    r.font.size = Pt(8)
    r.font.color.rgb = GREY


def generate_diagram(path, title, boxes, arrows):
    """boxes: list of (x,y,w,h,text,color); arrows: list of (x1,y1,x2,y2)"""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')
    for x, y, w, h, text, color in boxes:
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03", facecolor=color, edgecolor='black', linewidth=1.2)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha='center', va='center', fontsize=8, weight='bold', wrap=True)
    for x1, y1, x2, y2 in arrows:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle='->', color='#37474F', lw=1.5))
    ax.set_title(title, fontsize=13, weight='bold', pad=10)
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()


# ===================== TITLE =====================
para("SecureWealth Twin — Comprehensive PPT Source Document", size=24, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Every feature explained for PSB Hackathon Series 2026", size=13, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Technical judges / hackathon evaluators", size=11, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ===================== EXECUTIVE SUMMARY =====================
H1("Executive Summary")
para("SecureWealth Twin is a full-stack, AI-powered digital wealth intelligence system that acts as a virtual financial twin for retail customers and small businesses. It combines:")
bullet("Wealth Intelligence — personalized saving, investing, goal-planning, and market-aware recommendations.")
bullet("Mandatory Wealth Protection — a cyber-fraud layer that evaluates risk signals before every critical money action and decides Allow / Warn / Block.")
bullet("Responsible AI — consent-first data handling, explainable recommendations, audit logs, and regulatory alignment.")
bullet("CreditBridge AI — a flagship explainable credit-scoring engine for thin-file retail borrowers and MSMEs.")
para("The solution is deployed live, open-source, and built with a modern React/Node stack that can integrate with bank systems via APIs and Account Aggregator.")

# ===================== PROBLEM STATEMENT =====================
H1("1. Problem Statement")
para("The PSB Hackathon problem statement identifies a dual responsibility for banks:")
H2("1.1 Help customers grow wealth intelligently")
bullet("Young professionals and first-time investors need guided saving and investing habits.")
bullet("Families need goal-based planning for education, home purchase, and retirement.")
bullet("MSMEs need cash-flow visibility, surplus-fund management, and working-capital health.")
bullet("Customers expect personalized nudges instead of generic product pitches.")
H2("1.2 Protect wealth from fraud and risky actions")
bullet("Digital channels expose users to social-engineering fraud, OTP misuse, and coerced transactions.")
bullet("High-value or rushed actions need real-time risk evaluation.")
bullet("Protection must not add friction that confuses or scares users.")
H2("1.3 Regulatory and trust expectations")
bullet("DPDP Act 2023 and RBI guidelines require consent, explainability, and secure data handling.")
bullet("Banks need audit trails and transparent AI decision-making.")

# ===================== SOLUTION PILLARS =====================
H1("2. Solution Pillars")
feature_block(
    "Pillar 1: Digital Wealth Intelligence",
    "Customers receive generic advice and lack a holistic view of income, expenses, assets, and goals.",
    [
        "360° financial twin combining bank accounts, AA-linked accounts, physical assets, and goals.",
        "AI-driven next-best-actions with natural-language reasoning.",
        "Market-aware recommendations using real-world RBI, macro, and equity data.",
        "Scenario simulation and Monte-Carlo forecasting.",
    ],
    "React 19 + TypeScript frontend, Zustand global state, multi-provider AI orchestrator (Gemini, OpenAI, Anthropic, Groq, etc.), real-world data ingestion service.",
    "Users get actionable, personalized guidance that adapts to life changes and market conditions.",
    "Show the dashboard net-worth card, then click into Wealth Twin to demonstrate the twin 'knowing' the user.",
)
feature_block(
    "Pillar 2: Mandatory Wealth Protection Layer",
    "Critical money actions (large investments, SIP changes, fund transfers) are high-risk moments for fraud and impulsive mistakes.",
    [
        "Device trust check via browser fingerprinting.",
        "Timing check flags actions taken too quickly after login.",
        "Amount-vs-history anomaly detection.",
        "OTP retry and failure pattern analysis.",
        "Wealth Protection Risk Score maps to Allow / Warn (cooling-off) / Block.",
        "Duress PIN and transaction trap for coerced actions.",
    ],
    "fingerprintService.ts, timingCheck.js Express middleware, fraudDetectionService.ts rule engine, WebSocket fraud alerts, auditLogger middleware.",
    "Wealth actions are safeguarded proportionally without killing user experience.",
    "Start a high-value SIP modification on a new device and watch the risk meter escalate to cooling vault.",
)
feature_block(
    "Pillar 3: Responsible AI & Compliance",
    "Users and regulators demand transparency, consent, and safe data handling.",
    [
        "Consent modal before financial data use.",
        "Explainable tooltips and AI decision logs.",
        "KYC gating before real investment actions.",
        "Privacy audit panel and data transparency.",
        "No guaranteed-return or zero-risk claims.",
    ],
    "React Context for consent state, KYCModal component, backend /kyc routes, audit logs in SQLite.",
    "Trust is built through transparency; compliance is demonstrable.",
    "Hover over an AI recommendation to reveal the exact reasoning and data sources.",
)

# ===================== TECH ARCHITECTURE =====================
H1("3. Technology Architecture")
H2("3.1 Stack Overview")
table(
    ["Layer", "Technology"],
    [
        ["Frontend", "React 19.2.5, TypeScript, Vite 8, Tailwind CSS v4, Framer Motion, Recharts, Lucide React, Zustand"],
        ["Backend", "Node.js 22, Express, SQLite (better-sqlite3), JWT, bcryptjs, Helmet, express-rate-limit, Winston, WebSocket (ws)"],
        ["AI", "Multi-provider orchestrator: Gemini, Groq, OpenAI, Anthropic, Cohere, Mistral, DeepSeek, Grok, NVIDIA, Hugging Face, OpenRouter"],
        ["Security", "ML-KEM-768 post-quantum KEM, WebAuthn/Passkeys, FingerprintJS, behavioral biometrics, duress PIN, transaction trap"],
        ["Data Ingestion", "Yahoo Finance, RBI/World Bank, Alpha Vantage, Kaggle metadata"],
        ["Deployment", "Render (backend + static site), Surge (frontend), Docker, GitHub Actions CI/CD"],
    ],
    widths=[1.8, 5.2],
)
H2("3.2 System Architecture Diagram")
boxes = [
    (0.2, 5.0, 1.6, 0.7, 'Customer\nWeb / Mobile', '#E3F2FD'),
    (2.6, 5.0, 1.8, 0.7, 'React 19 SPA\nVite + Tailwind', '#FFF3E0'),
    (5.0, 5.0, 1.8, 0.7, 'Node.js API\nExpress', '#E8F5E9'),
    (7.4, 5.0, 1.4, 0.7, 'SQLite\nDatabase', '#F3E5F5'),
    (9.4, 5.0, 0.6, 0.7, 'AA', '#E8EAF6'),
    (2.6, 3.8, 1.8, 0.7, 'Protection\nRisk Engine', '#FFCDD2'),
    (5.0, 3.8, 1.8, 0.7, 'AI Orchestrator\nMulti-provider', '#E0F7FA'),
    (7.4, 3.8, 1.4, 0.7, 'Real-World\nData Feed', '#FFEBEE'),
    (2.6, 2.6, 1.8, 0.7, 'Admin Portal\nFraud Ops', '#F1F8E9'),
    (5.0, 2.6, 1.8, 0.7, 'Audit Logger\nImmutable', '#ECEFF1'),
]
arrows = [
    (1.8, 5.35, 2.6, 5.35),
    (4.4, 5.35, 5.0, 5.35),
    (6.8, 5.35, 7.4, 5.35),
    (8.8, 5.35, 9.4, 5.35),
    (5.9, 5.0, 5.9, 4.5),
    (8.1, 4.5, 8.1, 5.0),
    (3.5, 4.5, 3.5, 3.8),
    (5.9, 3.8, 5.9, 3.3),
    (6.8, 3.1, 5.4, 3.1),
]
generate_diagram("securewealth_arch.png", "SecureWealth Twin — End-to-End Architecture", boxes, arrows)
doc.add_picture("securewealth_arch.png", width=Inches(6.5))
H2("3.3 Data Flow")
table(
    ["Step", "Flow"],
    [
        ["1", "User logs in; frontend collects device fingerprint and stores JWT access token."],
        ["2", "Frontend requests user data from Node API with Authorization header and X-Device-Id."],
        ["3", "Backend validates JWT, scopes data per user, logs request to audit table."],
        ["4", "Account Aggregator consent flow pulls multi-bank data on user request."],
        ["5", "AI orchestrator generates recommendations using user profile + market data."],
        ["6", "Critical actions are evaluated by fraud/protection engine; risk score decides Allow/Warn/Block."],
        ["7", "Admin portal receives WebSocket fraud alerts and displays case management."],
    ],
    widths=[0.6, 6.4],
)

# ===================== WEALTH INTELLIGENCE FEATURES =====================
H1("4. Wealth Intelligence Features")
feature_block(
    "4.1 Dashboard",
    "Users need an at-a-glance view of their complete financial health.",
    [
        "Net worth summary with bank, investment, physical asset totals.",
        "Financial pulse score and wealth DNA visualization.",
        "Live market hero with indices and macro signals.",
        "Next-best-action cards and quick-action buttons.",
        "Notification center and stock ticker.",
    ],
    "React dashboard components, Zustand store, market data hooks, Framer Motion widgets.",
    "Users instantly understand where they stand financially and what to do next.",
    "Point to the net-worth number and explain it includes AA-linked + manual assets.",
)
feature_block(
    "4.2 Wealth Twin (AI Financial Twin)",
    "Generic robo-advisors do not adapt to individual life context and risk appetite.",
    [
        "Overview tab shows twin-derived financial health and anomalies.",
        "Goals tab aligns recommendations with user-defined goals.",
        "Rebalance tab suggests portfolio shifts based on market/risk changes.",
        "Retirement tab projects long-term corpus needs.",
        "Tax tab surfaces tax-saving opportunities.",
        "What-If tab simulates life changes (job loss, bonus, market crash).",
        "Conversational twin interface for natural-language queries.",
    ],
    "React feature-folder under ai/components/wealthTwin, multi-provider AI orchestrator, forecast engine, useRecommendationEngine hook.",
    "A personalized, evolving virtual twin that reasons like a financial advisor.",
    "Ask the twin: 'Can I afford a car?' and show the structured answer with assumptions.",
)
feature_block(
    "4.3 BHAVISHYA AI (Predictive Life-Cycle Engine)",
    "Traditional planning assumes a static future; real life has shocks and milestones.",
    [
        "Forecasts major life events (marriage, education, health, job change).",
        "Models financial shocks (market crash, inflation spike, medical emergency).",
        "Projects cash-flow impact of each scenario.",
        "Suggests proactive adjustments (buffer, insurance, rebalancing).",
        "Visual timeline of future financial states.",
    ],
    "InnovationLab BhavishyaEngine component, Monte-Carlo style simulation, scenario engine, Recharts timeline.",
    "Users prepare for uncertainty instead of reacting to it.",
    "Run a 'medical emergency at age 45' scenario and show the recommended emergency fund increase.",
)
feature_block(
    "4.4 Goal Tracker",
    "Goals fail because they are not tracked, prioritized, or connected to cash flow.",
    [
        "Create goals for education, home, retirement, travel, etc.",
        "AI coaching suggests monthly contribution adjustments.",
        "Conflict intelligence warns when goals compete for the same cash flow.",
        "Boost cards accelerate goals through extra income or expense cuts.",
        "Celebration animations on milestones.",
    ],
    "GoalTracker component, wealth store goals slice, useRecommendationEngine, Framer Motion celebrations.",
    "Users stay motivated and on track toward multiple life goals simultaneously.",
    "Show a retirement goal and the AI coach saying 'Save ₹2,000 more per month to be on track'.",
)
feature_block(
    "4.5 Portfolio View",
    "Users hold assets across institutions and cannot see allocation or drift.",
    [
        "Consolidated holdings view across banks, demat, and physical assets.",
        "Asset allocation pie chart and drift alerts.",
        "Rebalancing suggestions with reason.",
        "ESG scoring for investments.",
    ],
    "PortfolioView component, Recharts, market data service, ESG scoring logic.",
    "Users maintain desired risk/return profile and align investments with values.",
    "Show the allocation drift alert and the one-click rebalance suggestion.",
)
feature_block(
    "4.6 Assets Module",
    "Net worth is incomplete without physical and external assets.",
    [
        "Link bank accounts via RBI Account Aggregator.",
        "Manually add property, gold, vehicles, and other assets.",
        "Physical Asset Intelligence estimates valuations.",
        "Vision Appraisal uses image upload for asset appraisal.",
    ],
    "ManualAssetForm, LinkAccountModal, PhysicalAssetIntelligence, VisionAppraisalModal, backend /banking/assets routes.",
    "A true 360° net worth picture powers better advice.",
    "Add a gold asset via photo upload and watch net worth update instantly.",
)
feature_block(
    "4.7 Market View",
    "Retail investors lack timely, contextual market and macro insights.",
    [
        "Live Indian and global indices.",
        "Global macro radar (RBI rates, inflation, USD-INR, oil, gold).",
        "Market news and strategist signals.",
        "Smart triggers based on user portfolio exposure.",
    ],
    "MarketView component, real-world data ingestion (RBI, Yahoo Finance, Alpha Vantage), WebSocket/live polling.",
    "Users make informed decisions with contextual market intelligence.",
    "Show the RBI repo-rate card and explain how it affects FD vs. equity recommendations.",
)
feature_block(
    "4.8 Forecast & What-If Simulator",
    "Users cannot see the long-term impact of spending/saving decisions.",
    [
        "Monte Carlo simulation of wealth outcomes.",
        "What-if scenarios: salary change, expense shock, market correction.",
        "Probability-based goal achievement timeline.",
    ],
    "ForecastView, MonteCarloSimulator, scenario engine, Recharts projections.",
    "Users visualize trade-offs and choose better financial paths.",
    "Reduce monthly savings by ₹1,000 and show the retirement corpus impact.",
)

# ===================== PAYMENTS & PROTECTION =====================
H1("5. Payments, Transactions & Protection")
feature_block(
    "5.1 Payments Hub",
    "Users need a safe, unified place to pay and transfer money.",
    [
        "UPI, NEFT, card, and QR simulators.",
        "Bill splitter and group jar for shared expenses.",
        "Voice payment interface.",
        "MPIN and rewards/streaks/cashback wallet.",
    ],
    "PaymentsPage component, simulated payment flows, reward context, voice narration hook.",
    "Convenient payments with gamification and safety hooks.",
    "Make a high-value voice payment and trigger the protection layer.",
)
feature_block(
    "5.2 Transactions Intelligence",
    "Transactions are raw data; users need insight and control.",
    [
        "AI auto-categorization of spending.",
        "Emotion check-in at transaction time.",
        "Receipt scanning and duplicate detection.",
        "Transaction comparison and smart tagging.",
    ],
    "TransactionsView, AI extraction service, OCR via receipt scanning, smart tags.",
    "Users understand spending behavior and detect anomalies.",
    "Show the emotion tag cloud: 'How did you feel while spending this month?'.",
)
feature_block(
    "5.3 Protection Hub (Core Fraud Layer)",
    "Critical wealth actions need real-time risk evaluation without confusing users.",
    [
        "Wealth Protection Risk Meter combining multiple signals.",
        "OTP simulation with retry/failure analysis.",
        "Cooling vault for medium-risk actions.",
        "Duress mode: secret PIN triggers fake success and account lockdown.",
        "Panic button and lockdown overlay.",
        "Threat intelligence and behavioral biometrics.",
        "URL safety and secure checkout checks.",
    ],
    "fraudDetectionService.ts, timingCheck.js, RiskMeter, CoolingVaultModal, LockdownOverlay, duressService.ts, auditLogger.",
    "Wealth actions are protected proportionally; users feel safe, not blocked.",
    "Enter duress PIN during a fake high-value transfer and show silent alert + lockdown.",
)
feature_block(
    "5.4 Security Beast (10-Layer Security Dashboard)",
    "Banks need defense-in-depth against modern cyber threats.",
    [
        "TPM attestation and secure enclave simulation.",
        "eBPF-style monitor and honeytokens.",
        "Passkeys / WebAuthn login.",
        "Post-quantum cryptography (ML-KEM-768).",
        "Behavioral biometrics and decentralized identity.",
        "Transaction traps and blockchain audit trail.",
    ],
    "SecurityBeastView, mlkem library, WebAuthn service, blockchainService, FingerprintJS.",
    "A demonstrable, layered security posture for next-gen banking.",
    "Show the 10-layer radar and highlight post-quantum crypto as future-proofing.",
)
feature_block(
    "5.5 Privacy Center",
    "Users want control over their financial data.",
    [
        "Privacy audit panel showing what data is used and why.",
        "Consent controls and data transparency settings.",
        "One-click revoke consent flows.",
    ],
    "PrivacyView, ConsentModal, privacy audit components, backend consent records.",
    "Users trust the platform because they control their data.",
    "Open Privacy Center and list every data source with consent status.",
)
feature_block(
    "5.6 Audit Log",
    "Regulators and banks need immutable records of sensitive actions.",
    [
        "Logs every high-value action, protection decision, and consent event.",
        "Timestamped and user-scoped.",
        "Exportable for compliance review.",
    ],
    "AuditLog component, auditLogger middleware, SQLite audit table.",
    "Complete auditability for compliance and dispute resolution.",
    "Show the audit entry generated by the duress PIN event.",
)

# ===================== FINANCIAL TOOLS =====================
H1("6. Financial Tools")
feature_block(
    "6.1 Tax Planner",
    "Tax optimization is complex and last-minute for most users.",
    [
        "Old vs. new regime comparison.",
        "Section 80C tracker and deadline calendar.",
        "Tax-saving suggestions.",
        "Downloadable tax report.",
    ],
    "TaxView, tax calculation logic, backend /tax routes.",
    "Users save tax legally and on time.",
    "Compare old vs. new regime for a sample user and show savings.",
)
feature_block(
    "6.2 Calculators",
    "Users need quick answers to common financial questions.",
    [
        "EMI calculator for loans.",
        "SIP calculator for mutual funds.",
        "FD calculator.",
        "Retirement and rent-vs-buy calculators.",
    ],
    "CalculatorsView, math utilities, interactive sliders.",
    "Users make informed decisions without Excel.",
    "Run an SIP calculator and show corpus at 12% CAGR over 15 years.",
)
feature_block(
    "6.3 Bill Calendar",
    "Late fees and missed payments hurt credit health.",
    [
        "Upcoming due dates across credit cards, utilities, loans.",
        "Autopay reminders.",
        "Recurring bill tracking.",
    ],
    "BillCalendar component, bill reminders, backend /banking/bills routes.",
    "Users never miss a payment.",
    "Show the calendar with color-coded upcoming bills.",
)
feature_block(
    "6.4 Loan Center",
    "Loan tracking is fragmented across lenders.",
    [
        "Apply for new loans.",
        "Track existing loan status.",
        "EMI schedule and outstanding balance.",
    ],
    "LoanCenter component, backend /banking/loans routes.",
    "Centralized loan management inside the bank app.",
    "Show an active loan card with remaining tenure and interest saved on prepayment.",
)

# ===================== LOANS & CREDIT =====================
H1("7. Loans & Credit Features")
feature_block(
    "7.1 Credit Health",
    "Users need visibility into their CIBIL-style credit score and drivers.",
    [
        "Score display with color-coded risk band.",
        "Factor breakdown bars.",
        "Credit simulator for utilization/payoff scenarios.",
        "Improvement tips based on weakest factor.",
    ],
    "CreditHealth component, cibilScore/cibilFactors from Zustand store, Recharts.",
    "Users improve credit health proactively.",
    "Show the factor that is dragging the score and the one action to fix it.",
)
feature_block(
    "7.2 CreditBridge AI (Flagship)",
    "Traditional bureau-only models exclude thin-file Indians and MSMEs; black-box decisions create distrust.",
    [
        "Dual scoring for Retail and MSME using alternate data.",
        "Transparent 300–900 score with exposed weights.",
        "SHAP-style factor breakdown and score waterfall.",
        "Rejection reason codes with actionable steps.",
        "Counterfactual What-If simulator.",
        "Bureau-only comparison showing coverage gain.",
        "Ethics dashboard: governance, consent ledger, human review queue, adverse action notice.",
        "Lender marketplace with smart matching, filters, and EMI calculator.",
        "PDF explanation report export.",
    ],
    "creditBridgeEngine.ts, ethicsEngine.ts, lenderMarketplace.ts, React + Recharts, jsPDF export, Zustand.",
    "Fair, explainable credit access for thin-file borrowers and MSMEs.",
    "Switch to MSME mode, load 'Strong MSME' persona, and show matched lender offers.",
)
feature_block(
    "7.3 MSME CreditBridge",
    "Small businesses lack collateral and formal credit history.",
    [
        "Collateral-free MSME scoring using GST, bank cash-flow, and digital payment data.",
        "Document upload and verification flow.",
        "Offer generation and impact audit.",
    ],
    "MSMEcreditbridgeView, msmeScoreEngine.ts, backend /msme routes.",
    "MSMEs get faster, fairer credit decisions.",
    "Upload GST returns and show an instant score improvement.",
)
feature_block(
    "7.4 Loans Hub",
    "Users discover loan products across a cluttered interface.",
    [
        "One landing page for all loan/credit products.",
        "Cards linking to CreditBridge, MSME, social collateral, impact simulator, research.",
    ],
    "LoansHub component, navigation config.",
    "Easy product discovery and cross-sell.",
    "Click through Loans Hub to CreditBridge AI.",
)
feature_block(
    "7.5 Social Collateral Loan",
    "Thin-file borrowers can leverage community trust.",
    [
        "Trust-circle backed lending.",
        "Community vouching and risk sharing.",
        "Transparent terms and repayment tracking.",
    ],
    "SocialCollateralLoan component, trust-circle logic.",
    "Access to credit through social proof instead of collateral.",
    "Create a trust circle and show the enhanced loan offer.",
)
feature_block(
    "7.6 Loan Impact Simulator",
    "Banks and policymakers want to see credit impact beyond the borrower.",
    [
        "Estimates jobs created, GDP contribution, and income growth from MSME credit.",
        "Scenario-based impact visualization.",
    ],
    "LoanImpactSimulator component, impact model, Recharts.",
    "Demonstrates the real-economy value of MSME lending.",
    "Disburse ₹10L to a manufacturing MSME and show estimated jobs supported.",
)
feature_block(
    "7.7 Loan Research Showcase",
    "Academic grounding builds credibility for AI credit decisions.",
    [
        "Displays published research on algorithmic accountability.",
        "Links EAA framework to CreditBridge AI design.",
    ],
    "LoanResearchShowcase component, research content.",
    "Shows regulatory and ethical alignment.",
    "Open the research showcase and connect it to the bias-audit tab in CreditBridge.",
)

# ===================== FAMILY & LIFESTYLE =====================
H1("8. Family, Lifestyle & Gamification")
feature_block(
    "8.1 Family Dashboard",
    "Household finances are shared but banking apps are individual.",
    [
        "Shared wallets and allowances.",
        "Family goals and spending visibility.",
        "Parental controls and insights.",
    ],
    "FamilyDashboard component, shared state, permissions.",
    "Families plan and save together.",
    "Set a family vacation goal and split contributions.",
)
feature_block(
    "8.2 Digital Gold",
    "Users want simple, safe gold investment.",
    [
        "Buy, sell, and gift 24K gold digitally.",
        "Live gold price tracking.",
        "SIP in gold.",
    ],
    "DigitalGold component, market price integration.",
    "Convenient gold accumulation and gifting.",
    "Buy ₹500 worth of digital gold and show the vault balance.",
)
feature_block(
    "8.3 Subscriptions Manager",
    "Recurring subscriptions drain money unnoticed.",
    [
        "Track all subscriptions in one place.",
        "Detect unused subscriptions.",
        "Cancel reminders and annual-vs-monthly savings.",
    ],
    "SubscriptionTracker component, subscription list, savings calculator.",
    "Users stop paying for unused services.",
    "Highlight a subscription the user has not used in 3 months.",
)
feature_block(
    "8.4 Challenges",
    "Saving is hard without motivation.",
    [
        "Gamified savings challenges.",
        "Streaks, leaderboards, and rewards.",
        "Goal-linked challenge creation.",
    ],
    "ChallengesView component, gamification context.",
    "Users build habits through friendly competition.",
    "Join a 'No-Spend Weekend' challenge and track progress.",
)
feature_block(
    "8.5 NRI Center",
    "NRIs need specialized cross-border tools.",
    [
        "NRE/NRO account overview.",
        "Remittance planning and FEMA guidance.",
        "India investment tracking.",
    ],
    "NRIMode component, NRI-specific recommendations.",
    "NRIs manage India wealth compliantly.",
    "Show the remittance cost estimator for USD→INR.",
)
feature_block(
    "8.6 SME Centre / Business Mode",
    "Small businesses need treasury and cash-flow visibility.",
    [
        "Cash-flow timeline and surplus-fund advisor.",
        "Working-capital health score.",
        "Short-term investment suggestions for surplus cash.",
    ],
    "BusinessMode component, cash-flow engine, backend /business routes.",
    "MSMEs optimize liquidity and growth.",
    "Show the working-capital health score and a surplus-fund FD suggestion.",
)
feature_block(
    "8.7 Kids Mode",
    "Children need safe, educational money experiences.",
    [
        "Simplified UI with safe money lessons.",
        "Pocket money tracking and chores.",
        "Savings goals for kids.",
    ],
    "KidsMode component, child-friendly design.",
    "Early financial literacy in a protected environment.",
    "Set a 'new bicycle' goal and show the kid-friendly progress bar.",
)
feature_block(
    "8.8 Senior Mode",
    "Seniors need larger text, simpler flows, and safety.",
    [
        "Accessible UI with high contrast and large fonts.",
        "Simplified navigation.",
        "Voice narration support.",
    ],
    "SeniorMode component, accessibility hooks, voice narration.",
    "Elderly users bank independently and safely.",
    "Toggle Senior Mode and show the simplified dashboard.",
)
feature_block(
    "8.9 Fantasy League",
    "Investing education can be engaging and low-risk.",
    [
        "Portfolio fantasy league with virtual money.",
        "Learn by competing on returns.",
        "Leaderboards and portfolio analytics.",
    ],
    "FantasyLeague component, gamified portfolio engine.",
    "Users learn investing without risking real money.",
    "Show a fantasy portfolio outperforming the benchmark.",
)

# ===================== INNOVATION LAB & AI =====================
H1("9. Innovation Lab, AI & Administration")
feature_block(
    "9.1 Innovation Lab",
    "Banks need a sandbox to showcase futuristic capabilities.",
    [
        "Future Self Simulator — visualize future financial self.",
        "Crisis Predictor and Macro-Shock Simulator.",
        "Wealth Weather — market mood indicator.",
        "Emotional Resonance and Neural Network Visualizer.",
        "Time Machine, Generational Wealth, Prosperity Score.",
        "Quantum Lock and Sovereign Vault experiments.",
    ],
    "InnovationLabView, multiple experimental components, client-side simulation.",
    "Demonstrates innovation pipeline and WOW factor for judges.",
    "Open Future Self Simulator and show the age-progressed avatar with savings advice.",
)
feature_block(
    "9.2 AI Recommendations",
    "Generic product recommendations erode trust.",
    [
        "Next-best-actions based on user profile, goals, and market.",
        "Reasoning tooltips and AI decision logs.",
        "ELI5 explanations for non-technical users.",
    ],
    "AIRecommendationsView, aiOrchestrator.ts, explainability components.",
    "Users receive relevant, trustworthy advice.",
    "Hover over a recommendation to reveal the exact data points used.",
)
feature_block(
    "9.3 Values Alignment",
    "Investors increasingly want values-based portfolios.",
    [
        "Filter recommendations by ESG, ethical, or religious values.",
        "Values score for current portfolio.",
    ],
    "ValuesAlignment component, values scoring logic.",
    "Users invest in line with personal beliefs.",
    "Show the ESG alignment score of a sample portfolio.",
)
feature_block(
    "9.4 Admin Panel",
    "Bank ops and fraud teams need visibility and control.",
    [
        "Standalone /admin portal with secure login.",
        "Fraud Intelligence Center: case explorer, heatmap, trace graph.",
        "Live fraud alerts via WebSocket.",
        "System health, user activity, and audit exports.",
    ],
    "Admin portal components, backend /admin and /fraud routes, WebSocket alerts.",
    "Operations teams monitor and respond to threats in real time.",
    "Show a live fraud alert arriving in the admin panel.",
)
feature_block(
    "9.5 Onboarding Wizard",
    "First-time users need guided setup.",
    [
        "Step-by-step income, goals, risk profiling.",
        "Consent collection.",
        "Account Aggregator linking prompt.",
    ],
    "OnboardingWizard component, multi-step form.",
    "Users complete setup quickly and compliantly.",
    "Walk through onboarding and show the AA linking animation.",
)
feature_block(
    "9.6 Demo Mode & Judge Tour",
    "Demo environments must be judge-friendly and fault-tolerant.",
    [
        "Public /demo showcase.",
        "Judge tour with highlighted steps.",
        "Cinematic intro and demo personas.",
        "Cursor spotlight and fraud radar demo.",
    ],
    "DemoMode component, JudgeTour, demo personas, notification demo.",
    "Smooth, guided demo experience for evaluators.",
    "Launch demo mode and follow the judge tour.",
)
feature_block(
    "9.7 Accessibility & i18n",
    "Banking must be inclusive.",
    [
        "Accessibility settings (font size, contrast, voice).",
        "Hinglish/English i18n support.",
        "Screen-reader friendly toggle states and labels.",
    ],
    "AccessibilitySettings, i18n files, aria attributes.",
    "Users of all abilities and languages can use the app.",
    "Switch to Hinglish and increase font size.",
)

# ===================== ADDITIONAL / GRANULAR FEATURE INVENTORY =====================
H1("10. Additional & Granular Feature Inventory")
para("This section captures sub-components and capabilities that are part of the modules above but deserve explicit mention so that no feature is missed in the PPT.")

H2("10.1 Dashboard Intelligence Components")
para("The Dashboard is composed of many intelligent widgets, each delivering a specific insight:")
mini_list([
    "AdaptiveInsight — contextual insight based on user behavior.",
    "BehavioralNudges — small behavioral prompts to improve habits.",
    "FinancialPulse — one-number financial health score.",
    "FinancialWeather — market/regime mood indicator.",
    "KYCStatusCard — KYC progress and pending actions.",
    "MarketIntelligenceHero — top macro/market callout.",
    "MonthlyNarrative — auto-generated month-in-review story.",
    "NBAInsights — next-best-action cards.",
    "NetWorthCard — total net worth with breakdown.",
    "NotificationCenter — alerts, reminders, and fraud warnings.",
    "PhysicalAssetsPromo — prompt to add property/gold/vehicles.",
    "PredictiveShieldBadge — proactive risk status.",
    "QuickActions — one-tap frequent actions.",
    "StockTicker — live market ticker strip.",
    "WealthBenchmark — compare net worth to peers.",
    "WealthDNA — personalized financial personality.",
    "WealthTwinHero — twin summary card on dashboard.",
])

H2("10.2 AI Explainability & Chat Components")
para("Wealth Twin and AI Recommendations are supported by a suite of explainability and conversational components:")
mini_list([
    "AIDecisionLog — full trace of why a recommendation was made.",
    "ExplainableTooltip — hover-to-see-reason on any AI card.",
    "ELI5Tooltip — simple explanation for non-technical users.",
    "FinancialTwinChat — chat interface with the twin.",
    "WealthChat — general wealth Q&A assistant.",
    "SmartActionOrchestrator — executes multi-step actions autonomously.",
    "BehavioralEngine — detects patterns and adapts nudges.",
    "FinancialLiteracyCards — byte-sized education cards.",
    "AgenticActionCard — AI-initiated action suggestions with user approval.",
    "RecommendationCard — standardized recommendation UI with reasoning.",
])

H2("10.3 Payments Ecosystem Detail")
para("Payments is not just a transfer screen; it is a full ecosystem:")
mini_list([
    "PaymentHub — unified payments landing.",
    "UPIPaymentSimulator — simulate UPI collect/pay.",
    "QrScannerSimulator — scan-and-pay flow.",
    "VoicePayment — voice-command initiated payment.",
    "MPINInput — secure PIN entry.",
    "BillSplitter — split bills with contacts.",
    "GroupJar — shared money jar.",
    "CashbackWallet / CashbackPiggy — reward wallet and visual piggy.",
    "StreakTracker / StreakFire — payment streak gamification.",
    "SpinWheel — reward spin.",
    "AdReward — reward for watching educational ads.",
    "MoodMeter — emotion capture at payment time.",
    "PaymentRequests — request money from contacts.",
    "ReferralSection — invite-and-earn.",
    "RewardsDashboardCard — rewards summary.",
    "AnimatedTransactionToast — delightful confirmation.",
    "FloatingPayButton — quick-access pay FAB.",
])

H2("10.4 Protection Layer Components (Exhaustive)")
para("The protection layer contains many specialized guards:")
mini_list([
    "RiskMeter — composite risk score visualizer.",
    "TransactionGuardModal — blocks/warns on high-risk actions.",
    "CoolingVaultModal — forced cooling-off period.",
    "OTPSimulation — OTP retry/failure analysis.",
    "DuressMode / DuressPinSetup — secret PIN for coerced actions.",
    "PanicButton — instant account lock.",
    "LockdownOverlay — full-account lock screen.",
    "FamilySafeWord / FamilyApprovalGate — family-based safety checks.",
    "CoercedModeBanner — detected coercion warning.",
    "ScamCallerID — identify suspicious caller patterns.",
    "URLSafetyChecker — block phishing links.",
    "SecureCheckout — merchant checkout risk check.",
    "MoneyMuleGraph — visualize mule-account connections.",
    "ThreatIntel — live threat feed.",
    "BehavioralBiometrics — typing/mouse behavior anomaly.",
    "FraudDetectionEngine — core rule engine.",
    "ProtectionModal / ProtectionView — protection hub UI.",
    "RakshakInterventionChat — conversational intervention.",
    "CounterfactualPanel — 'what if I had allowed' analysis.",
    "StressTestSimulator — simulate attack scenarios.",
    "SecurityLog — user-facing protection event log.",
])

H2("10.5 Security Beast Components (Exhaustive)")
para("The 10-layer security dashboard is composed of:")
mini_list([
    "TpmAttestation — hardware trust check.",
    "EbpfMonitor — kernel-level activity monitor (simulated).",
    "HoneytokenManager — decoy credentials/assets.",
    "PasskeyAuth / WebAuthn — passwordless login.",
    "PostQuantumCrypto — ML-KEM key exchange demo.",
    "BehavioralBiometrics — behavior-based identity.",
    "DecentralizedId — self-sovereign identity concepts.",
    "TransactionTrap — honeypot transaction code.",
    "SecureEnclaveCheck — secure enclave simulation.",
    "BlockchainAudit — immutable event log.",
    "AntiScamShield — scam protection dashboard.",
    "GhostMode — invisible/hidden account mode.",
    "DeadMansSwitch — inheritance trigger.",
    "DecoyAccountView — fake account under duress.",
    "SecurityScoreDashboard — overall security score.",
    "VoiceCommandBar — voice-controlled security actions.",
    "DeviceFingerprintPanel — device trust details.",
    "DuressTrigger — manual duress activation.",
])

H2("10.6 Architecture & System Views")
para("The app also includes meta-views for judges and architecture documentation:")
mini_list([
    "FeaturesUniverse — galaxy/map view of all 40+ features.",
    "SystemArchitecture — interactive architecture diagram.",
    "PerformanceMetrics — frontend/backend performance stats.",
])

H2("10.7 Profile, Reports & Miscellaneous")
para("Other capabilities that complete the product:")
mini_list([
    "ProfileSettings — user profile, preferences, security settings.",
    "ReportGeneratorModal + FinancialReport — generate PDF/HTML financial reports.",
    "ParametricInsurance — weather/event-triggered insurance demo.",
    "GigIncomeSmoother — income smoothing for freelancers.",
    "LocationVerifier — location-based risk/safety check.",
    "InvestmentQuiz — financial literacy quiz.",
    "AddSalaryModal — salary/income input.",
    "PSB Components (psb/) — PSB-specific welcome banner, schemes card, quick pay, recent transactions, security health widget, accessible footer.",
    "Salary module — salary tracking and analysis.",
])

H2("10.8 Innovation Lab — Complete Component List")
para("The Innovation Lab contains 30+ experimental prototypes. Every component is listed below to ensure nothing is omitted:")
mini_list([
    "AIFutureTwin — AI-generated future self twin.",
    "AIInsightsAggregator — consolidated AI insight feed.",
    "AutonomousAgent — self-directed wealth agent.",
    "AutoInstrumentGenerator — auto-create financial instruments.",
    "BhavishyaEngine — predictive life-cycle engine (flagship).",
    "ChakraBalance — financial wellness balance visualization.",
    "CollectiveImmuneSystem — crowd-sourced fraud immunity.",
    "CommunityDNA — community financial behavior patterns.",
    "CrisisPredictor — predict and prepare for crises.",
    "DigitalInheritance — digital asset inheritance planning.",
    "DreamVisualizer — visualize financial dreams.",
    "EmotionalHeatmap — emotional spending map.",
    "EmotionalResonance — align recommendations with emotional state.",
    "FestivalAwareEngine — festival-based spending/investing nudges.",
    "FinancialDNAHelix — DNA-style financial identity.",
    "FutureSelfSimulator — see your future financial self.",
    "GenerationalWealth — inter-generational wealth planning.",
    "InnovationLabView — main lab shell.",
    "InnovationOverview — lab landing page.",
    "LifeEventPredictor — predict major life events.",
    "LifeShockSimulator — simulate financial shocks.",
    "MacroShockSimulator — simulate macro-economic shocks.",
    "MarketIntelligence — advanced market analysis.",
    "MonteCarloSimulator — probabilistic future simulator.",
    "NeuralNetworkViz — visualize recommendation neural net.",
    "NeuroFrictionWidget — reduce cognitive friction.",
    "ParticleConstellation — artistic data visualization.",
    "PreparednessScore — emergency preparedness score.",
    "ProsperityScore — holistic prosperity indicator.",
    "QuantumLock — quantum-resistant lock demo.",
    "SovereignVault — sovereign data vault concept.",
    "TemporalWealth — time-based wealth view.",
    "TimeMachine — go back/forward in financial time.",
    "WealthWeather — market mood as weather.",
])

H2("10.9 Backend Algorithms & Services")
para("The backend contains specialized algorithms that power the frontend:")
mini_list([
    "gstinValidator, itcRiskScanner, missingITCRecovery — GST compliance.",
    "msmeCreditScore.js — MSME credit scoring algorithm.",
    "shellCompanyDetector.js — detect suspicious business entities.",
    "taxOptimizer.js, taxRateErrorDetector.js — tax optimization and validation.",
    "ai-provider.js, aiRouter — multi-provider LLM routing.",
    "anomaly-detector.js — behavioral anomaly detection.",
    "dataIngestion.js, market-data.js — real-world data ingestion.",
    "scenario-engine.js, screener-service.js, dcf-engine.js, chart-engine.js — analytics engines.",
    "websocket.js — live fraud alert broadcast.",
    "middleware/timingCheck.js — rushed-action detection.",
    "middleware/auditLogger.js — immutable audit logging.",
])

# ===================== WHY / HOW / BENEFIT =====================
H1("11. Why, How, and Benefit Summary")
H2("11.1 Why This Tool Was Built")
table(
    ["Step", "Journey"],
    [
        ["Problem", "Digital wealth users face fraud, impulsive actions, generic advice, and thin-file credit exclusion."],
        ["Responsibility", "Banks must grow customer wealth AND protect it during every critical action."],
        ["Regulation", "DPDP Act 2023, RBI digital lending guidelines, and customer demand for explainable AI."],
        ["Solution", "SecureWealth Twin — a unified AI financial twin with mandatory cyber-fraud protection."],
        ["Outcome", "Wealth created responsibly, protected proportionally, and explained transparently."],
    ],
    widths=[1.0, 5.8],
)
H2("11.2 How It Was Built")
para("Frontend and backend are built as modular, feature-folder architectures. The frontend uses React 19 with TypeScript, Vite, Tailwind CSS v4, Framer Motion, Recharts, Lucide icons, and Zustand for state. The backend is Node.js + Express with SQLite (better-sqlite3) and a one-line adapter for PostgreSQL.")
para("Key implementation highlights:")
bullet("Protection layer: rule-based scoring in fraudDetectionService.ts plus server timingCheck.js.")
bullet("AI: multi-provider orchestrator with circuit breaker, caching, quota, and fallback.")
bullet("Data: real-world ingestion of RBI macro, NSE/BSE stocks, commodities, forex via Yahoo Finance/Alpha Vantage.")
bullet("Credit: deterministic scoring engine with SHAP-style additive decomposition.")
bullet("Deployment: Docker + Render + Surge + GitHub Actions CI/CD.")
H2("11.3 Key Benefits")
table(
    ["Stakeholder", "Benefit"],
    [
        ["Retail Customers", "Personalized guidance, goal tracking, fraud protection, accessible UI."],
        ["MSMEs", "Cash-flow visibility, surplus-fund advisor, collateral-free credit scoring."],
        ["Families", "Shared goals, allowances, transparent AI reasoning."],
        ["First-Time Investors", "Guided journey, gamified learning, simple nudges."],
        ["Bank", "Engagement, cross-sell, retention, safer digital wealth adoption, auditability."],
        ["Regulators", "Consent-first data, explainable AI, audit logs, no misleading claims."],
    ],
    widths=[2.0, 4.8],
)

# ===================== PYTHON SCRIPTS =====================
H1("12. Python Scripts Used")
para("Three Python utilities generate supporting Word documents for the demo team:")
H3("11.1 build_cuecard.py")
para("One-page printable 10-minute demo cue card.")
add_code_block("build_cuecard.py", lines=40)
H3("11.2 build_playbook.py")
para("Full team demo playbook with pitch, architecture, protection deep-dive, run-of-show, and Q&A.")
add_code_block("build_playbook.py", lines=40)
H3("11.3 build_tooltable.py")
para("Landscape demo tool table mapping each component to judging metrics.")
add_code_block("build_tooltable.py", lines=45)

# ===================== COMPLIANCE & EVALUATION =====================
H1("13. Compliance & Evaluation Mapping")
H2("13.1 Compliance Checklist")
table(
    ["Requirement", "Implementation"],
    [
        ["Customer Data Privacy & Consent", "ConsentModal, PrivacyCenter, data provenance ledger in CreditBridge AI."],
        ["Secure Data Handling", "JWT auth, bcrypt hashing, Helmet CSP, audit logs, no plain-text secrets."],
        ["Transparent & Explainable AI", "ExplainableTooltip, AIDecisionLog, CreditBridge factor breakdown, reason codes."],
        ["No Misleading Claims", "Demo-only badges, disclaimers, no guaranteed returns."],
        ["Basic Financial Regulations", "KYCModal before investment actions, regulatory disclaimers."],
    ],
    widths=[2.4, 4.6],
)
H2("13.2 Hackathon Evaluation Metrics Mapping")
table(
    ["Metric", "Features That Satisfy It"],
    [
        ["Quality of Wealth Intelligence", "Wealth Twin, AI Recommendations, Goals, Portfolio, Market, Forecast, BHAVISHYA."],
        ["Effectiveness of Wealth Protection", "Protection hub, Risk Meter, OTP simulation, Duress mode, Cooling Vault, Security Beast."],
        ["Simplicity & UX", "Dashboard, Command Palette, Senior/Kids mode, Toasts, Accessibility, Demo Mode."],
        ["Use of AI / Data Analytics", "Multi-provider AI orchestrator, spending persona, BHAVISHYA, CreditBridge scoring, real-world data ingestion."],
        ["Innovation", "Innovation Lab, CreditBridge AI, Social Collateral Loan, Behavioral biometrics, Quantum lock, Sovereign vault."],
        ["Technical Design & Architecture", "Feature-folder React, Node microservices, SQLite/PG adapter, Docker, CI/CD."],
        ["Scalability & Practicality", "Docker container, Render/Surge deployment, AI fallback, offline queue, WebSocket alerts."],
    ],
    widths=[2.4, 4.6],
)

# ===================== DEMO FLOW =====================
H1("14. Suggested 10-Minute Demo Flow")
table(
    ["Time", "Block", "What to Show", "Key Message"],
    [
        ["0:00-0:45", "Hook + Architecture", "Live app + architecture slide", "We grow wealth AND protect it — dual mandate."],
        ["0:45-3:00", "Wealth Twin", "Dashboard → AA → Assets → Goals → AI Recommendations", "The twin builds a full picture and gives advice WITH a reason."],
        ["3:00-6:00", "Protection Layer", "High-value action → device/timing/amount/OTP → Risk Score → Warn/Block + Duress PIN", "Every critical action passes a proportional protection layer."],
        ["6:00-7:30", "Responsible AI", "Explainable tooltip → Consent → KYC → Audit log", "Consent-first, explainable, auditable."],
        ["7:30-9:00", "Innovation / CreditBridge", "BHAVISHYA scenario or CreditBridge AI scoring", "40+ modules; flagship CreditBridge for thin-file credit."],
        ["9:00-10:00", "Scale + Close", "Tech stack + deployment + team roles", "React + Node + Docker + AA + mock CBS; ready to scale."],
    ],
    widths=[0.9, 1.4, 3.2, 2.5],
)

# ===================== APPENDIX: EXHAUSTIVE CHECKLIST =====================
H1("Appendix A: Exhaustive Component Checklist")
para("This checklist is derived directly from the project structure. If a component exists in the codebase, it is listed here so the PPT can reference it if needed.")

H2("A.1 Frontend Features (client/src/features/)")

def dir_block(title, files):
    H3(title)
    mini_list(files)

dir_block("aa", ["AACallbackHandler.tsx", "AAFetchAnimation.tsx", "AccountAggregatorFull.tsx"])
dir_block("accessibility", ["AccessibilitySettings.tsx"])
dir_block("admin", [
    "AdminActivityTab.tsx", "AdminDashboard.tsx", "AdminLoginArchitecture.tsx", "AlertHistoryTab.tsx",
    "AlertToast.tsx", "DemoTour.tsx", "FraudCaseDetail.tsx", "FraudCaseExplorer.tsx",
    "FraudCorrelationPanel.tsx", "FraudExportPanel.tsx", "FraudHeatmap.tsx", "FraudIntelligenceCenter.tsx",
    "FraudMapView.tsx", "FraudRiskExplainer.tsx", "FraudRulesPanel.tsx", "FraudTimeline.tsx",
    "FraudTraceGraph.tsx", "SystemHealthTab.tsx", "useFraudCases.ts", "fraudDataGenerator.ts",
    "fraudTypes.ts", "permissions.ts", "fraudService.ts"
])
dir_block("ai", [
    "AIDecisionLog.tsx", "AIRecommendationsView.tsx", "AgenticActionCard.tsx", "BehavioralEngine.tsx",
    "ELI5Tooltip.tsx", "ExplainableTooltip.tsx", "FinancialLiteracyCards.tsx", "FinancialTwinChat.tsx",
    "RecommendationCard.tsx", "SmartActionOrchestrator.tsx", "WealthChat.tsx", "WealthTwinView.tsx",
    "ExplainablePanel.tsx", "GoalsTab.tsx", "OverviewTab.tsx", "RebalanceTab.tsx", "RetirementTab.tsx",
    "TaxTab.tsx", "TwinTabs.tsx", "WealthTwinContext.tsx", "WhatIfTab.tsx", "useWealthTwinData.ts", "utils.ts"
])
dir_block("architecture", ["FeaturesUniverse.tsx", "PerformanceMetrics.tsx", "SystemArchitecture.tsx"])
dir_block("assets", ["AccountAggregatorWidget.tsx", "LinkAccountModal.tsx", "ManualAssetForm.tsx", "PhysicalAssetIntelligence.tsx", "VisionAppraisalModal.tsx"])
dir_block("auth", ["BiometricAuth.tsx", "CreateAccountModal.tsx", "CreateAccountModal.test.tsx", "FaceLoginModal.tsx", "LoginPage.tsx", "LoginPortal.tsx"])
dir_block("banking", ["AccountStatement.tsx", "AuditLog.tsx", "LoanCenter.tsx", "RecurringPayments.tsx"])
dir_block("bills", ["BillCalendar.tsx"])
dir_block("business", ["BusinessMode.tsx", "CashFlowTimeline.tsx", "SMEDashboard.tsx", "SurplusFundAdvisor.tsx", "WorkingCapitalHealth.tsx"])
dir_block("calculators", ["CalculatorsView.tsx", "RentVsBuyCalculator.tsx"])
dir_block("challenges", ["ChallengesView.tsx"])
dir_block("compliance", ["ComplianceBadges.tsx", "ComplianceBar.tsx", "ConsentModal.tsx", "DemoControls.tsx", "KYCModal.tsx", "PrivacyCenter.tsx"])
dir_block("credit", ["CreditBridgeAI.tsx", "CreditHealth.tsx", "DemoCreditCard.tsx", "creditBridgeEngine.ts", "ethicsEngine.ts", "lenderMarketplace.ts"])
dir_block("dashboard", [
    "AdaptiveInsight.tsx", "BehavioralNudges.tsx", "DashboardView.tsx", "DashboardWidget.tsx",
    "DeviceStatusCard.tsx", "FinancialPulse.tsx", "FinancialWeather.tsx", "KYCStatusCard.tsx",
    "MarketIntelligenceHero.tsx", "MonthlyNarrative.tsx", "NBAInsights.tsx", "NetWorthCard.tsx",
    "NotificationCenter.tsx", "PhysicalAssetsPromo.tsx", "PredictiveShieldBadge.tsx", "QuickActions.tsx",
    "StockTicker.tsx", "WealthBenchmark.tsx", "WealthDNA.tsx", "WealthTwinHero.tsx"
])
dir_block("demo", [
    "BlockchainAudit.tsx", "CinematicIntro.tsx", "CursorSpotlight.tsx", "DemoAssistant.tsx",
    "DemoMode.tsx", "DemoPersonas.tsx", "DemoShowcase.tsx", "FeatureUniverse.tsx",
    "FraudRadar.tsx", "JudgeTour.tsx", "NotificationDemo.tsx"
])
dir_block("family", ["FamilyDashboard.tsx"])
dir_block("forecast", ["ForecastView.tsx", "ScenarioSimulator.tsx", "WhatIfSimulator.tsx"])
dir_block("gamification", ["BadgeStreak.tsx", "FantasyLeague.tsx", "FantasyPortfolio.tsx"])
dir_block("goals", [
    "AddGoalModal.tsx", "BoostCard.tsx", "BoostsManager.tsx", "DynamicCompass.tsx",
    "GoalCelebration.tsx", "GoalConflictIntelligence.tsx", "GoalConflictModal.tsx", "GoalTracker.tsx"
])
dir_block("gold", ["DigitalGold.tsx"])
dir_block("income", ["GigIncomeSmoother.tsx"])
dir_block("innovation", [
    "AIFutureTwin.tsx", "AIInsightsAggregator.tsx", "AutoInstrumentGenerator.tsx", "AutonomousAgent.tsx",
    "BhavishyaEngine.tsx", "ChakraBalance.tsx", "CollectiveImmuneSystem.tsx", "CommunityDNA.tsx",
    "CrisisPredictor.tsx", "DigitalInheritance.tsx", "DreamVisualizer.tsx", "EmotionalHeatmap.tsx",
    "EmotionalResonance.tsx", "FestivalAwareEngine.tsx", "FinancialDNAHelix.tsx", "FutureSelfSimulator.tsx",
    "GenerationalWealth.tsx", "InnovationLabView.tsx", "InnovationOverview.tsx", "LifeEventPredictor.tsx",
    "LifeShockSimulator.tsx", "MacroShockSimulator.tsx", "MarketIntelligence.tsx", "MonteCarloSimulator.tsx",
    "NeuralNetworkViz.tsx", "NeuroFrictionWidget.tsx", "ParticleConstellation.tsx", "PreparednessScore.tsx",
    "ProsperityScore.tsx", "QuantumLock.tsx", "SovereignVault.tsx", "TemporalWealth.tsx", "TimeMachine.tsx",
    "WealthWeather.tsx"
])
dir_block("insurance", ["ParametricInsurance.tsx"])
dir_block("kids", ["KidsMode.tsx"])
dir_block("loans", ["LoanImpactSimulator.tsx", "LoanResearchShowcase.tsx", "LoansHub.tsx", "SocialCollateralLoan.tsx"])
dir_block("location", ["LocationVerifier.tsx"])
dir_block("market", ["GlobalMacroRadar.tsx", "MacroSignalTower.tsx", "MarketNewsFeed.tsx", "MarketStrategist.tsx", "MarketView.tsx", "SmartTriggers.tsx", "useMacroFeed.ts"])
dir_block("msme", ["MSMEcreditbridgeView.tsx", "msmeScoreEngine.ts"])
dir_block("nri", ["NRIMode.tsx"])
dir_block("onboarding", ["OnboardingWizard.tsx"])
dir_block("payments", [
    "AdReward.tsx", "AnimatedTransactionToast.tsx", "BillSplitter.tsx", "CashbackPiggy.tsx",
    "CashbackWallet.tsx", "FloatingPayButton.tsx", "GroupJar.tsx", "MPINInput.tsx", "MoodMeter.tsx",
    "PaymentHub.tsx", "PaymentRequests.tsx", "PaymentsPage.tsx", "QrScannerSimulator.tsx",
    "ReferralSection.tsx", "RewardsDashboardCard.tsx", "SpinWheel.tsx", "StreakFire.tsx",
    "StreakTracker.tsx", "UPIPaymentSimulator.tsx", "VoicePayment.tsx"
])
dir_block("pitch", ["PitchDeckView.tsx", "PitchMode.tsx"])
dir_block("portfolio", ["ESGScore.tsx", "PortfolioView.tsx"])
dir_block("privacy", ["PrivacyAuditPanel.tsx", "PrivacyView.tsx"])
dir_block("profile", ["ProfileSettings.tsx"])
dir_block("protection", [
    "RakshakInterventionChat.tsx", "BehavioralBiometrics.tsx", "CoercedModeBanner.tsx",
    "CoolingVaultModal.tsx", "CounterfactualPanel.tsx", "DuressMode.tsx", "DuressPinSetup.tsx",
    "FamilyApprovalGate.tsx", "FamilySafeWord.tsx", "FraudDetectionEngine.tsx", "LockdownOverlay.tsx",
    "MoneyMuleGraph.tsx", "OTPSimulation.tsx", "OTPSimulation.test.tsx", "PanicButton.tsx",
    "PaymentGuard.tsx", "ProtectionModal.tsx", "ProtectionView.tsx", "RiskMeter.tsx",
    "ScamCallerID.tsx", "SecureCheckout.tsx", "SecurityLog.tsx", "StressTestSimulator.tsx",
    "ThreatIntel.tsx", "TransactionGuardModal.tsx", "URLSafetyChecker.tsx"
])
dir_block("psb", [
    "AccessibleFooter.tsx", "PSBLogo.tsx", "PSBSchemesCard.tsx", "QuickPayCard.tsx",
    "RecentTransactionsTable.tsx", "SecurityHealthWidget.tsx", "WelcomeBanner.tsx"
])
dir_block("quiz", ["InvestmentQuiz.tsx"])
dir_block("report", ["FinancialReport.tsx", "ReportGeneratorModal.tsx"])
dir_block("salary", ["AddSalaryModal.tsx"])
dir_block("security", [
    "AntiScamShield.tsx", "BlockchainAudit.tsx", "DeadMansSwitch.tsx", "DecentralizedId.tsx",
    "DecoyAccountView.tsx", "DeviceFingerprintPanel.tsx", "DuressTrigger.tsx", "EbpfMonitor.tsx",
    "GhostMode.tsx", "HoneytokenManager.tsx", "PasskeyAuth.tsx", "PostQuantumCrypto.tsx",
    "SecureEnclaveCheck.tsx", "SecurityBeastView.tsx", "SecurityScoreDashboard.tsx",
    "TpmAttestation.tsx", "TransactionTrap.tsx", "VoiceCommandBar.tsx"
])
dir_block("senior", ["SeniorMode.tsx"])
dir_block("subscriptions", ["AdminAgent.tsx", "CultFitTracker.tsx", "SubscriptionTracker.tsx"])
dir_block("tax", ["OldVsNewRegime.tsx", "Section80CTracker.tsx", "TaxCalculator.tsx", "TaxDeadlineCalendar.tsx", "TaxView.tsx"])
dir_block("transactions", [
    "AICategorization.tsx", "EmotionCheckin.tsx", "ScanReceipt.tsx", "SmartDuplicateDetection.tsx",
    "TransactionComparison.tsx", "TransactionDetailModal.tsx", "TransactionTagger.tsx", "TransactionsView.tsx"
])
dir_block("values", ["ValuesAlignment.tsx"])

H2("A.2 In-App Views (AuthenticatedApp.tsx)")
mini_list([
    "dashboard", "wealth-twin", "ai-recommendations", "goals", "portfolio", "family", "assets",
    "market", "forecast", "protection", "privacy", "tax", "calculators", "transactions", "features",
    "architecture", "bills", "credit-health", "creditbridge-ai", "notification-demo", "digital-gold",
    "challenges", "kids-mode", "subscriptions", "accessibility", "nri-mode", "business-mode",
    "values-alignment", "fantasy-league", "boosts", "security-beast", "bhavishya", "innovation-lab",
    "pitch-deck", "payments", "loan-center", "msme-creditbridge", "loans-hub", "loan-research",
    "loan-impact", "social-collateral-loan", "recurring-payments", "account-statement", "audit-log", "profile"
])

H2("A.3 Backend API Routes")
mini_list([
    "admin.js", "ai.js", "analytics.js", "auth.js", "banking.js", "business.js", "charts.js",
    "documents.js", "export.js", "extract.js", "financial-model.js", "fraud.js", "gallery.js",
    "gst.js", "kyc.js", "market-data.js", "msme.js", "nlp-query.js", "realData.js", "scenarios.js",
    "screener.js", "tax.js"
])

H2("A.4 Backend Algorithms")
mini_list([
    "gstinValidator.js", "itcRiskScanner.js", "missingITCRecovery.js", "msmeCreditScore.js",
    "shellCompanyDetector.js", "taxOptimizer.js", "taxRateErrorDetector.js"
])

H2("A.5 Major Shared Services")
mini_list([
    "aiOrchestrator.ts", "fraudDetectionService.ts", "fingerprintService.ts", "useRecommendationEngine.ts",
    "useProtectionEngine.ts", "auditLogger.ts", "securityLogger.ts", "offlineQueue.ts",
    "twinService.ts", "scenarioEngine.ts", "nbaService.ts", "voiceService.ts", "passkeyService.ts",
    "postQuantumService.ts", "duressService.ts", "blockchainService.ts", "notificationService.ts",
    "cashbackEngine.ts", "leagueService.ts", "streakService.ts", "subscriptionPredictor.ts",
    "privacyAuditService.ts", "attestationService.ts", "behavioralBiometricsService.ts",
    "browserThreatService.ts", "urlSafetyService.ts", "sentimentService.ts", "newsService.ts"
])

H2("A.6 Python / Doc-Generation Scripts")
mini_list([
    "build_cuecard.py — 10-minute demo cue card",
    "build_playbook.py — full team demo playbook",
    "build_tooltable.py — Hinglish demo tool table",
    "build_comprehensive_doc.py — this comprehensive PPT source document"
])

# Save
doc.save(OUTPUT)
print(f"Saved {OUTPUT} ({os.path.getsize(OUTPUT)} bytes)")
